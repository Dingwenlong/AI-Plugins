#!/usr/bin/env python3
from __future__ import annotations

import argparse
import datetime as dt
import difflib
import getpass
import hashlib
import json
import posixpath
import re
import sys
import unicodedata
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from decimal import Decimal
from pathlib import Path
from typing import Any, Iterable
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from jsonschema import Draft202012Validator
from openpyxl import load_workbook
from openpyxl.worksheet.worksheet import Worksheet

from reference_support import (
    format_reference_locator,
    find_api_codes,
    find_external_keywords,
    find_framework_keywords,
    find_table_keys,
    normalize_match_key,
    REFERENCE_INDEX_FILENAMES,
    reference_catalog_path,
    reference_index_path,
    reference_roots_for_read,
    slugify as reference_slugify,
)
from chain_workspace import update_chain_status
from runtime import (
    API_SPEC_SCHEMA_VERSION,
    BATCH_SCHEMA_VERSION,
    STATE_SCHEMA_VERSION,
    ExecutionContext,
    append_progress,
    build_api_id,
    default_batch_file,
    configure_stdio,
    dump_json,
    extract_function_code,
    extract_version_token_from_tsd_path,
    load_json,
    load_batch_file,
    load_schema,
    normalize_author_name,
    normalize_persisted_path,
    now_iso,
    resolve_context_root,
    resolve_agent_dir,
    resolve_docx_path,
    resolve_project_root,
    save_batch_file,
    upsert_batch_item,
)


UNSUPPORTED_LEGACY_FLAGS = {
    "--progress-db",
    "--artifact-profile",
    "--baseline-run-id",
    "--manual-api-id",
    "--attempt",
    "--project-profile-id",
    "--rule-pack-id",
    "--rule-pack-version",
    "--git-profile-id",
    "--expert-id",
    "--workflow-id",
    "--skill-id",
    "--no-track",
}

SIMPLE_DECIMAL_RE = re.compile(r"^\d+(?:\.\d+)?$")
ROW_KEY_RE = re.compile(r"^\d+(?:\.\d+)*$")
FORMULA_INCREMENT_RE = re.compile(r"^=\s*[A-Z]+\d+\s*\+\s*(?P<increment>\d+(?:\.\d+)?)\s*$", re.IGNORECASE)
BACKEND_API_LINE_RE = re.compile(
    r"^(?:\d+\.)?\s*(?P<system>[A-Za-z0-9_./ -]+?)\s*[-=]>\s*(?P<target>.+)$",
    re.IGNORECASE,
)
BACKEND_API_CONTINUATION_RE = re.compile(r"^(?:\d+\.)?\s*[-=]>\s*(?P<target>.+)$", re.IGNORECASE)
UNCERTAIN_DEPENDENCY_RE = re.compile(r"(todo|待確認|待确定|是否還需要|是否还需要|未找到|麻煩提供|请提供)", re.IGNORECASE)
BUSINESS_DEPENDENCY_RE = re.compile(
    r"\b(?P<system>IRIS|CommonFunc|CommonUtil|Backend|DB|Redis|JWT|Header)\s*(?:->|[-=]>)\s*(?P<target>[^\n;；]+)",
    re.IGNORECASE,
)
METHOD_CALL_RE = re.compile(r"\b(?P<owner>CommonFunc|CommonUtil|Backend)\s*\.\s*(?P<method>[A-Za-z_][A-Za-z0-9_]*)\s*\(", re.IGNORECASE)
IRIS_CODE_RE = re.compile(r"\b(?P<code>[A-Z]{2}\d{4})\b")
MAIL_PUSH_RE = re.compile(r"\b(?P<name>SendToMonitorMail_Push)\s*\(", re.IGNORECASE)
REFERENCE_CATEGORIES = ("db_schema", "external_api")
REFERENCE_CATEGORY_LABELS = {
    "db_schema": "DB Schema",
    "external_api": "External API",
}
REFERENCE_HINT_LIMIT = 3
PROJECT_HARD_CONSTRAINTS_FILENAME = "project-hard-constraints.json"
DB_SOURCE_LABEL_PRIORITY = {
    "db-tableschema": 40,
    "db-relation": 30,
    "db-dictionary": 20,
    "db-traceability": 10,
}
DRAWING_RELATIONSHIP_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/drawing"
PACKAGE_REL_NS = "{http://schemas.openxmlformats.org/package/2006/relationships}"
SPREADSHEET_DRAWING_NS = "{http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing}"


class SkillError(RuntimeError):
    def __init__(self, message: str, *, status: str = "error") -> None:
        super().__init__(message)
        self.status = status


class ZhArgumentParser(argparse.ArgumentParser):
    def __init__(self, *args: object, **kwargs: object) -> None:
        kwargs.setdefault("add_help", False)
        super().__init__(*args, **kwargs)
        self._positionals.title = "位置参数"
        self._optionals.title = "可选参数"

    def format_usage(self) -> str:
        return super().format_usage().replace("usage:", "用法：", 1)

    def format_help(self) -> str:
        text = super().format_help()
        return (
            text.replace("usage:", "用法：", 1)
            .replace("positional arguments:", "位置参数：")
            .replace("optional arguments:", "可选参数：")
            .replace("options:", "可选参数：")
            .replace("show this help message and exit", "显示此帮助并退出")
        )

    def error(self, message: str) -> None:
        self.print_usage(sys.stderr)
        self.exit(2, f"{self.prog}: 错误：{message}\n")


@dataclass(frozen=True)
class ApiEntry:
    category: str
    name: str
    description: str = ""


@dataclass(frozen=True)
class SheetMatch:
    api: ApiEntry
    workbook_path: Path
    sheet_name: str


@dataclass(frozen=True)
class SheetVisualSignal:
    sheet_name: str
    image_count: int
    drawing_anchor_count: int


def reject_unsupported_legacy_flags(argv: list[str]) -> None:
    matches = []
    for token in argv:
        name = token.split("=", 1)[0]
        if name in UNSUPPORTED_LEGACY_FLAGS:
            matches.append(name)
    if matches:
        rendered = ", ".join(sorted(set(matches)))
        raise SystemExit(
            f"api-spec-writer does not support legacy orchestration flags: {rendered}. "
            "Use the legacy tracker/audit/review workflow for tracker/audit/review features."
        )


def parse_args() -> argparse.Namespace:
    reject_unsupported_legacy_flags(sys.argv[1:])
    parser = ZhArgumentParser(description="初始化或恢复 api-spec-writer 执行面，并逐 API 生成规格交接产物。")
    parser.add_argument("-h", "--help", action="help", help="显示此帮助并退出")
    parser.add_argument("docx_ref", nargs="?", default=None, help="兼容输入：DOCX 路径、精确文件名，或 `.agent/TSD` 下的唯一部分文件名。默认仍要求先有 development-handoff.json。")
    parser.add_argument("--project-root", required=True)
    parser.add_argument("--agent-dir", default=".agent")
    parser.add_argument("--agent-root", default=None, help="集中式 .agent 根目录；优先级高于 --agent-dir。")
    parser.add_argument("--workspace-root", default=None, help="集中式工作区根目录，例如 D:\\Repo\\Project。")
    parser.add_argument("--workspace-key", default=None, help="local-workspaces.json 中的工作区 key。")
    parser.add_argument("--rules-root", help="专案规则库根目录；优先级高于环境变量与 workspace 配置。")
    parser.add_argument("--context-root", default=None)
    parser.add_argument("--design-handoff", default=None, help="梳理技能生成的 development-handoff.json；可跳过第 01 步。")
    parser.add_argument("--allow-legacy-input", action="store_true", help="兼容旧流程：允许没有梳理 handoff 时直接消费 docx_ref / execution-batch。默认不启用。")
    parser.add_argument("--api-id", default=None)
    parser.add_argument("--function-code", default=None)
    parser.add_argument("--new-author", default=None)
    return parser.parse_args()


def normalize_batch_docx_ref(project_root: Path, docx_path: Path) -> str:
    return normalize_persisted_path(docx_path, project_root=project_root) or docx_path.name


def resolve_design_handoff_path(agent_dir: Path, explicit: str | None, function_code: str | None) -> Path | None:
    if explicit:
        path = Path(explicit).expanduser()
        if not path.is_absolute():
            path = (Path.cwd() / path).resolve()
        if not path.exists():
            raise SkillError(f"指定的 development-handoff.json 不存在：{path.as_posix()}", status="blocked")
        return path
    if not function_code:
        return None
    path = agent_dir / "functions" / function_code / "handoff" / "development-handoff.json"
    return path if path.exists() else None


def load_design_handoff(path: Path | None) -> dict[str, Any] | None:
    if path is None:
        return None
    payload = load_json(path)
    if not isinstance(payload, dict):
        raise SkillError(f"development-handoff.json must be a JSON object: {path.as_posix()}", status="blocked")
    return payload


def handoff_blocks_spec(payload: dict[str, Any]) -> str | None:
    readiness = payload.get("readiness") if isinstance(payload.get("readiness"), dict) else {}
    status = clean_text(payload.get("status") or readiness.get("status") or readiness.get("designStatus"))
    if status in {"blocked", "not_ready", "needs_design"}:
        return clean_text(payload.get("blockReason") or readiness.get("blockReason")) or "功能设计尚未达到可进入开发门槛"
    ready = readiness.get("developmentReady")
    if ready is False:
        return clean_text(readiness.get("blockReason")) or "功能设计尚未达到可进入开发门槛"
    return None


def resolve_handoff_relative_path(agent_dir: Path, value: object) -> Path | None:
    text = clean_text(value)
    if not text:
        return None
    path = Path(text).expanduser()
    if not path.is_absolute():
        path = agent_dir / path
    return path.resolve() if path.exists() else None


def select_handoff_file(agent_dir: Path, payload: dict[str, Any], *kinds: str) -> Path | None:
    wanted = {kind.casefold() for kind in kinds}
    for item in list(payload.get("sourceFiles") or []):
        if not isinstance(item, dict):
            continue
        kind = clean_text(item.get("kind") or item.get("type") or item.get("category")).casefold()
        if kind not in wanted:
            continue
        for key in ("copiedRelativePath", "targetRelativePath", "relativePath", "path"):
            path = resolve_handoff_relative_path(agent_dir, item.get(key))
            if path is not None:
                return path
    return None


def infer_requested_function_code(
    *,
    requested_function_code: str | None,
    selected_item: dict[str, Any] | None,
    docx_ref: str | None,
) -> str | None:
    if requested_function_code:
        return clean_text(requested_function_code)
    if selected_item is not None:
        function_code = clean_text(selected_item.get("functionCode"))
        if function_code:
            return function_code
    if docx_ref:
        return extract_function_code(Path(docx_ref).name)
    return None


def missing_handoff_message(agent_dir: Path, function_code: str | None) -> str:
    if function_code:
        expected = agent_dir / "functions" / function_code / "handoff" / "development-handoff.json"
        return (
            "02 未找到梳理产物 development-handoff.json，不能直接回退到 01 或 legacy TSD。"
            f"请先运行 `专案需求接口设计梳理`，并执行 materialize_design_handoff.py 生成：{expected.as_posix()}"
        )
    return (
        "02 未找到可用的梳理产物，也无法确定功能编号。"
        "请先提供 --function-code，或先运行 `专案需求接口设计梳理` 生成 `.agent/functions/<functionCode>/handoff/development-handoff.json`。"
    )


def resolve_batch_target(
    project_root: Path,
    agent_dir: Path,
    context_root: Path,
    *,
    docx_ref: str | None,
    requested_function_code: str | None,
    design_handoff: str | None = None,
    allow_legacy_input: bool = False,
) -> tuple[Path, str | None, dict[str, Any], dict[str, Any] | None, Path]:
    batch_file = default_batch_file(context_root)
    batch_payload = load_batch_file(batch_file)
    selected_item: dict[str, Any] | None = None

    if requested_function_code:
        for item in list(batch_payload.get("items") or []):
            if clean_text(item.get("functionCode")) == clean_text(requested_function_code):
                selected_item = dict(item)
                break
    elif batch_payload.get("activeFunctionCode"):
        active = clean_text(batch_payload.get("activeFunctionCode"))
        for item in list(batch_payload.get("items") or []):
            if clean_text(item.get("functionCode")) == active:
                selected_item = dict(item)
                break
    elif len(list(batch_payload.get("items") or [])) == 1:
        selected_item = dict(list(batch_payload.get("items") or [])[0])

    inferred_function_code = infer_requested_function_code(
        requested_function_code=requested_function_code,
        selected_item=selected_item,
        docx_ref=docx_ref,
    )

    handoff_path = resolve_design_handoff_path(agent_dir, design_handoff, inferred_function_code)
    handoff_payload = load_design_handoff(handoff_path)
    if handoff_payload is not None:
        block_reason = handoff_blocks_spec(handoff_payload)
        if block_reason:
            raise SkillError(block_reason, status="blocked")
        function_code = requested_function_code or clean_text(handoff_payload.get("functionCode")) or None
        docx_path = select_handoff_file(agent_dir, handoff_payload, "tsd", "docx")
        if docx_path is None:
            raise SkillError(f"development-handoff.json 未提供可用 TSD 文件：{handoff_path}", status="blocked")
        if not function_code:
            function_code = extract_function_code(docx_path.name)
        if not function_code:
            raise SkillError(f"无法从 development-handoff.json 或 TSD 文件名提取功能编号：{docx_path.name}", status="blocked")
        return docx_path, function_code, batch_payload, selected_item, batch_file

    if not allow_legacy_input:
        raise SkillError(missing_handoff_message(agent_dir, inferred_function_code), status="blocked")

    if docx_ref:
        docx_path = resolve_docx_path(project_root, agent_dir, docx_ref)
        function_code = requested_function_code or extract_function_code(docx_path.name)
        if not function_code:
            raise SkillError(f"无法从 TSD 文件名提取功能编号：{docx_path.name}", status="blocked")
        return docx_path, function_code, batch_payload, selected_item, batch_file

    if selected_item is not None:
        function_code = clean_text(selected_item.get("functionCode")) or None
        docx_path = resolve_docx_path(project_root, agent_dir, clean_text(selected_item.get("docxRef")))
        if requested_function_code and function_code and function_code != requested_function_code:
            raise SkillError(f"--function-code 与 execution-batch.json 不一致：expected {function_code}, got {requested_function_code}")
        return docx_path, function_code, batch_payload, selected_item, batch_file

    docx_path = resolve_docx_path(project_root, agent_dir, "")
    function_code = requested_function_code or extract_function_code(docx_path.name)
    if not function_code:
        raise SkillError(f"无法从 TSD 文件名提取功能编号：{docx_path.name}", status="blocked")
    return docx_path, function_code, batch_payload, None, batch_file


def refresh_batch_pointer(context: ExecutionContext, *, preferred_function_code: str | None = None) -> None:
    payload = load_batch_file(context.batch_file)
    items = sorted(list(payload.get("items") or []), key=lambda item: (int(item.get("order") or 0), clean_text(item.get("functionCode"))))
    if not items:
        return

    def execution_status(function_code: str) -> str | None:
        path = context.context_root / function_code / "execution-state.json"
        if not path.exists():
            return None
        payload = load_json(path)
        return clean_text(payload.get("status")) or None

    active = clean_text(preferred_function_code) or clean_text(payload.get("activeFunctionCode"))
    if not active and items:
        active = clean_text(items[0].get("functionCode"))
    if active and execution_status(active) == "done":
        for item in items:
            candidate = clean_text(item.get("functionCode"))
            if candidate and execution_status(candidate) != "done":
                active = candidate
                break
    payload["activeFunctionCode"] = active or None
    save_batch_file(context.batch_file, payload, updated_by=context.new_author)


def normalize_token(value: str) -> str:
    text = unicodedata.normalize("NFKC", value or "")
    text = re.sub(r"\s+", "", text)
    return text.casefold()


def clean_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    lines = [segment.strip() for segment in text.split("\n")]
    return "\n".join(segment for segment in lines if segment)


def normalize_display_version(value: str | None) -> str | None:
    if not value:
        return None
    return value[1:] if value.lower().startswith("v") else value


def summarize_checklist(items: list[dict[str, Any]]) -> dict[str, int]:
    counts = {"total": len(items), "pending": 0, "in_progress": 0, "done": 0, "blocked": 0, "error": 0, "retired": 0}
    for item in items:
        status = item.get("status")
        if status in counts:
            counts[status] += 1
    return counts


def derive_execution_status(checklist_items: list[dict[str, Any]]) -> str:
    if any(item.get("status") in {"pending", "in_progress", "error"} for item in checklist_items):
        return "waiting_resume"
    if any(item.get("status") == "blocked" for item in checklist_items):
        return "blocked"
    return "done"


def default_code_status(spec_status: str) -> str:
    return "pending" if spec_status == "done" else "waiting_spec"


def default_code_phase(spec_status: str) -> str:
    return "pending" if spec_status == "done" else "waiting_spec"


def summarize_code_status(items: list[dict[str, Any]]) -> dict[str, int]:
    counts = {
        "total": len(items),
        "waiting_spec": 0,
        "pending": 0,
        "in_progress": 0,
        "tests_passed": 0,
        "tests_failed": 0,
        "blocked": 0,
        "error": 0,
    }
    for item in items:
        status = clean_text(item.get("codeStatus"))
        if status in counts:
            counts[status] += 1
    return counts


def build_code_artifacts_payload(context: ExecutionContext, api_id: str) -> dict[str, Any]:
    api_dir = context.paths.api_dir(api_id)
    return {
        "changePlan": normalize_persisted_path(api_dir / "change-plan.json", project_root=context.project_root) if (api_dir / "change-plan.json").exists() else None,
        "implementationReport": normalize_persisted_path(api_dir / "implementation-report.md", project_root=context.project_root) if (api_dir / "implementation-report.md").exists() else None,
        "diagnosisReport": normalize_persisted_path(api_dir / "diagnosis-report.json", project_root=context.project_root) if (api_dir / "diagnosis-report.json").exists() else None,
    }


def derive_code_execution_projection(
    checklist_items: list[dict[str, Any]],
    existing_payload: dict[str, Any],
) -> tuple[str, str, dict[str, int]]:
    summary = summarize_code_status(checklist_items)
    code_reset = any(bool(item.get("codeReset")) for item in checklist_items)
    previous_status = clean_text(existing_payload.get("codeStatus"))
    previous_phase = clean_text(existing_payload.get("codePhase"))
    ready_items = [item for item in checklist_items if item.get("status") == "done"]

    for item in ready_items:
        if clean_text(item.get("codeStatus")) == "in_progress":
            return "running", clean_text(item.get("codePhase")) or "in_progress", summary

    if ready_items and all(clean_text(item.get("codeStatus")) == "tests_passed" for item in ready_items):
        return "done", previous_phase or "validated", summary

    for item in ready_items:
        item_status = clean_text(item.get("codeStatus"))
        if item_status in {"pending", "tests_failed", "error"}:
            if code_reset:
                return "waiting_resume", clean_text(item.get("codePhase")) or "pending", summary
            if previous_status in {"", "waiting_spec"}:
                return previous_status or "waiting_spec", previous_phase or "waiting_spec", summary
            return "waiting_resume", clean_text(item.get("codePhase")) or "pending", summary

    for item in ready_items:
        if clean_text(item.get("codeStatus")) == "blocked":
            return "blocked", clean_text(item.get("codePhase")) or "blocked", summary

    if any(item.get("status") in {"pending", "in_progress"} for item in checklist_items):
        return "waiting_spec", previous_phase or "waiting_spec", summary

    return previous_status or "waiting_spec", previous_phase or "waiting_spec", summary


def aggregate_execution_status(spec_status: str, code_status: str) -> str:
    if spec_status == "done" and code_status == "done":
        return "done"
    if spec_status != "done":
        if spec_status in {"running", "in_progress"} or code_status in {"running", "in_progress"}:
            return "running"
        if spec_status == "blocked" or code_status == "blocked":
            return "blocked"
        return "waiting_spec"
    if code_status in {"running", "in_progress"}:
        return "running"
    if code_status == "blocked":
        return "blocked"
    return "waiting_code"


def aggregate_execution_phase(spec_status: str, code_phase: str | None) -> str:
    phase = clean_text(code_phase)
    return phase or spec_status


def latest_timestamp(*values: object) -> str:
    parsed: list[tuple[str, dt.datetime]] = []
    for value in values:
        text = clean_text(value)
        if not text:
            continue
        try:
            parsed.append((text, dt.datetime.fromisoformat(text)))
        except ValueError:
            continue
    if not parsed:
        return now_iso()
    parsed.sort(key=lambda item: item[1])
    return parsed[-1][0]


def build_api_source(
    context: ExecutionContext,
    *,
    workbook_path: Path | None,
    sheet_names: list[str] | None,
) -> dict[str, Any]:
    return {
        "tsdFile": context.docx_path.name,
        "workbookFile": workbook_path.name if workbook_path else None,
        "sheetNames": list(sheet_names or []),
    }


def load_reference_library(agent_dir: Path, function_code: str | None = None) -> dict[str, list[dict[str, Any]]]:
    library: dict[str, list[dict[str, Any]]] = {}
    for root in reference_roots_for_read(agent_dir, function_code):
        if not (root / "catalog.json").exists():
            continue
        for category in REFERENCE_CATEGORIES:
            index_path = root / "indexes" / REFERENCE_INDEX_FILENAMES[category]
            if not index_path.exists():
                continue
            payload = load_json(index_path)
            items = list(payload.get("items") or [])
            for item in items:
                item["matchKeys"] = [normalize_match_key(value) for value in item.get("matchKeys") or [] if clean_text(value)]
                item.setdefault("referenceRoot", root.as_posix())
            library.setdefault(category, []).extend(items)
    return library


def base_relative_reference_path(agent_dir: Path, target: Path) -> str:
    try:
        return target.relative_to(agent_dir).as_posix()
    except ValueError:
        return target.as_posix()


def build_reference_library_warning(agent_dir: Path, function_code: str | None = None) -> str:
    roots = reference_roots_for_read(agent_dir, function_code)
    active_root = next((root for root in roots if (root / "catalog.json").exists()), None)
    if active_root is None:
        catalog_ref = base_relative_reference_path(agent_dir, reference_catalog_path(agent_dir))
        return (
            f"未找到 `{catalog_ref}`；本次会继续生成 API Spec，但不会带入第 01 步参考索引。"
            "若此 API 需要外部 API 或 DB Schema 证据，请补运行可选第 01 步 `$reference-index-importer`。开发规范请改用 `专案规则分析器 --category code-guidelines` 接入 project-rules。"
        )

    missing_indexes: list[str] = []
    for category in REFERENCE_CATEGORIES:
        index_path = active_root / "indexes" / REFERENCE_INDEX_FILENAMES[category]
        if not index_path.exists():
            label = REFERENCE_CATEGORY_LABELS.get(category, category)
            missing_indexes.append(f"{label}: `{base_relative_reference_path(agent_dir, index_path)}`")

    if missing_indexes:
        return (
            "`.agent/reference/global` 索引不完整，缺少 "
            + "、".join(missing_indexes)
            + "；本次会继续生成 API Spec，但相关 `referenceHints` 可能不足。"
            "建议先重新运行第 01 步 `$reference-index-importer`。"
        )
    return ""


def build_db_match_candidates(text: str) -> list[dict[str, object]]:
    full_keys = sorted(key for key in find_table_keys(text) if "." in key)
    keyword_keys = sorted(key for key in find_table_keys(text) if "." not in key)
    candidates: list[dict[str, object]] = []
    seen: set[tuple[str, str]] = set()
    for key in full_keys:
        token = ("sql_table", normalize_match_key(key))
        if token in seen:
            continue
        candidates.append({"matchSource": "sql_table", "matchKey": token[1], "baseScore": 300})
        seen.add(token)
    for key in keyword_keys:
        token = ("sql_keyword", normalize_match_key(key))
        if token in seen:
            continue
        candidates.append({"matchSource": "sql_keyword", "matchKey": token[1], "baseScore": 220})
        seen.add(token)
    return candidates


def build_external_match_candidates(text: str) -> list[dict[str, object]]:
    candidates: list[dict[str, object]] = []
    seen: set[tuple[str, str]] = set()
    for key in sorted(find_api_codes(text)):
        token = ("api_code", normalize_match_key(key))
        if token in seen:
            continue
        candidates.append({"matchSource": "api_code", "matchKey": token[1], "baseScore": 320})
        seen.add(token)
    for key in sorted(find_external_keywords(text)):
        token = ("iris_keyword", normalize_match_key(key))
        if token in seen:
            continue
        candidates.append({"matchSource": "iris_keyword", "matchKey": token[1], "baseScore": 160})
        seen.add(token)
    return candidates


def build_framework_match_candidates(text: str) -> list[dict[str, object]]:
    candidates: list[dict[str, object]] = []
    seen: set[str] = set()
    for key in sorted(find_framework_keywords(text)):
        normalized = normalize_match_key(key)
        if normalized in seen:
            continue
        candidates.append({"matchSource": "framework_keyword", "matchKey": normalized, "baseScore": 180})
        seen.add(normalized)
    return candidates


def reference_source_priority(item: dict[str, Any]) -> int:
    if clean_text(item.get("category")) != "db_schema":
        return 0
    return DB_SOURCE_LABEL_PRIORITY.get(clean_text(item.get("sourceLabel")), 0)


def reference_sort_key(entry: dict[str, Any]) -> tuple[int, int, int, str]:
    return (
        int(entry.get("score") or 0),
        reference_source_priority(entry.get("item") or {}),
        int(clean_text(entry.get("item", {}).get("versionDateToken")) or 0),
        clean_text(entry.get("item", {}).get("title")).casefold(),
    )


def select_reference_match_entries(
    items: list[dict[str, Any]],
    candidates: list[dict[str, object]],
    *,
    collapse_match_keys_to_best_source: bool = False,
) -> list[dict[str, Any]]:
    if not items or not candidates:
        return []

    best_by_ref: dict[tuple[str, str, str], dict[str, Any]] = {}
    for item in items:
        item_match_keys = {normalize_match_key(value) for value in item.get("matchKeys") or [] if clean_text(value)}
        if not item_match_keys:
            continue
        title_token = normalize_token(clean_text(item.get("title")))
        api_codes = {normalize_match_key(value) for value in item.get("apiCodes") or [] if clean_text(value)}
        table_names = {normalize_match_key(value) for value in item.get("tableNames") or [] if clean_text(value)}
        for candidate in candidates:
            match_key = clean_text(candidate.get("matchKey"))
            if match_key not in item_match_keys:
                continue
            score = int(candidate.get("baseScore") or 0)
            if normalize_token(match_key) in title_token:
                score += 20
            if candidate.get("matchSource") == "api_code" and match_key in api_codes:
                score += 20
            if candidate.get("matchSource") == "sql_table" and match_key in table_names:
                score += 20
            entry = {
                "item": item,
                "candidate": candidate,
                "score": score,
            }
            if collapse_match_keys_to_best_source:
                entry_key = (
                    clean_text(candidate.get("matchSource")),
                    match_key,
                    "",
                )
            else:
                entry_key = (
                    clean_text(item.get("refId")),
                    clean_text(candidate.get("matchSource")),
                    match_key,
                )
            existing = best_by_ref.get(entry_key)
            if existing is None or reference_sort_key(entry) > reference_sort_key(existing):
                best_by_ref[entry_key] = entry
    return sorted(best_by_ref.values(), key=reference_sort_key, reverse=True)


def pick_reference_locator(item: dict[str, Any], match_key: str) -> dict[str, str]:
    normalized = normalize_match_key(match_key)
    sheet_values = list((item.get("sheetMatchKeys") or {}).get(normalized) or [])
    if sheet_values:
        return {"sheetName": clean_text(sheet_values[0])}
    section_values = list((item.get("sectionMatchKeys") or {}).get(normalized) or [])
    if section_values:
        return {"sectionTitle": clean_text(section_values[0])}
    page_values = list((item.get("pageMatchKeys") or {}).get(normalized) or [])
    if page_values:
        return {"pageHint": clean_text(page_values[0])}
    sheet_names = list(item.get("sheetNames") or [])
    if len(sheet_names) == 1:
        return {"sheetName": clean_text(sheet_names[0])}
    return {}


def build_reference_reason(category: str, match_source: str, match_key: str, *, source_label: str = "") -> str:
    if category == "db_schema":
        db_reason_map = {
            "db-tableschema": "可读取字段定义与约束",
            "db-relation": "可读取字段定义与约束",
            "db-dictionary": "可读取数据字典与枚举说明",
            "db-traceability": "可读取表映射与来源追踪",
        }
        detail = db_reason_map.get(source_label, "可读取数据库参考资料")
        return f"命中 SQL 表 {match_key}，{detail}"
    if category == "external_api" and match_source == "api_code":
        return f"命中外部接口代号 {match_key}，可读取 IRIS/OpenAPI 设计"
    if category == "external_api":
        return f"命中外部接口关键词 {match_key}，可读取 IRIS/OpenAPI 设计"
    return f"命中参考关键词 {match_key}，可读取外部参考资料"


def build_reference_hint(entry: dict[str, Any], *, category: str) -> dict[str, Any]:
    item = entry["item"]
    candidate = entry["candidate"]
    match_key = clean_text(candidate.get("matchKey"))
    locator = pick_reference_locator(item, match_key)
    prefix = {
        "db_schema": "db",
        "external_api": "external",
    }[category]
    return {
        "id": f"{prefix}.{reference_slugify(match_key)}.{reference_slugify(clean_text(item.get('title')))}",
        "category": category,
        "matchSource": clean_text(candidate.get("matchSource")),
        "matchKey": match_key,
        "title": clean_text(item.get("title")),
        "relativePath": clean_text(item.get("relativePath")),
        "locator": locator,
        "reason": build_reference_reason(
            category,
            clean_text(candidate.get("matchSource")),
            match_key,
            source_label=clean_text(item.get("sourceLabel")),
        ),
        "authority": "reference_imported",
    }


def dedupe_reference_hints(hints: list[dict[str, Any]]) -> list[dict[str, Any]]:
    ordered: list[dict[str, Any]] = []
    seen: set[str] = set()
    for hint in hints:
        hint_id = clean_text(hint.get("id"))
        if not hint_id or hint_id in seen:
            continue
        ordered.append(hint)
        seen.add(hint_id)
    return ordered


def format_reference_hint_brief(hint: dict[str, Any]) -> str:
    locator = format_reference_locator(hint.get("locator") or {})
    locator_suffix = f"#{locator}" if locator else ""
    return f"{clean_text(hint.get('title'))} -> {clean_text(hint.get('relativePath'))}{locator_suffix}"


def append_reference_hint_line(text: str, hints: list[dict[str, Any]]) -> str:
    rendered = clean_text(text)
    if not hints:
        return rendered
    hint_line = "可读取外援：" + "；".join(format_reference_hint_brief(hint) for hint in hints)
    if not rendered:
        return hint_line
    return f"{rendered}\n{hint_line}"


def collect_backend_reference_texts(backend_apis: dict[str, list[str]], raw_appendix: dict[str, object] | None) -> list[str]:
    texts: list[str] = []
    for system, targets in sorted(backend_apis.items()):
        texts.append(system)
        texts.extend(targets)
    if raw_appendix:
        texts.extend(clean_text(line) for line in raw_appendix.get("backendApiLines") or [])
    return [text for text in texts if clean_text(text)]


def apply_reference_hints_to_business_logic(
    business_logic_payload: dict[str, Any],
    *,
    reference_library: dict[str, list[dict[str, Any]]],
    backend_apis: dict[str, list[str]],
    raw_appendix: dict[str, object] | None,
) -> dict[str, Any]:
    if not reference_library:
        return business_logic_payload

    payload = json.loads(json.dumps(business_logic_payload, ensure_ascii=False))
    aggregated_hints: list[dict[str, Any]] = []

    db_items = list(reference_library.get("db_schema") or [])
    external_items = list(reference_library.get("external_api") or [])
    for sql_spec in payload.get("sqlSpecs") or []:
        matches = select_reference_match_entries(
            db_items,
            build_db_match_candidates(clean_text(sql_spec.get("queryText"))),
            collapse_match_keys_to_best_source=True,
        )
        hints = [build_reference_hint(entry, category="db_schema") for entry in matches[:REFERENCE_HINT_LIMIT]]
        if hints:
            sql_spec["referenceHints"] = hints
            sql_spec["queryText"] = append_reference_hint_line(clean_text(sql_spec.get("queryText")), hints)
            aggregated_hints.extend(hints)

    for step in payload.get("steps") or []:
        step_text = "\n".join(part for part in [clean_text(step.get("title")), clean_text(step.get("details"))] if part)
        match_entries = select_reference_match_entries(external_items, build_external_match_candidates(step_text))
        sorted_entries = sorted(match_entries, key=reference_sort_key, reverse=True)[:REFERENCE_HINT_LIMIT]
        hints = dedupe_reference_hints(
            [
                build_reference_hint(entry, category=entry["item"]["category"])
                for entry in sorted_entries
            ]
        )
        if hints:
            step["referenceHints"] = hints
            step["details"] = append_reference_hint_line(clean_text(step.get("details")), hints)
            aggregated_hints.extend(hints)

    for data_source in payload.get("dataSources") or []:
        matches = select_reference_match_entries(
            db_items,
            build_db_match_candidates(clean_text(data_source.get("name"))),
            collapse_match_keys_to_best_source=True,
        )
        aggregated_hints.extend(build_reference_hint(entry, category="db_schema") for entry in matches[:REFERENCE_HINT_LIMIT])

    backend_text = "\n".join(collect_backend_reference_texts(backend_apis, raw_appendix))
    backend_matches = select_reference_match_entries(external_items, build_external_match_candidates(backend_text))
    aggregated_hints.extend(build_reference_hint(entry, category="external_api") for entry in backend_matches[:REFERENCE_HINT_LIMIT])

    deduped = dedupe_reference_hints(aggregated_hints)
    if deduped:
        payload["referenceHints"] = deduped
    return payload


def stable_payload_hash(payload: dict[str, Any]) -> str:
    rendered = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return f"sha256:{hashlib.sha256(rendered.encode('utf-8')).hexdigest()}"


def write_execution_state(
    context: ExecutionContext,
    checklist_items: list[dict[str, Any]],
    *,
    status: str,
) -> dict[str, Any]:
    existing_payload = load_json(context.paths.execution_state_path) if context.paths.execution_state_path.exists() else {}
    code_status, code_phase, code_summary = derive_code_execution_projection(checklist_items, existing_payload)
    spec_updated_at = now_iso()
    fixture_statuses = [clean_text(item.get("fixtureStatus")) for item in checklist_items if clean_text(item.get("fixtureStatus"))]
    if any(value == "blocked" for value in fixture_statuses):
        fixture_status = "blocked"
    elif any(value == "error" for value in fixture_statuses):
        fixture_status = "error"
    elif fixture_statuses and all(value in {"done", "skipped"} for value in fixture_statuses):
        fixture_status = "done"
    elif any(value == "pending" for value in fixture_statuses):
        fixture_status = "pending"
    else:
        fixture_status = clean_text(existing_payload.get("fixtureStatus")) or None
    fixture_phase = clean_text(existing_payload.get("fixturePhase")) or ""
    if fixture_status == "done" and fixture_phase in {"", "pending", "waiting_spec", "waiting_fixture"}:
        fixture_phase = "applied"
    elif fixture_status == "pending" and fixture_phase in {"", "waiting_spec", "waiting_fixture", "applied"}:
        fixture_phase = "pending"
    elif not fixture_phase:
        fixture_phase = fixture_status
    fixture_summary = {
        "total": len(checklist_items),
        "pending": sum(1 for item in checklist_items if clean_text(item.get("fixtureStatus")) == "pending"),
        "in_progress": sum(1 for item in checklist_items if clean_text(item.get("fixtureStatus")) == "in_progress"),
        "done": sum(1 for item in checklist_items if clean_text(item.get("fixtureStatus")) == "done"),
        "skipped": sum(1 for item in checklist_items if clean_text(item.get("fixtureStatus")) == "skipped"),
        "blocked": sum(1 for item in checklist_items if clean_text(item.get("fixtureStatus")) == "blocked"),
        "error": sum(1 for item in checklist_items if clean_text(item.get("fixtureStatus")) == "error"),
    }
    payload = {
        "schemaVersion": STATE_SCHEMA_VERSION,
        "executionId": context.execution_id,
        "functionCode": context.function_code,
        "status": aggregate_execution_status(status, code_status),
        "phase": aggregate_execution_phase(status, code_phase),
        "updatedAt": latest_timestamp(spec_updated_at, existing_payload.get("codeUpdatedAt")),
        "specStatus": status,
        "specUpdatedAt": spec_updated_at,
        "specSummary": summarize_checklist(checklist_items),
        "specDocxPath": normalize_persisted_path(context.docx_path, project_root=context.project_root),
        "specLastMessage": clean_text(existing_payload.get("specLastMessage")),
        "codeStatus": code_status,
        "codePhase": code_phase,
        "codeUpdatedAt": existing_payload.get("codeUpdatedAt"),
        "codeCurrentApiId": existing_payload.get("codeCurrentApiId"),
        "codeSummary": code_summary,
        "codeProjectRoot": existing_payload.get("codeProjectRoot"),
        "codeSolutionPath": existing_payload.get("codeSolutionPath"),
        "codeLastMessage": existing_payload.get("codeLastMessage"),
        "fixtureStatus": fixture_status,
        "fixturePhase": fixture_phase,
        "fixtureUpdatedAt": existing_payload.get("fixtureUpdatedAt"),
        "fixtureCurrentApiId": existing_payload.get("fixtureCurrentApiId"),
        "fixtureSummary": fixture_summary,
        "fixtureLastMessage": existing_payload.get("fixtureLastMessage"),
        "artifacts": {
            "batchFile": normalize_persisted_path(context.batch_file, project_root=context.project_root),
            "checklist": normalize_persisted_path(context.paths.checklist_path, project_root=context.project_root),
            "specProgress": normalize_persisted_path(context.paths.progress_path, project_root=context.project_root),
            "codeProgress": normalize_persisted_path(context.paths.root / "code-progress.md", project_root=context.project_root),
            "repoSnapshot": normalize_persisted_path(context.paths.root / "repo-snapshot.json", project_root=context.project_root),
        },
    }
    dump_json(context.paths.execution_state_path, payload)
    update_chain_status(
        agent_root=context.agent_dir,
        function_code=context.function_code,
        stage="spec",
        status=payload.get("specStatus"),
        phase=payload.get("phase"),
        message=payload.get("specLastMessage") or payload.get("status"),
        project_root=context.project_root,
        artifacts={
            "executionState": normalize_persisted_path(context.paths.execution_state_path, project_root=context.agent_dir),
            "apiChecklist": normalize_persisted_path(context.paths.checklist_path, project_root=context.agent_dir),
            "specProgress": normalize_persisted_path(context.paths.progress_path, project_root=context.agent_dir),
        },
    )
    return payload


def validate_manifest_payload(payload: dict[str, Any]) -> None:
    schema = load_schema("manifest.schema.json")
    validator = Draft202012Validator(schema)
    errors = sorted(validator.iter_errors(payload), key=lambda item: list(item.absolute_path))
    if errors:
        first = errors[0]
        raise SkillError(f"manifest schema 校验失败：{format_validation_path(first.absolute_path)} - {first.message}")


def build_context(args: argparse.Namespace) -> ExecutionContext:
    project_root = resolve_project_root(args.project_root)
    agent_dir = resolve_agent_dir(project_root, args.agent_dir, args.agent_root, args.workspace_root, args.workspace_key, args.rules_root)
    context_root = resolve_context_root(project_root, agent_dir, args.context_root)
    docx_path, function_code, batch_payload, _, batch_file = resolve_batch_target(
        project_root,
        agent_dir,
        context_root,
        docx_ref=args.docx_ref,
        requested_function_code=args.function_code,
        design_handoff=args.design_handoff,
        allow_legacy_input=args.allow_legacy_input,
    )
    new_author = normalize_author_name(args.new_author or getpass.getuser())
    if not function_code:
        raise SkillError(f"无法从 TSD 文件名提取功能编号：{docx_path.name}", status="blocked")
    execution_id = function_code
    state_root = (context_root / function_code).resolve()
    batch_payload = upsert_batch_item(
        batch_payload,
        function_code=function_code,
        docx_ref=normalize_batch_docx_ref(project_root, docx_path),
        make_active=True,
    )
    save_batch_file(batch_file, batch_payload, updated_by=new_author)
    context = ExecutionContext(
        project_root=project_root,
        agent_dir=agent_dir,
        context_root=context_root,
        batch_file=batch_file,
        state_root=state_root,
        docx_path=docx_path,
        execution_id=execution_id,
        function_code=function_code,
        new_author=new_author,
    )
    return context


def header_kind(value: str) -> str | None:
    token = normalize_token(value)
    if token in {"api類別", "api类别", "apicategory"}:
        return "category"
    if token in {"api名稱", "api名称", "apiname"}:
        return "name"
    if token in {"功能說明", "功能说明", "說明", "说明", "description"}:
        return "description"
    return None


def extract_api_entries(docx_path: Path) -> list[ApiEntry]:
    try:
        document = Document(docx_path)
    except (BadZipFile, PackageNotFoundError, OSError) as exc:
        raise SkillError(f"无法读取 TSD 文档：{docx_path.as_posix()}") from exc

    found: list[ApiEntry] = []
    seen: set[tuple[str, str]] = set()
    for table in document.tables:
        if not table.rows:
            continue
        header_map: dict[str, int] = {}
        for index, cell in enumerate(table.rows[0].cells):
            kind = header_kind(cell.text)
            if kind and kind not in header_map:
                header_map[kind] = index
        if "category" not in header_map or "name" not in header_map:
            continue

        last_category = ""
        for row in table.rows[1:]:
            values = [clean_text(cell.text) for cell in row.cells]
            if not any(values):
                continue
            category = values[header_map["category"]].strip()
            name = values[header_map["name"]].strip()
            description = values[header_map.get("description", -1)].strip() if "description" in header_map else ""
            if category:
                last_category = category
            elif last_category:
                category = last_category
            if not category or not name:
                continue
            key = (category.casefold(), name.casefold())
            if key in seen:
                continue
            found.append(ApiEntry(category=category, name=name, description=description))
            seen.add(key)

    if not found:
        raise SkillError(f"未能从文档中抽取 API 清单，请检查 {docx_path.name} 是否包含 API類別/API名稱 表格", status="blocked")
    return found


def extract_date_token(path: Path) -> tuple[int, float]:
    match = re.search(r"_(\d{8})(?=\.xlsx$)", path.name, flags=re.IGNORECASE)
    return int(match.group(1)) if match else 0, path.stat().st_mtime


def is_common_category(category: str) -> bool:
    token = normalize_token(category)
    return "common" in token or "backend" in token


def derive_embedded_primary_category(category: str) -> str | None:
    match = re.match(r"^(?P<base>.+?)\s*CommonUtil$", clean_text(category), flags=re.IGNORECASE)
    if not match:
        return None
    base = clean_text(match.group("base"))
    return base or None


def find_workbook_candidates(search_dir: Path, category: str) -> list[Path]:
    if not search_dir.exists():
        return []
    candidates = [
        path
        for path in search_dir.glob("*.xlsx")
        if path.is_file() and path.name.casefold().startswith(f"newda_api_detail_{category}_".casefold())
    ]
    if not candidates:
        candidates = [
            path for path in search_dir.glob("*.xlsx") if path.is_file() and normalize_token(category) in normalize_token(path.stem)
        ]
    candidates.sort(key=extract_date_token, reverse=True)
    return candidates


def choose_workbook_for_category(agent_dir: Path, category: str, function_code: str | None = None) -> Path:
    input_root = agent_dir / "functions" / function_code / "inputs" if function_code else None
    search_dirs: list[Path] = []
    if input_root is not None:
        search_dirs.append(input_root / ("common" if is_common_category(category) else "api-spec"))
    search_dirs.append(agent_dir / ("Common" if is_common_category(category) else "API Spec"))
    if is_common_category(category):
        missing_message = f"请导入共用{category}文件"
    else:
        missing_message = f"请导入{category} Excel Spec文件"

    candidates: list[Path] = []
    for search_dir in search_dirs:
        candidates = find_workbook_candidates(search_dir, category)
        if candidates:
            break
    if not candidates:
        embedded_category = derive_embedded_primary_category(category)
        if embedded_category:
            embedded_dirs: list[Path] = []
            if input_root is not None:
                embedded_dirs.append(input_root / "api-spec")
            embedded_dirs.append(agent_dir / "API Spec")
            for search_dir in embedded_dirs:
                candidates = find_workbook_candidates(search_dir, embedded_category)
                if candidates:
                    break
    if not candidates:
        raise SkillError(missing_message, status="blocked")
    return candidates[0]


def sheet_name_is_exact_direct_match(sheet_name: str, api: ApiEntry) -> bool:
    return sheet_name.casefold().startswith(api.name.casefold())


def leading_ascii_identifier(value: str) -> str:
    match = re.match(r"^[A-Za-z0-9_.-]+", clean_text(value))
    return normalize_token(match.group(0)) if match else ""


def sheet_name_is_fuzzy_match(sheet_name: str, api: ApiEntry) -> bool:
    if sheet_name_is_exact_direct_match(sheet_name, api):
        return True

    api_token = normalize_token(api.name)
    sheet_prefix = leading_ascii_identifier(sheet_name)
    if not api_token or not sheet_prefix:
        return False

    if api_token.startswith(sheet_prefix) and len(sheet_prefix) >= max(8, len(api_token) - 2):
        return True
    if sheet_prefix.startswith(api_token) and len(api_token) >= max(8, len(sheet_prefix) - 2):
        return True

    similarity = difflib.SequenceMatcher(a=api_token, b=sheet_prefix).ratio()
    return similarity >= 0.94


def sheet_name_is_direct_match(sheet_name: str, api: ApiEntry) -> bool:
    return sheet_name_is_fuzzy_match(sheet_name, api)


def score_sheet_api_mentions(sheet: Worksheet, api_name: str) -> int:
    token = normalize_token(api_name)
    if not token:
        return 0
    score = 0
    for row in sheet.iter_rows(values_only=True):
        row_text = normalize_token(" ".join("" if cell is None else str(cell) for cell in row))
        if token not in row_text:
            continue
        score += max(1, row_text.count(token))
    return score


def find_indirect_sheet_names(workbook_path: Path, api: ApiEntry) -> list[str]:
    workbook = load_workbook(workbook_path, read_only=True, data_only=False)
    try:
        scored_matches: list[tuple[str, int]] = []
        for sheet_name in workbook.sheetnames:
            score = score_sheet_api_mentions(workbook[sheet_name], api.name)
            if score > 0:
                scored_matches.append((sheet_name, score))
    finally:
        workbook.close()

    if not scored_matches:
        return []

    non_summary_matches = [
        item
        for item in scored_matches
        if normalize_token(item[0]) not in {normalize_token("Api_List"), normalize_token("API List")}
    ]
    scored_matches = non_summary_matches or scored_matches
    scored_matches.sort(key=lambda item: (-item[1], item[0].casefold()))
    return [sheet_name for sheet_name, _ in scored_matches]


def find_sheet_names(workbook_path: Path, api: ApiEntry) -> list[str]:
    workbook = load_workbook(workbook_path, data_only=False)
    try:
        matches = [sheet for sheet in workbook.sheetnames if sheet_name_is_fuzzy_match(sheet, api)]
    finally:
        workbook.close()
    if not matches and derive_embedded_primary_category(api.category):
        matches = find_indirect_sheet_names(workbook_path, api)
    if not matches:
        raise SkillError(f"{api.category}的{api.name}不存在，请检查", status="blocked")
    return matches


def unique_non_empty(values: Iterable[str]) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for value in values:
        text = clean_text(value)
        if not text or text in seen:
            continue
        seen.add(text)
        ordered.append(text)
    return ordered


def normalize_zip_member_name(value: str) -> str:
    return posixpath.normpath(clean_text(value).replace("\\", "/")).lstrip("./")


def build_sheet_relationship_member(sheet_path: str) -> str:
    normalized = normalize_zip_member_name(sheet_path)
    directory, filename = posixpath.split(normalized)
    if not directory or not filename:
        return ""
    return posixpath.join(directory, "_rels", f"{filename}.rels")


def resolve_zip_relationship_target(source_member: str, target: str) -> str:
    return posixpath.normpath(posixpath.join(posixpath.dirname(source_member), clean_text(target).replace("\\", "/"))).lstrip("./")


def inspect_drawing_part(archive: ZipFile, drawing_member: str) -> tuple[int, int]:
    root = ET.fromstring(archive.read(drawing_member))
    anchor_count = 0
    for anchor_tag in ("oneCellAnchor", "twoCellAnchor", "absoluteAnchor"):
        anchor_count += len(root.findall(f".//{SPREADSHEET_DRAWING_NS}{anchor_tag}"))
    image_count = len(root.findall(f".//{SPREADSHEET_DRAWING_NS}pic"))
    return anchor_count, image_count


def inspect_sheet_visual_signals(workbook_path: Path, sheet_names: list[str]) -> list[SheetVisualSignal]:
    requested_sheet_names = unique_non_empty(sheet_names)
    if not requested_sheet_names:
        return []

    workbook = load_workbook(workbook_path, data_only=False)
    try:
        sheet_members: list[tuple[str, str, int]] = []
        for sheet_name in requested_sheet_names:
            if sheet_name not in workbook.sheetnames:
                continue
            sheet = workbook[sheet_name]
            sheet_members.append(
                (
                    sheet_name,
                    normalize_zip_member_name(getattr(sheet, "path", "")),
                    len(getattr(sheet, "_images", []) or []),
                )
            )
    finally:
        workbook.close()

    signals: list[SheetVisualSignal] = []
    with ZipFile(workbook_path) as archive:
        archive_members = set(archive.namelist())
        for sheet_name, sheet_member, image_count in sheet_members:
            drawing_anchor_count = 0
            relationship_member = build_sheet_relationship_member(sheet_member)
            if relationship_member in archive_members:
                rel_root = ET.fromstring(archive.read(relationship_member))
                for relationship in rel_root.findall(f"{PACKAGE_REL_NS}Relationship"):
                    if clean_text(relationship.attrib.get("Type")) != DRAWING_RELATIONSHIP_TYPE:
                        continue
                    target_member = resolve_zip_relationship_target(relationship_member, clean_text(relationship.attrib.get("Target")))
                    if target_member not in archive_members:
                        continue
                    anchor_count, drawing_image_count = inspect_drawing_part(archive, target_member)
                    drawing_anchor_count += anchor_count
                    image_count = max(image_count, drawing_image_count)
            if image_count > 0 or drawing_anchor_count > 0:
                signals.append(
                    SheetVisualSignal(
                        sheet_name=sheet_name,
                        image_count=image_count,
                        drawing_anchor_count=drawing_anchor_count,
                    )
                )
    return signals


KNOWN_EXCEL_VISUAL_RULES: list[dict[str, Any]] = [
    {
        "workbookPrefix": "NEWDA_API_DETAIL_Setting_",
        "sheetName": "UpdateUserAlias變更暱稱",
        "blocking": False,
        "notes": [
            "嵌入图片转录：现有测试资料示意表的列顺序为 CUSTID、IP、DEVICEINFO、APPVERSION、OSVERSION、ALIAS、TIMESTAMP，对应 DA_USER_ALIAS 的现有资料样式。",
            "嵌入图片转录：User-Agent 拆解映射为 APPName、APPVer、UDID、DeviceOS、DeviceOSVer、DeviceModel；分别表示 APP 名称、APP 版本号、装置 ID、装置 OS 名称、装置 OS 版本、装置型号。",
            "视觉规则判定：上述图片仅补充样式示意，核心映射规则已在工作表正文中明文描述，因此不再作为 blocking unresolved。",
        ],
    },
    {
        "workbookPrefix": "NEWDA_API_DETAIL_Login_",
        "sheetName": "GetAnnounce停机及台风天公告获取",
        "blocking": False,
        "notes": [
            "嵌入图片转录：主要读取栏位为 ItemType、Content、StartDate、EndDate、Enabled、UpdateDT。",
            "视觉规则判定：这些栏位已在同 sheet 的 FrontContent_Announcement 表结构区块中明文列出，该图片仅作重复提醒，不再作为 blocking unresolved。",
        ],
    },
    {
        "workbookPrefix": "NEWDA_API_DETAIL_Setting_",
        "sheetName": "QueryUserLoginLog登入記錄查詢",
        "blocking": False,
        "notes": [
            "嵌入图片转录：该图片为 USER_LOGIN_LOG 样本资料截图，展示 TIMESTAMP、IP、SOURCE、OS、DEVICE、LOGINTYPE、COUNTRY 等现有资料样式。",
            "视觉规则判定：SQL、字段说明、旧代码逻辑与前端显示规则已在工作表正文中明文描述，该截图仅作示意，不再作为 blocking unresolved。",
        ],
    },
]


def find_excel_visual_rule(workbook_name: str, sheet_name: str) -> dict[str, Any] | None:
    normalized_workbook_name = clean_text(workbook_name)
    normalized_sheet_name = clean_text(sheet_name)
    for rule in KNOWN_EXCEL_VISUAL_RULES:
        workbook_prefix = clean_text(rule.get("workbookPrefix"))
        expected_sheet_name = clean_text(rule.get("sheetName"))
        if workbook_prefix and not normalized_workbook_name.startswith(workbook_prefix):
            continue
        if expected_sheet_name and normalized_sheet_name != expected_sheet_name:
            continue
        return rule
    return None


def collect_excel_visual_warnings(workbook_path: Path, sheet_names: list[str]) -> tuple[list[str], list[dict[str, Any]]]:
    if not sheet_names:
        return [], []
    try:
        signals = inspect_sheet_visual_signals(workbook_path, sheet_names)
    except (BadZipFile, ET.ParseError, KeyError, OSError, ValueError) as exc:
        message = (
            f"无法检查 `{workbook_path.name}` 的 drawing layer：{clean_text(str(exc)) or exc.__class__.__name__}；"
            "当前解析仅提取单元格值，图片、图形或图表中的内容可能遗漏。"
        )
        return [message], [{"topic": "excel_visual.scan_failed", "reason": message, "blocking": True}]

    notes: list[str] = []
    unresolved: list[dict[str, Any]] = []
    for index, signal in enumerate(signals, start=1):
        visual_rule = find_excel_visual_rule(workbook_path.name, signal.sheet_name)
        if visual_rule is not None:
            rule_notes = [clean_text(note) for note in list(visual_rule.get("notes") or []) if clean_text(note)]
            notes.extend(rule_notes)
            if not bool(visual_rule.get("blocking", True)):
                continue
        summary_parts: list[str] = []
        if signal.image_count > 0:
            summary_parts.append(f"{signal.image_count} 张嵌入图片")
        if signal.drawing_anchor_count > 0:
            summary_parts.append(f"{signal.drawing_anchor_count} 个 drawing anchor")
        summary = "、".join(summary_parts)
        message = (
            f"sheet `{signal.sheet_name}` 含有 {summary}；当前解析仅提取单元格值，"
            "图片、图形、图表或 drawing layer 内的规则/文字需要人工复核。"
        )
        topic_slug = reference_slugify(signal.sheet_name) or f"sheet_{index}"
        notes.append(message)
        unresolved.append(
            {
                "topic": f"excel_visual.{topic_slug}",
                "reason": f"`{workbook_path.name}` 的 sheet `{signal.sheet_name}` 含有 {summary}，当前规格抽取不会解析这些视觉对象内的内容。",
                "blocking": True,
            }
        )
    return notes, unresolved


def build_merged_anchor_map(sheet: Worksheet) -> dict[tuple[int, int], tuple[int, int]]:
    merged_anchors: dict[tuple[int, int], tuple[int, int]] = {}
    for merged_range in sheet.merged_cells.ranges:
        for row in range(merged_range.min_row, merged_range.max_row + 1):
            for col in range(merged_range.min_col, merged_range.max_col + 1):
                merged_anchors[(row, col)] = (merged_range.min_row, merged_range.min_col)
    return merged_anchors


def worksheet_to_rows(sheet: Worksheet) -> list[list[str]]:
    merged_anchors = build_merged_anchor_map(sheet)
    rows: list[list[str]] = []
    actual_max_col = 0
    actual_max_row = 0
    for row_index in range(1, sheet.max_row + 1):
        row_values: list[str] = []
        row_has_value = False
        for col_index in range(1, sheet.max_column + 1):
            anchor = merged_anchors.get((row_index, col_index))
            if anchor and anchor != (row_index, col_index):
                raw = ""
            else:
                raw = clean_text(sheet.cell(row_index, col_index).value)
            if raw:
                row_has_value = True
                actual_max_col = max(actual_max_col, col_index)
            row_values.append(raw)
        if row_has_value:
            actual_max_row = row_index
        rows.append(row_values)
    return [row[:actual_max_col] for row in rows[:actual_max_row]]


def decimal_scale(value: str) -> int:
    text = value.strip()
    if "." not in text:
        return 0
    return len(text.split(".", 1)[1])


def format_decimal_with_scale(value: Decimal, scale: int) -> str:
    quantized = value.quantize(Decimal(1).scaleb(-scale)) if scale > 0 else value.quantize(Decimal("1"))
    rendered = format(quantized, "f")
    if scale == 0:
        return rendered.split(".", 1)[0]
    integer, _, fraction = rendered.partition(".")
    return f"{integer}.{fraction[:scale].ljust(scale, '0')}"


def normalize_sequence_tokens(rows: list[list[str]]) -> list[list[str]]:
    normalized = [row[:] for row in rows]
    previous_number: Decimal | None = None
    previous_scale = 0
    for row in normalized:
        if not row:
            continue
        first_cell = clean_text(row[0])
        if not first_cell:
            continue
        if SIMPLE_DECIMAL_RE.match(first_cell):
            previous_number = Decimal(first_cell)
            previous_scale = decimal_scale(first_cell)
            continue
        formula_match = FORMULA_INCREMENT_RE.match(first_cell)
        if formula_match and previous_number is not None:
            increment_text = formula_match.group("increment")
            increment = Decimal(increment_text)
            previous_scale = max(previous_scale, decimal_scale(increment_text))
            previous_number += increment
            row[0] = format_decimal_with_scale(previous_number, previous_scale)
    return normalized


def load_sheet_rows(match: SheetMatch) -> list[list[str]]:
    workbook = load_workbook(match.workbook_path, data_only=False)
    try:
        return normalize_sequence_tokens(worksheet_to_rows(workbook[match.sheet_name]))
    finally:
        workbook.close()


def first_non_empty_cell(row: list[str]) -> str:
    return next((cell for cell in row if cell), "")


def second_non_empty_cell(row: list[str]) -> str:
    non_empty = [cell for cell in row if cell]
    return non_empty[1] if len(non_empty) > 1 else ""


def join_non_empty_cells(row: list[str]) -> str:
    return "\n".join(cell for cell in row if cell)


def extract_table_names(value: str) -> list[str]:
    text = clean_text(value)
    if not text:
        return []

    patterns = [
        re.compile(r"\[([A-Za-z0-9_]+)\]\.\[([A-Za-z0-9_]+)\]\.\[([A-Za-z0-9_]+)\]"),
        re.compile(r"\b([A-Za-z0-9_]+)\.([A-Za-z0-9_]+)\.([A-Za-z0-9_]+)\b"),
        re.compile(
            r"\b(?:FROM|UPDATE|INTO|JOIN|MERGE\s+INTO|DELETE\s+FROM)\s+([A-Za-z_][A-Za-z0-9_]*(?:\.[A-Za-z_][A-Za-z0-9_]*){0,2})\b",
            re.IGNORECASE,
        ),
    ]

    names: list[str] = []
    seen: set[str] = set()

    def remember(candidate: str) -> None:
        candidate_text = clean_text(candidate)
        if re.match(r"^\d+(?:\.\d+)+$", candidate_text) or re.match(r"^V\d+(?:\.\d+)+$", candidate_text, re.IGNORECASE):
            return
        normalized = clean_text(candidate).upper()
        if not normalized or normalized in seen:
            return
        seen.add(normalized)
        names.append(normalized)

    for match in patterns[0].finditer(text):
        remember(".".join(match.groups()))
    for match in patterns[1].finditer(text):
        remember(".".join(match.groups()))
    for match in patterns[2].finditer(text):
        remember(match.group(1))
    return names


def extract_section_rows(rows: list[list[str]], *, start_title: str, stop_titles: tuple[str, ...]) -> list[list[str]]:
    section_rows: list[list[str]] = []
    in_section = False
    for row in rows:
        first = first_non_empty_cell(row)
        if not in_section:
            if matches_title(first, start_title):
                in_section = True
            continue
        if any(matches_title(first, title) for title in stop_titles):
            break
        section_rows.append(row)
    return section_rows


def find_header_index(headers: list[str], *candidates: str) -> int | None:
    tokens = [normalize_token(header) for header in headers]
    candidate_tokens = {normalize_token(candidate) for candidate in candidates}
    for index, token in enumerate(tokens):
        if token in candidate_tokens or any(candidate and candidate in token for candidate in candidate_tokens):
            return index
    return None


def matches_title(value: str, title: str) -> bool:
    token = normalize_token(value)
    expected = normalize_token(title)
    return token == expected or token.startswith(expected)


def parse_required_flag(value: str) -> bool | None:
    token = normalize_token(value)
    if token in {"y", "yes", "true", "required"}:
        return True
    if token in {"n", "no", "false", "optional"}:
        return False
    return None


def is_datetime_like_field(*, field_name: str, description: str, notes: str, example: str) -> bool:
    token_source = " ".join([field_name, description, notes, example]).casefold()
    if any(marker in token_source for marker in ("yyyy/mm/dd", "yyyy-mm-dd", "hh:mm:ss", "timestamp")):
        return True
    normalized_field_name = re.sub(r"[^a-z0-9]+", "", field_name.casefold())
    return normalized_field_name.endswith(("time", "date", "dt", "datetime", "timestamp"))


def normalize_datetime_text(value: str) -> str:
    def replace(match: re.Match[str]) -> str:
        return f"{match.group('year')}-{match.group('month')}-{match.group('day')}{match.group('suffix') or ''}"

    return re.sub(
        r"(?P<year>\d{4})/(?P<month>\d{2})/(?P<day>\d{2})(?P<suffix>(?:[ T]\d{2}:\d{2}:\d{2})?)",
        replace,
        value,
    )


def normalize_datetime_format_tokens(value: str) -> str:
    return re.sub(
        r"(?i)\b(yyyy)[/-](mm)[/-](dd)\b",
        lambda match: f"{match.group(1)}-{match.group(2)}-{match.group(3)}",
        value,
    )


def normalize_datetime_value(value: str) -> str:
    return normalize_datetime_format_tokens(normalize_datetime_text(value))


def normalize_contract_notes(value: str) -> str:
    return clean_text(value)


def build_mapping_note_index(field_mappings: list[dict[str, str]]) -> dict[str, str]:
    index: dict[str, str] = {}
    for item in field_mappings:
        row_key = clean_text(item.get("rowKey"))
        source = clean_text(item.get("source"))
        if row_key and source:
            index[row_key] = source
    return index


def prune_mapped_notes(fields: list[dict[str, object]], mapping_note_index: dict[str, str]) -> list[dict[str, object]]:
    def prune(node: dict[str, object]) -> dict[str, object]:
        row_key = clean_text(node.get("rowKey"))
        note = clean_text(node.get("notes"))
        source = mapping_note_index.get(row_key, "")
        if note and source and normalize_token(note) == normalize_token(source):
            node = dict(node)
            node.pop("notes", None)
        children = node.get("properties") or []
        if children:
            node["properties"] = [prune(child) for child in children]
        return node

    return [prune(field) for field in fields]


def normalize_contract_example(value: str, *, field_name: str, data_type: str, has_children: bool) -> str:
    text = clean_text(value)
    if not text:
        return ""
    if has_children and data_type in {"object", "json", "array"}:
        return re.sub(r'":\s*(\{\}|\[\])', r'": \1', text)
    return text


def normalize_contract_data_type(
    value: str,
    *,
    has_children: bool,
    field_name: str,
    description: str,
    notes: str,
    example: str,
) -> str:
    text = clean_text(value)
    token = normalize_token(text)
    if token in {"integer"}:
        return "int"
    return text


def build_field_tree(flat_fields: list[dict[str, object]]) -> list[dict[str, object]]:
    roots: list[dict[str, object]] = []
    nodes: dict[str, dict[str, object]] = {}
    for item in flat_fields:
        row_key = clean_text(item.get("rowKey"))
        node = {
            "rowKey": row_key,
            "fieldName": clean_text(item.get("fieldName")),
            "dataType": clean_text(item.get("dataType")),
            "required": item.get("required"),
            "description": clean_text(item.get("description")),
            "example": clean_text(item.get("example")),
            "notes": clean_text(item.get("notes")),
        }
        nodes[row_key] = node
        parent_key = row_key.rsplit(".", 1)[0] if "." in row_key else None
        if parent_key and parent_key in nodes:
            nodes[parent_key].setdefault("properties", []).append(node)
        else:
            roots.append(node)

    def finalize(node: dict[str, object]) -> dict[str, object]:
        children = [finalize(child) for child in node.get("properties", []) or []]
        if children:
            node["properties"] = children
        else:
            node.pop("properties", None)
        normalized_notes = normalize_contract_notes(clean_text(node.get("notes")))
        normalized_data_type = normalize_contract_data_type(
            clean_text(node.get("dataType")),
            has_children=bool(children),
            field_name=clean_text(node.get("fieldName")),
            description=clean_text(node.get("description")),
            notes=normalized_notes,
            example=clean_text(node.get("example")),
        )
        node["dataType"] = normalized_data_type
        node["example"] = normalize_contract_example(
            clean_text(node.get("example")),
            field_name=clean_text(node.get("fieldName")),
            data_type=normalized_data_type,
            has_children=bool(children),
        )
        node["notes"] = normalized_notes
        if node.get("required") is None:
            node.pop("required", None)
        for key in ("description", "example", "notes", "dataType", "fieldName", "rowKey"):
            if node.get(key) == "":
                node.pop(key, None)
        return node

    return [finalize(node) for node in roots]


def extract_structured_fields(
    rows: list[list[str]],
    *,
    section_title: str,
    stop_titles: tuple[str, ...],
) -> tuple[list[dict[str, object]], list[str]]:
    section_rows = extract_section_rows(rows, start_title=section_title, stop_titles=stop_titles)
    if not section_rows:
        return [], []
    header_index = next(
        (index for index, row in enumerate(section_rows) if normalize_token(first_non_empty_cell(row)) == normalize_token("#")),
        None,
    )
    if header_index is None:
        return [], [join_non_empty_cells(row) for row in section_rows if any(row)]

    headers = section_rows[header_index]
    field_name_index = find_header_index(headers, "欄位名稱")
    data_type_index = find_header_index(headers, "資料型態")
    required_index = find_header_index(headers, "必填")
    description_index = find_header_index(headers, "說明", "欄位說明")
    example_index = find_header_index(headers, "範例")
    notes_index = find_header_index(headers, "備註", "備註(包含Source Description，需要說明每個Response欄位的來源及處理邏輯)")

    flat_fields: list[dict[str, object]] = []
    notes: list[str] = []
    for row in section_rows[header_index + 1 :]:
        if not any(row):
            continue
        row_key = clean_text(row[0]) if row else ""
        if not row_key or not ROW_KEY_RE.match(row_key):
            notes.append(join_non_empty_cells(row))
            continue
        flat_fields.append(
            {
                "rowKey": row_key,
                "fieldName": clean_text(row[field_name_index]) if field_name_index is not None and field_name_index < len(row) else "",
                "dataType": clean_text(row[data_type_index]) if data_type_index is not None and data_type_index < len(row) else "",
                "required": parse_required_flag(row[required_index]) if required_index is not None and required_index < len(row) else None,
                "description": clean_text(row[description_index]) if description_index is not None and description_index < len(row) else "",
                "example": clean_text(row[example_index]) if example_index is not None and example_index < len(row) else "",
                "notes": clean_text(row[notes_index]) if notes_index is not None and notes_index < len(row) else "",
            }
        )
    return build_field_tree(flat_fields), notes


def try_parse_json_payload(value: str) -> object | None:
    text = clean_text(value)
    if not text:
        return None
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return None


def build_field_name_index(fields: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    return {clean_text(field.get("fieldName")): field for field in fields if clean_text(field.get("fieldName"))}


def normalize_payload_against_contract(payload: object, fields: list[dict[str, Any]]) -> object:
    if not isinstance(payload, dict):
        return payload

    field_index = build_field_name_index(fields)
    normalized: dict[str, object] = {}
    for key, value in payload.items():
        field = field_index.get(clean_text(key))
        if field is None:
            normalized[key] = value
            continue

        child_fields = list(field.get("properties") or [])
        data_type = clean_text(field.get("dataType"))
        if data_type == "array" and child_fields:
            if isinstance(value, dict):
                if not value:
                    normalized[key] = []
                else:
                    normalized[key] = [normalize_payload_against_contract(value, child_fields)]
                continue
            if isinstance(value, list):
                normalized[key] = [
                    normalize_payload_against_contract(item, child_fields) if isinstance(item, dict) else item for item in value
                ]
                continue
        if isinstance(value, dict) and child_fields:
            normalized[key] = normalize_payload_against_contract(value, child_fields)
            continue
        if isinstance(value, list) and child_fields:
            normalized[key] = [
                normalize_payload_against_contract(item, child_fields) if isinstance(item, dict) else item for item in value
            ]
            continue
        if data_type == "DateTime" and isinstance(value, str):
            normalized[key] = normalize_datetime_value(value)
            continue
        normalized[key] = value
    return normalized


def mark_field_tree_optional(field: dict[str, Any]) -> dict[str, Any]:
    node = dict(field)
    if node.get("required") is True:
        node["required"] = False
    children = list(node.get("properties") or [])
    if children:
        node["properties"] = [mark_field_tree_optional(child) for child in children]
    return node


def relax_field_requirements_from_examples(fields: list[dict[str, Any]], payload_samples: list[object]) -> list[dict[str, Any]]:
    relaxed_fields: list[dict[str, Any]] = []
    for field in fields:
        node = dict(field)
        field_name = clean_text(node.get("fieldName"))
        child_fields = list(node.get("properties") or [])
        if not child_fields or not field_name:
            relaxed_fields.append(node)
            continue

        child_payloads = [
            sample[field_name]
            for sample in payload_samples
            if isinstance(sample, dict) and field_name in sample and isinstance(sample[field_name], dict)
        ]
        empty_object_present = any(not payload for payload in child_payloads)
        normalized_children = relax_field_requirements_from_examples(child_fields, child_payloads)
        if empty_object_present:
            normalized_children = [mark_field_tree_optional(child) for child in normalized_children]
        node["properties"] = normalized_children
        relaxed_fields.append(node)
    return relaxed_fields


def normalize_mock_examples_against_contract(
    mock_examples: list[dict[str, object]],
    *,
    request_fields: list[dict[str, Any]],
    response_fields: list[dict[str, Any]],
) -> list[dict[str, object]]:
    normalized_examples: list[dict[str, object]] = []
    for item in mock_examples:
        normalized_item = dict(item)
        if "requestPayload" in normalized_item:
            normalized_item["requestPayload"] = normalize_payload_against_contract(normalized_item.get("requestPayload"), request_fields)
        if "responsePayload" in normalized_item:
            normalized_item["responsePayload"] = normalize_payload_against_contract(normalized_item.get("responsePayload"), response_fields)
        normalized_examples.append(normalized_item)
    return normalized_examples


def extract_mock_examples(rows: list[list[str]]) -> tuple[list[dict[str, object]], dict[str, object]]:
    example_rows = extract_section_rows(rows, start_title="範例", stop_titles=("For中台開發人員", "API 內部業務邏輯"))
    if not example_rows:
        return [], {}
    notes: list[str] = []
    header_index = next(
        (
            index
            for index, row in enumerate(example_rows)
            if normalize_token(first_non_empty_cell(row)) in {normalize_token("情境說明"), normalize_token("scenario")}
        ),
        None,
    )
    if header_index is None:
        return [], {"mockExampleNotes": [join_non_empty_cells(row) for row in example_rows if any(row)]}

    for row in example_rows[:header_index]:
        if any(row):
            notes.append(join_non_empty_cells(row))

    headers = example_rows[header_index]
    scenario_index = find_header_index(headers, "情境說明", "scenario")
    request_index = find_header_index(headers, "request")
    response_index = find_header_index(headers, "response")
    unparsed_examples: list[dict[str, object]] = []
    mock_examples: list[dict[str, object]] = []
    for row in example_rows[header_index + 1 :]:
        if not any(row):
            continue
        scenario = clean_text(row[scenario_index]) if scenario_index is not None and scenario_index < len(row) else first_non_empty_cell(row)
        if not scenario:
            continue
        payload: dict[str, object] = {"scenario": scenario}
        has_structured_payload = False
        if request_index is not None and request_index < len(row):
            request_text = clean_text(row[request_index])
            request_payload = try_parse_json_payload(request_text)
            if request_payload is not None:
                payload["requestPayload"] = request_payload
                has_structured_payload = True
            elif request_text:
                unparsed_examples.append({"scenario": scenario, "requestPayloadText": request_text})
        if response_index is not None and response_index < len(row):
            response_text = clean_text(row[response_index])
            response_payload = try_parse_json_payload(response_text)
            if response_payload is not None:
                payload["responsePayload"] = response_payload
                has_structured_payload = True
            elif response_text:
                unparsed_examples.append({"scenario": scenario, "responsePayloadText": response_text})
        if has_structured_payload:
            mock_examples.append(payload)

    external: dict[str, object] = {}
    if notes:
        external["mockExampleNotes"] = notes
    if unparsed_examples:
        external["unparsedMockExamples"] = unparsed_examples
    return mock_examples, external


def normalize_example_text(item: dict[str, object]) -> str:
    payload = item.get("responsePayload")
    parts = [clean_text(item.get("scenario"))]
    if isinstance(payload, dict):
        parts.append(clean_text(payload.get("responseCode")))
        parts.append(clean_text(payload.get("responseMessage")))
    return "\n".join(part for part in parts if part)


def classify_failure_example(item: dict[str, object]) -> str:
    token = normalize_token(normalize_example_text(item))
    if any(marker in token for marker in ("颱風", "休市", "營業時間", "營業日", "交易時間")):
        return "time_window"
    if any(marker in token for marker in ("foreignstop", "foreignalert", "外幣帳戶", "化整為零", "匯")):
        return "business_guard"
    if any(marker in token for marker in ("查無資料", "未輸入", "必填", "參數", "查詢")):
        return "data_validation"
    return "other_failure"


def is_success_example(item: dict[str, object]) -> bool:
    payload = item.get("responsePayload")
    if not isinstance(payload, dict):
        return False
    if payload.get("isSuccess") is True:
        return True
    return clean_text(payload.get("responseCode")) == "0000"


def dedupe_mock_examples_by_code(mock_examples: list[dict[str, object]]) -> list[dict[str, object]]:
    deduped: list[dict[str, object]] = []
    seen_codes: set[str] = set()
    seen_scenarios: set[str] = set()
    for item in mock_examples:
        payload = item.get("responsePayload")
        code = clean_text(payload.get("responseCode")) if isinstance(payload, dict) else ""
        scenario = clean_text(item.get("scenario"))
        unique_key = code or scenario
        if not unique_key:
            continue
        if code:
            if code in seen_codes:
                continue
            seen_codes.add(code)
        elif scenario in seen_scenarios:
            continue
        seen_scenarios.add(scenario)
        deduped.append(item)
    return deduped


def select_representative_mock_examples(mock_examples: list[dict[str, object]]) -> list[dict[str, object]]:
    selected: list[dict[str, object]] = []
    seen_examples: set[str] = set()
    for item in mock_examples:
        scenario = clean_text(item.get("scenario"))
        if not scenario:
            continue
        key = json.dumps(
            {
                "scenario": scenario,
                "requestPayload": item.get("requestPayload"),
                "responsePayload": item.get("responsePayload"),
            },
            ensure_ascii=False,
            sort_keys=True,
            default=str,
        )
        if key in seen_examples:
            continue
        seen_examples.add(key)
        selected.append(item)
    return selected


def build_error_code_rules_from_examples(mock_examples: list[dict[str, object]]) -> list[dict[str, str]]:
    rules: list[dict[str, str]] = []
    seen_rules: set[tuple[str, str, str]] = set()
    for item in mock_examples:
        payload = item.get("responsePayload")
        if not isinstance(payload, dict):
            continue
        code = clean_text(payload.get("responseCode"))
        scenario = clean_text(item.get("scenario"))
        message = clean_text(payload.get("responseMessage"))
        rule_key = (code, scenario, message)
        if not code or code == "0000" or rule_key in seen_rules:
            continue
        rules.append(
            {
                "code": code,
                "scenario": scenario,
                "message": message,
            }
        )
        seen_rules.add(rule_key)
    return rules


def merge_error_code_rules(*rule_sets: list[dict[str, str]]) -> list[dict[str, str]]:
    merged: list[dict[str, str]] = []
    seen_rules: set[tuple[str, str, str]] = set()
    for rule_set in rule_sets:
        for item in rule_set:
            code = clean_text(item.get("code"))
            scenario = clean_text(item.get("scenario") or item.get("condition"))
            message = clean_text(item.get("message"))
            rule_key = (code, scenario, message)
            if not code or code == "0000" or rule_key in seen_rules:
                continue
            merged.append(
                {
                    "code": code,
                    "scenario": scenario,
                    "message": message,
                }
            )
            seen_rules.add(rule_key)
    return merged


def add_backend_api_target(backend_apis: dict[str, list[str]], system: object, target: object) -> None:
    system_text = clean_text(system)
    target_text = clean_text(target)
    if not system_text or not target_text:
        return
    backend_apis.setdefault(system_text, [])
    if target_text not in backend_apis[system_text]:
        backend_apis[system_text].append(target_text)


def extract_backend_apis(logic_entries: list[dict[str, str]]) -> tuple[dict[str, list[str]], list[str]]:
    backend_apis: dict[str, list[str]] = {}
    unmatched_lines: list[str] = []
    for entry in logic_entries:
        if normalize_token("backendapi") not in normalize_token(entry["key"]):
            continue
        current_system = ""
        for raw_line in clean_text(entry["value"]).split("\n"):
            line = re.sub(r"^涉及BackendAPI[:：]?\s*", "", raw_line.strip(), flags=re.IGNORECASE).strip()
            if not line:
                continue
            match = BACKEND_API_LINE_RE.match(line)
            if match:
                system = clean_text(match.group("system"))
                target = clean_text(match.group("target"))
                if not system or not target:
                    unmatched_lines.append(line)
                    continue
                if dependency_is_uncertain(target):
                    unmatched_lines.append(line)
                    current_system = system
                    continue
                add_backend_api_target(backend_apis, system, target)
                current_system = system
                continue
            continuation_match = BACKEND_API_CONTINUATION_RE.match(line)
            if continuation_match and current_system:
                if dependency_is_uncertain(continuation_match.group("target")):
                    unmatched_lines.append(line)
                    continue
                add_backend_api_target(backend_apis, current_system, continuation_match.group("target"))
                continue
            dependency_matches = list(BUSINESS_DEPENDENCY_RE.finditer(line))
            if dependency_matches:
                for dependency_match in dependency_matches:
                    system = clean_text(dependency_match.group("system"))
                    target = clean_text(dependency_match.group("target"))
                    if dependency_is_uncertain(target):
                        unmatched_lines.append(line)
                        current_system = system or current_system
                        continue
                    add_backend_api_target(backend_apis, system, target)
                    current_system = system or current_system
                continue
            if not match:
                unmatched_lines.append(line)
                continue
    return backend_apis, unmatched_lines


def normalize_dependency_id(*parts: object) -> str:
    return re.sub(r"[^a-z0-9]+", "_", normalize_token("_".join(clean_text(part) for part in parts if clean_text(part)))).strip("_")


def dependency_is_uncertain(text: object) -> bool:
    return UNCERTAIN_DEPENDENCY_RE.search(clean_text(text)) is not None


def add_runtime_dependency(
    dependencies: list[dict[str, object]],
    *,
    dependency_id: str,
    dependency_type: str,
    description: str,
    authority: str = "backend_contract",
    required: bool = True,
) -> None:
    if not dependency_id or not description:
        return
    dependencies.append(
        {
            "id": dependency_id,
            "type": dependency_type,
            "authority": authority,
            "required": required,
            "description": description,
        }
    )


def dependency_type_for_system(system: object, target: object) -> str:
    normalized_system = normalize_token(clean_text(system))
    normalized_target = normalize_token(clean_text(target))
    if "db" in normalized_system:
        return "sql_table"
    if normalized_system in {"redis", "jwt", "header"}:
        return "runtime_context"
    if "mail" in normalized_target or "push" in normalized_target:
        return "notification"
    return "external_api"


def append_runtime_dependencies_from_backend_apis(
    dependencies: list[dict[str, object]],
    backend_apis: dict[str, list[str]],
) -> None:
    for system, targets in backend_apis.items():
        for target in targets:
            dep_type = dependency_type_for_system(system, target)
            add_runtime_dependency(
                dependencies,
                dependency_id=normalize_dependency_id(dep_type, system, target),
                dependency_type=dep_type,
                description=f"Excel 依赖声明提到 {clean_text(system)}->{clean_text(target)}。",
            )


def extract_inline_business_dependencies(text: str) -> tuple[list[dict[str, object]], list[dict[str, Any]]]:
    dependencies: list[dict[str, object]] = []
    unresolved: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()

    def remember(system: object, target: object, *, source_text: object) -> None:
        system_text = clean_text(system)
        target_text = clean_text(target)
        if not system_text or not target_text:
            return
        normalized_system = normalize_token(system_text)
        key = (normalized_system, normalize_token(target_text))
        if key in seen:
            return
        seen.add(key)
        uncertain = dependency_is_uncertain(source_text)
        if uncertain:
            unresolved.append(
                {
                    "topic": f"dependency.{normalize_dependency_id(system_text, target_text) or 'unknown'}",
                    "reason": f"业务逻辑正文提到依赖 {system_text}->{target_text}，但同一段存在 todo / 待确认语义，不能作为已确认实现或 fixture 需求。",
                    "blocking": True,
                    "blockedReason": "dependency contract unresolved",
                    "missingFacts": [
                        f"{system_text}->{target_text} 是否仍需接线",
                        "该依赖的权威来源、调用条件与失败处理",
                    ],
                    "suggestedOwner": "spec",
                    "nextDecisionNeeded": f"确认 {system_text}->{target_text} 是实现必需依赖、可忽略旧逻辑，还是仅作为历史备注。",
                }
            )
            return
        dep_type = dependency_type_for_system(system_text, target_text)
        add_runtime_dependency(
            dependencies,
            dependency_id=normalize_dependency_id(dep_type, system_text, target_text),
            dependency_type=dep_type,
            description=f"业务逻辑正文提到 {system_text}->{target_text} 依赖。",
        )

    for match in BUSINESS_DEPENDENCY_RE.finditer(text):
        remember(match.group("system"), match.group("target"), source_text=match.group(0))
    for match in METHOD_CALL_RE.finditer(text):
        remember(match.group("owner"), f"{match.group('method')}()", source_text=match.group(0))
    for match in IRIS_CODE_RE.finditer(text):
        remember("IRIS", match.group("code"), source_text=match.group(0))
    for match in MAIL_PUSH_RE.finditer(text):
        remember("CommonFunc", f"{match.group('name')}()", source_text=match.group(0))
    return dependencies, unresolved


def merge_backend_apis_from_business_text(backend_apis: dict[str, list[str]], logic_entries: list[dict[str, str]]) -> None:
    for entry in logic_entries:
        entry_text = "\n".join(part for part in [clean_text(entry.get("key")), clean_text(entry.get("value"))] if part)
        if not entry_text or dependency_is_uncertain(entry_text):
            continue
        for match in BUSINESS_DEPENDENCY_RE.finditer(entry_text):
            add_backend_api_target(backend_apis, match.group("system"), match.group("target"))
        for match in METHOD_CALL_RE.finditer(entry_text):
            add_backend_api_target(backend_apis, match.group("owner"), f"{match.group('method')}()")
        for match in IRIS_CODE_RE.finditer(entry_text):
            add_backend_api_target(backend_apis, "IRIS", match.group("code"))
        for match in MAIL_PUSH_RE.finditer(entry_text):
            add_backend_api_target(backend_apis, "CommonFunc", f"{match.group('name')}()")
        for table_name in extract_table_names(entry_text):
            add_backend_api_target(backend_apis, "DB", table_name)


def extract_field_mapping_rules(rows: list[list[str]]) -> list[dict[str, str]]:
    rules: list[dict[str, str]] = []
    in_response_table = False
    for row in rows:
        first = first_non_empty_cell(row)
        if normalize_token(first) == normalize_token("Response"):
            in_response_table = True
            continue
        if not in_response_table:
            continue
        if first.startswith("5.") and len(row) >= 7:
            field = clean_text(row[1])
            source = clean_text(row[6])
            rule = clean_text(row[4])
            if not field or not any((source, rule)):
                continue
            rules.append({"rowKey": clean_text(row[0]), "field": field, "source": source, "rule": rule})
            continue
        if first.startswith("範例"):
            break
    return rules


def extract_error_code_rules(rows: list[list[str]]) -> list[dict[str, str]]:
    rules: list[dict[str, str]] = []
    code_pattern = re.compile(r'"responseCode"\s*:\s*"(?P<code>\d+)"', re.IGNORECASE)
    message_pattern = re.compile(r'"responseMessage"\s*:\s*"(?P<message>[^"]*)"', re.IGNORECASE)
    for row in rows:
        scenario = first_non_empty_cell(row)
        payload = row[3] if len(row) > 3 else ""
        if not scenario or not payload:
            continue
        code_match = code_pattern.search(payload)
        if not code_match:
            continue
        message_match = message_pattern.search(payload)
        rules.append(
            {
                "code": code_match.group("code"),
                "scenario": scenario,
                "message": message_match.group("message") if message_match else "",
            }
        )
    return rules


def extract_source_mapping_rules(rows: list[list[str]]) -> list[dict[str, str]]:
    mappings: list[dict[str, str]] = []
    in_mapping_table = False
    for row in rows:
        first = first_non_empty_cell(row)
        if normalize_token(first) == normalize_token("新網銀顯示文字"):
            in_mapping_table = True
            continue
        if not in_mapping_table:
            continue
        if normalize_token(first) == normalize_token("source") or not first:
            continue
        source = clean_text(row[1]) if len(row) > 1 else ""
        if not source:
            continue
        mappings.append(
            {
                "title": first,
                "source": source,
                "loginWay": clean_text(row[2]) if len(row) > 2 else "",
                "loginType": clean_text(row[3]) if len(row) > 3 else "",
            }
        )
    return mappings


def extract_business_logic_section(rows: list[list[str]]) -> list[dict[str, str]]:
    entries: list[dict[str, str]] = []
    in_section = False
    for row in rows:
        first = first_non_empty_cell(row)
        if normalize_token(first) == normalize_token("API 內部業務邏輯"):
            in_section = True
            continue
        if not in_section:
            continue
        if normalize_token(first) == normalize_token("登入通路說明"):
            break
        if normalize_token(first) == normalize_token("#"):
            continue
        key = first_non_empty_cell(row)
        value = second_non_empty_cell(row) or join_non_empty_cells(row)
        if key and value:
            entries.append({"key": key, "value": value})
    return entries


def normalize_lookup_login_type(value: str) -> str:
    cleaned = clean_text(value)
    if not cleaned:
        return ""
    if "null" in cleaned.casefold():
        return ""
    if "/" in cleaned:
        return ""
    return cleaned


def build_lookup_tables(source_mappings: list[dict[str, str]]) -> list[dict[str, Any]]:
    if not source_mappings:
        return []

    entry_map: dict[str, dict[str, Any]] = {}
    entry_order: list[str] = []
    for mapping in source_mappings:
        key = normalize_mapping_source_text(clean_text(mapping.get("source")))
        title = clean_text(mapping.get("title")) or key or "lookup"
        if not key:
            continue
        entry = entry_map.setdefault(
            key,
            {
                "key": key,
                "title": title,
                "mappedValues": {},
                "rule": "",
            },
        )
        if key not in entry_order:
            entry_order.append(key)
        if not clean_text(entry.get("title")):
            entry["title"] = title
        mapped_values = dict(entry.get("mappedValues") or {})
        login_way = clean_text(mapping.get("loginWay"))
        login_type = normalize_lookup_login_type(clean_text(mapping.get("loginType")))
        if login_way:
            mapped_values["loginWay"] = login_way
        if login_type:
            mapped_values["loginType"] = login_type
        entry["mappedValues"] = mapped_values

    entries = [entry_map[key] for key in entry_order if entry_map[key].get("mappedValues") or entry_map[key].get("title")]
    if not entries:
        return []
    return [
        {
            "id": "source-lookup-table",
            "sourceField": "source",
            "description": "依 source 對照登入通路與登入方式等顯示值。",
            "entries": entries,
        }
    ]


def extract_business_logic_structure_from_rows(rows: list[list[str]]) -> dict[str, object]:
    logic_entries = extract_business_logic_section(rows)
    field_mapping_rules = extract_field_mapping_rules(rows)
    error_code_rules = extract_error_code_rules(rows)
    source_mappings = extract_source_mapping_rules(rows)
    data_sources: list[dict[str, object]] = []
    sql_specs: list[dict[str, object]] = []
    runtime_dependencies: list[dict[str, object]] = []
    dependency_unresolved: list[dict[str, Any]] = []
    legacy_references: list[dict[str, Any]] = []
    backend_apis, _ = extract_backend_apis(logic_entries)
    merge_backend_apis_from_business_text(backend_apis, logic_entries)
    append_runtime_dependencies_from_backend_apis(runtime_dependencies, backend_apis)
    for entry in logic_entries:
        key = entry["key"]
        value = entry["value"]
        normalized_key = normalize_token(key)
        entry_text = "\n".join(part for part in [key, value] if clean_text(part))
        inline_dependencies, inline_unresolved = extract_inline_business_dependencies(entry_text)
        runtime_dependencies.extend(inline_dependencies)
        dependency_unresolved.extend(inline_unresolved)

        table_names = extract_table_names(value)
        uncertain_dependency = dependency_is_uncertain(entry_text)
        if table_names and uncertain_dependency:
            for table_name in table_names:
                dependency_unresolved.append(
                    {
                        "topic": f"dependency.db.{normalize_dependency_id(table_name) or 'table'}",
                        "reason": f"业务逻辑正文提到 DB 表 {table_name}，但同一段存在 todo / 待确认语义，不能作为已确认 SQL fixture 需求。",
                        "blocking": True,
                        "blockedReason": "db dependency unresolved",
                        "missingFacts": [
                            f"{table_name} 是否仍需读写",
                            "若需要，该表的权威 schema、读写条件与 seed 场景",
                        ],
                        "suggestedOwner": "spec",
                        "nextDecisionNeeded": f"确认 {table_name} 是实现必需表、旧逻辑遗留表，还是无需接线。",
                    }
                )
        if table_names and not uncertain_dependency:
            for table_name in extract_table_names(value):
                data_sources.append({"name": table_name, "type": "sql_table", "authority": "backend_contract", "required": True})
        has_sql_text = bool(table_names) and any(
            re.search(rf"\b{keyword}\b", value, re.IGNORECASE)
            for keyword in ("SELECT", "INSERT", "UPDATE", "DELETE", "FROM", "JOIN", "INTO")
        )
        if ("sql" in normalized_key or has_sql_text) and not uncertain_dependency:
            sql_specs.append(
                {
                    "id": re.sub(r"[^a-z0-9]+", "_", normalized_key).strip("_") or "sql_spec",
                    "title": key,
                    "authority": "backend_contract",
                    "required": True,
                    "dataSources": extract_table_names(value),
                    "mustContain": build_sql_must_contain(value),
                    "queryText": value,
                }
            )
        if "redis" in normalize_token(value) and "custid" in normalize_token(value):
            runtime_dependencies.append(
                {
                    "id": "current_customer_context",
                    "type": "service",
                    "authority": "backend_contract",
                    "required": True,
                    "description": "Resolve keyId/custId from the current request context and Redis.",
                }
            )
        if ("sql" in normalized_key or table_names) and not uncertain_dependency:
            runtime_dependencies.append(
                {
                    "id": "mma_sql_connection",
                    "type": "service",
                    "authority": "backend_contract",
                    "required": True,
                    "description": describe_sql_runtime_dependency(value),
                }
            )
        if "舊代碼" in key or "参考旧代碼" in key or "參考舊代碼" in key:
            legacy_references.append(
                {
                    "id": build_handoff_identifier("legacy", key),
                    "title": key,
                    "kind": "legacyReference",
                    "origin": key,
                    "authority": "legacy_reference",
                    "nonAuthoritative": True,
                    "summary": summarize_snippet(value, fallback=key),
                    "snippet": clean_text(value),
                    "symbols": extract_symbol_candidates(value),
                }
            )

    deduped_data_sources = sorted({item["name"]: item for item in data_sources}.values(), key=lambda item: item["name"])
    deduped_runtime_dependencies = sorted(
        {item["id"]: item for item in runtime_dependencies}.values(),
        key=lambda item: item["id"],
    )
    prohibited_shortcuts: list[str | dict[str, object]] = []
    if sql_specs:
        prohibited_shortcuts.extend(["mock_response_payload", "hardcoded_custid"])
    prohibited_shortcuts.append("block_comment_header_substitution")
    return {
        "dataSources": deduped_data_sources,
        "sqlSpecs": sql_specs,
        "fieldMappings": [{**rule, "kind": "response_field", "authority": "backend_contract"} for rule in field_mapping_rules],
        "lookupTables": build_lookup_tables(source_mappings),
        "errorCodeRules": [{**rule, "authority": "backend_contract"} for rule in error_code_rules],
        "runtimeDependencies": deduped_runtime_dependencies,
        "dependencyUnresolved": dependency_unresolved,
        "prohibitedShortcuts": prohibited_shortcuts,
        "legacyReferences": legacy_references,
    }


def build_handoff_identifier(prefix: str, raw_value: object, *, fallback: str = "item") -> str:
    normalized = clean_text(raw_value)
    slug = re.sub(r"[^a-z0-9]+", "_", normalize_token(normalized)).strip("_")
    return f"{prefix}_{slug or fallback}"


def project_hard_constraints_path(agent_dir: Path) -> Path:
    return (agent_dir / "Common" / PROJECT_HARD_CONSTRAINTS_FILENAME).resolve()


def load_project_hard_constraints(agent_dir: Path | None) -> tuple[dict[str, Any] | None, Path | None]:
    if agent_dir is None:
        return None, None

    path = project_hard_constraints_path(agent_dir)
    if not path.exists():
        return None, None

    payload = load_json(path)
    if not isinstance(payload, dict):
        raise SkillError(
            f"{PROJECT_HARD_CONSTRAINTS_FILENAME} 必须是 JSON object：{path}",
            status="blocked",
        )

    schema = load_schema("project-hard-constraints.schema.json")
    validator = Draft202012Validator(schema)
    errors = sorted(validator.iter_errors(payload), key=lambda item: list(item.absolute_path))
    if errors:
        first = errors[0]
        raise SkillError(
            f"project-hard-constraints schema 校验失败：{format_validation_path(first.absolute_path)} - {first.message}",
            status="blocked",
        )
    return payload, path


def normalized_keyword_list(values: object) -> list[str]:
    if isinstance(values, str):
        candidates = [values]
    elif isinstance(values, list):
        candidates = [clean_text(value) for value in values]
    else:
        return []

    normalized_values: list[str] = []
    seen: set[str] = set()
    for candidate in candidates:
        normalized = normalize_token(clean_text(candidate))
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        normalized_values.append(normalized)
    return normalized_values


def build_project_constraint_reference_text(
    *,
    function_code: str | None,
    api_category: str,
    api_name: str,
    request_fields: list[dict[str, object]],
    response_fields: list[dict[str, object]],
    business_logic_payload: dict[str, object],
) -> str:
    parts: list[str] = [clean_text(function_code), clean_text(api_category), clean_text(api_name)]
    for field in walk_field_nodes(request_fields):
        parts.extend(
            [
                clean_text(field.get("fieldName")),
                clean_text(field.get("description")),
                clean_text(field.get("notes")),
            ]
        )
    for field in walk_field_nodes(response_fields):
        parts.extend(
            [
                clean_text(field.get("fieldName")),
                clean_text(field.get("description")),
                clean_text(field.get("notes")),
            ]
        )
    for step in list(business_logic_payload.get("steps") or []):
        if isinstance(step, dict):
            parts.extend([clean_text(step.get("title")), clean_text(step.get("details"))])
    for dependency in list(business_logic_payload.get("runtimeDependencies") or []):
        if isinstance(dependency, dict):
            parts.extend([clean_text(dependency.get("id")), clean_text(dependency.get("description"))])
    for source in list(business_logic_payload.get("dataSources") or []):
        if isinstance(source, dict):
            parts.extend([clean_text(source.get("name")), clean_text(source.get("type"))])
    for sql_spec in list(business_logic_payload.get("sqlSpecs") or []):
        if isinstance(sql_spec, dict):
            parts.extend(
                [
                    clean_text(sql_spec.get("title")),
                    clean_text(sql_spec.get("sqlText")),
                    clean_text(sql_spec.get("queryText")),
                ]
            )
    for shortcut in list(business_logic_payload.get("prohibitedShortcuts") or []):
        if isinstance(shortcut, dict):
            parts.append(json.dumps(shortcut, ensure_ascii=False, sort_keys=True))
        else:
            parts.append(clean_text(shortcut))
    return normalize_token("\n".join(part for part in parts if part))


def project_constraint_rule_matches(
    rule: dict[str, Any],
    *,
    function_code: str | None,
    api_category: str,
    api_name: str,
    normalized_reference_text: str,
) -> bool:
    conditions = rule.get("conditions")
    if not isinstance(conditions, dict):
        return True

    normalized_function_code = normalize_token(function_code or "")
    normalized_api_category = normalize_token(api_category)
    normalized_api_name = normalize_token(api_name)

    function_codes = normalized_keyword_list(conditions.get("functionCodes"))
    if function_codes and normalized_function_code not in function_codes:
        return False

    api_categories = normalized_keyword_list(conditions.get("apiCategories"))
    if api_categories and normalized_api_category not in api_categories:
        return False

    api_names = normalized_keyword_list(conditions.get("apiNames"))
    if api_names and normalized_api_name not in api_names:
        return False

    business_keywords = normalized_keyword_list(conditions.get("businessKeywordsAny"))
    if business_keywords and not any(keyword in normalized_reference_text for keyword in business_keywords):
        return False

    return True


def should_emit_project_unresolved(
    unresolved_payload: dict[str, Any],
    *,
    normalized_reference_text: str,
) -> bool:
    trigger_keywords = normalized_keyword_list(unresolved_payload.get("triggerKeywordsAny"))
    if trigger_keywords and not any(keyword in normalized_reference_text for keyword in trigger_keywords):
        return False

    satisfied_keywords = normalized_keyword_list(unresolved_payload.get("satisfiedByKeywordsAny"))
    if satisfied_keywords and any(keyword in normalized_reference_text for keyword in satisfied_keywords):
        return False

    return True


def build_project_policy_handoff(
    *,
    project_root: Path | None,
    agent_dir: Path | None,
    function_code: str | None,
    api_category: str,
    api_name: str,
    request_fields: list[dict[str, object]],
    response_fields: list[dict[str, object]],
    business_logic_payload: dict[str, object],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]]]:
    project_constraints, profile_path = load_project_hard_constraints(agent_dir)
    if not project_constraints or profile_path is None:
        return [], [], []

    profile_origin = normalize_persisted_path(profile_path, project_root=project_root) or profile_path.name
    normalized_reference_text = build_project_constraint_reference_text(
        function_code=function_code,
        api_category=api_category,
        api_name=api_name,
        request_fields=request_fields,
        response_fields=response_fields,
        business_logic_payload=business_logic_payload,
    )

    legacy_evidence: list[dict[str, Any]] = []
    constraints: list[dict[str, Any]] = []
    unresolved: list[dict[str, Any]] = []
    for rule in list(project_constraints.get("rules") or []):
        if not isinstance(rule, dict):
            continue
        if not project_constraint_rule_matches(
            rule,
            function_code=function_code,
            api_category=api_category,
            api_name=api_name,
            normalized_reference_text=normalized_reference_text,
        ):
            continue

        rule_id = clean_text(rule.get("ruleId")) or build_handoff_identifier("project_policy", rule.get("instruction"), fallback="rule")
        instruction = clean_text(rule.get("instruction"))
        if not instruction:
            continue

        evidence_id = build_handoff_identifier("project_policy", rule_id, fallback="rule")
        evidence_symbols = [
            symbol
            for symbol in [
                clean_text(rule.get("scope")),
                clean_text(rule.get("fileRole")),
                *[clean_text(value) for value in list(rule.get("appliesTo") or [])],
            ]
            if symbol
        ]
        legacy_evidence.append(
            {
                "evidenceId": evidence_id,
                "kind": "projectConstraint",
                "origin": f"{profile_origin}#{rule_id}",
                "authority": "project_hard_constraints",
                "symbols": evidence_symbols,
                "summary": summarize_snippet(instruction, fallback=rule_id),
                "snippet": instruction,
            }
        )
        constraints.append(
            {
                "constraintType": f"project_policy::{clean_text(rule.get('ruleType')) or 'rule'}",
                "rule": instruction,
                "severity": clean_text(rule.get("severity")) or "error",
                "evidenceIds": [evidence_id],
            }
        )

        unresolved_payload = rule.get("unresolved")
        if not isinstance(unresolved_payload, dict):
            continue
        if not should_emit_project_unresolved(unresolved_payload, normalized_reference_text=normalized_reference_text):
            continue

        unresolved_item = {
            "topic": clean_text(unresolved_payload.get("topic")) or f"project_policy.{rule_id}",
            "reason": clean_text(unresolved_payload.get("reason")) or instruction,
            "blocking": bool(unresolved_payload.get("blocking")),
        }
        for key in ("blockedReason", "suggestedOwner", "nextDecisionNeeded"):
            value = clean_text(unresolved_payload.get(key))
            if value:
                unresolved_item[key] = value
        missing_facts = [clean_text(value) for value in list(unresolved_payload.get("missingFacts") or []) if clean_text(value)]
        if missing_facts:
            unresolved_item["missingFacts"] = missing_facts
        unresolved.append(unresolved_item)

    return legacy_evidence, constraints, unresolved


def summarize_snippet(value: object, *, fallback: str) -> str:
    text = clean_text(value)
    if not text:
        return fallback
    first_line = clean_text(text.split("\n", 1)[0])
    return first_line or fallback


def extract_symbol_candidates(value: object) -> list[str]:
    text = clean_text(value)
    if not text:
        return []
    matches = re.findall(r"[A-Za-z_][A-Za-z0-9_.]{2,}", text)
    symbols: list[str] = []
    for match in matches:
        if match not in symbols:
            symbols.append(match)
        if len(symbols) >= 8:
            break
    return symbols


def extract_parameter_hints(query_text: object) -> list[str]:
    text = clean_text(query_text)
    if not text:
        return []
    hints: list[str] = []
    for match in re.findall(r"@[A-Za-z0-9_]+", text):
        if match not in hints:
            hints.append(match)
    return hints


def build_sql_must_contain(query_text: object) -> list[str]:
    text = clean_text(query_text)
    if not text:
        return []

    hints: list[str] = []
    seen: set[str] = set()

    def remember(value: str) -> None:
        cleaned = clean_text(value)
        if not cleaned:
            return
        normalized = cleaned.upper()
        if normalized in seen:
            return
        seen.add(normalized)
        hints.append(cleaned)

    for table_name in extract_table_names(text):
        remember(table_name)
    for parameter in extract_parameter_hints(text):
        remember(parameter)

    upper_text = text.upper()
    for keyword in ("SELECT", "INSERT", "UPDATE", "DELETE", "FROM", "WHERE", "SET", "VALUES", "ORDER BY", "TOP"):
        if keyword == "ORDER BY":
            present = "ORDER BY" in upper_text
        else:
            present = re.search(rf"\b{re.escape(keyword)}\b", upper_text) is not None
        if present:
            remember(keyword)
    return hints


def describe_sql_runtime_dependency(query_text: object) -> str:
    table_names = extract_table_names(clean_text(query_text))
    if table_names:
        preview = ", ".join(table_names[:3])
        return f"Execute authoritative SQL declared in the workbook for {preview}."
    return "Execute authoritative SQL declared in the workbook."


def normalize_mapping_rule_value(value: object) -> str | None:
    text = clean_text(value)
    return text or None


def guess_logic_action_type(title: str, details: str) -> str:
    haystack = f"{clean_text(title)} {clean_text(details)}".casefold()
    if "sql" in haystack or "查詢" in haystack or "query" in haystack or "select" in haystack:
        return "query"
    if "mapping" in haystack or "轉" in haystack or "對應" in haystack:
        return "mapping"
    if "驗證" in haystack or "check" in haystack or "validate" in haystack:
        return "validation"
    if "回傳" in haystack or "response" in haystack or "return" in haystack:
        return "return"
    if "舊代碼" in haystack or "legacy" in haystack:
        return "legacy_reference"
    if "api" in haystack:
        return "external_call"
    return "process"


def find_matching_error_rule(
    error_code_rules: list[dict[str, Any]],
    *,
    predicate: callable,
) -> dict[str, Any] | None:
    for rule in error_code_rules:
        if predicate(rule):
            return rule
    return None


def extract_text_element_limit(text: str) -> int | None:
    match = re.search(r"最多\s*(\d+)\s*(?:個)?(?:字元|字符|字|位)", clean_text(text))
    if not match:
        return None
    try:
        return int(match.group(1))
    except ValueError:
        return None


def build_request_validation_constraints(
    *,
    request_fields: list[dict[str, object]],
    business_steps: list[dict[str, Any]],
    error_code_rules: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    constraints: list[dict[str, Any]] = []
    unresolved: list[dict[str, Any]] = []
    business_text = "\n".join(
        clean_text(step.get("title")) + "\n" + clean_text(step.get("details"))
        for step in business_steps
        if isinstance(step, dict)
    )
    for field in walk_field_nodes(request_fields):
        field_name = clean_text(field.get("fieldName"))
        if not field_name:
            continue
        field_description = clean_text(field.get("description"))
        field_notes = clean_text(field.get("notes"))
        field_text = "\n".join(part for part in [field_description, field_notes, business_text] if part)

        if bool(field.get("required")):
            required_rule = find_matching_error_rule(
                error_code_rules,
                predicate=lambda rule: (
                    "必填" in clean_text(rule.get("scenario"))
                    or "未輸入" in clean_text(rule.get("scenario"))
                    or "請輸入" in clean_text(rule.get("message"))
                ),
            )
            if required_rule is None:
                unresolved.append(
                    {
                        "topic": f"request_validation.{field_name}.required",
                        "reason": f"欄位 {field_name} 為必填，但未找到對應的 spec 錯誤碼/訊息。",
                        "blocking": True,
                    }
                )
            else:
                constraints.append(
                    {
                        "constraintType": "request_field_validation",
                        "field": field_name,
                        "validationLayer": "dto_attribute",
                        "validationType": "required",
                        "expectedCode": clean_text(required_rule.get("code")),
                        "expectedMessage": clean_text(required_rule.get("message")),
                        "customValidationAttributeNeeded": False,
                        "rule": (
                            f"{field_name} | dto_attribute | required | "
                            f"{clean_text(required_rule.get('code'))} | {clean_text(required_rule.get('message'))}"
                        ),
                        "severity": "warning",
                        "evidenceIds": [],
                    }
                )

        text_element_limit = extract_text_element_limit(field_text)
        if text_element_limit is not None:
            length_rule = find_matching_error_rule(
                error_code_rules,
                predicate=lambda rule: (
                    str(text_element_limit) in clean_text(rule.get("scenario"))
                    or str(text_element_limit) in clean_text(rule.get("message"))
                )
                and any(
                    keyword in f"{clean_text(rule.get('scenario'))} {clean_text(rule.get('message'))}"
                    for keyword in ("最多", "超出", "字元", "字符", "位")
                ),
            )
            if length_rule is None:
                unresolved.append(
                    {
                        "topic": f"request_validation.{field_name}.text_element_max_length",
                        "reason": (
                            f"欄位 {field_name} 存在最多 {text_element_limit} 個文字元素規則，"
                            "但未找到對應的 spec 錯誤碼/訊息。"
                        ),
                        "blocking": True,
                    }
                )
            else:
                constraints.append(
                    {
                        "constraintType": "request_field_validation",
                        "field": field_name,
                        "validationLayer": "dto_attribute",
                        "validationType": "text_element_max_length",
                        "expectedCode": clean_text(length_rule.get("code")),
                        "expectedMessage": clean_text(length_rule.get("message")),
                        "customValidationAttributeNeeded": True,
                        "maxTextElements": text_element_limit,
                        "rule": (
                            f"{field_name} | dto_attribute | text_element_max_length({text_element_limit}) | "
                            f"{clean_text(length_rule.get('code'))} | {clean_text(length_rule.get('message'))}"
                        ),
                        "severity": "warning",
                        "evidenceIds": [],
                    }
                )

    return constraints, unresolved


def build_runtime_dependency_unresolved(
    *,
    runtime_dependencies: list[dict[str, Any]],
    business_steps: list[dict[str, Any]],
    request_paths: list[str],
    identity_contract_available: bool = False,
) -> list[dict[str, Any]]:
    unresolved: list[dict[str, Any]] = []
    normalized_request_paths = {normalize_token(path) for path in request_paths if clean_text(path)}
    if any(token in normalized_request_paths for token in {"custid", "cust_id"}):
        return unresolved

    dependency_ids = {normalize_token(clean_text(item.get("id"))) for item in runtime_dependencies if isinstance(item, dict)}
    if "current_customer_context" not in dependency_ids:
        return unresolved
    if identity_contract_available:
        return unresolved

    business_text = "\n".join(
        part
        for step in business_steps
        if isinstance(step, dict)
        for part in (clean_text(step.get("title")), clean_text(step.get("details")))
        if part
    )
    normalized_business_text = normalize_token(business_text)
    source_contract_markers = (
        "auth_sn",
        "memberhash",
        "member_hash",
        "access_token",
        "accesstoken",
        "authorization",
        "bearer",
        "claim",
    )
    if not any(marker in normalized_business_text for marker in source_contract_markers):
        return unresolved

    unresolved.append(
        {
            "topic": "identity.current_customer_context",
            "reason": (
                "Workbook references authenticated/session identity details, but the business-logic sheet does not "
                "fully define the request-side identity contract."
            ),
            "blocking": True,
            "blockedReason": "identity model blocked",
            "missingFacts": [
                "authenticated identity source for the current request",
                "session scope / lifecycle for the resolved customer context",
                "exact Redis key or claim chain that turns current login state into CustId",
            ],
            "suggestedOwner": "upstream auth",
            "nextDecisionNeeded": "Confirm how the authenticated request resolves to CustId before code implementation.",
        }
    )
    return unresolved


def build_code_handoff(
    *,
    project_root: Path | None,
    agent_dir: Path | None,
    function_code: str | None,
    api_category: str,
    api_name: str,
    request_fields: list[dict[str, object]],
    response_fields: list[dict[str, object]],
    business_logic_payload: dict[str, object],
    additional_unresolved: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    business_steps = [step for step in list(business_logic_payload.get("steps") or []) if isinstance(step, dict)]
    field_mappings = [mapping for mapping in list(business_logic_payload.get("fieldMappings") or []) if isinstance(mapping, dict)]
    lookup_tables = [table for table in list(business_logic_payload.get("lookupTables") or []) if isinstance(table, dict)]
    runtime_dependencies = [dependency for dependency in list(business_logic_payload.get("runtimeDependencies") or []) if isinstance(dependency, dict)]
    data_sources = [source for source in list(business_logic_payload.get("dataSources") or []) if isinstance(source, dict)]
    sql_specs = [spec for spec in list(business_logic_payload.get("sqlSpecs") or []) if isinstance(spec, dict)]
    legacy_references = [reference for reference in list(business_logic_payload.get("legacyReferences") or []) if isinstance(reference, dict)]
    prohibited_shortcuts = list(business_logic_payload.get("prohibitedShortcuts") or [])
    error_code_rules = [rule for rule in list(business_logic_payload.get("errorCodeRules") or []) if isinstance(rule, dict)]

    response_paths = [clean_text(field.get("fieldName")) for field in walk_field_nodes(response_fields) if clean_text(field.get("fieldName"))]
    request_paths = [clean_text(field.get("fieldName")) for field in walk_field_nodes(request_fields) if clean_text(field.get("fieldName"))]

    legacy_evidence: list[dict[str, Any]] = []
    evidence_ids: set[str] = set()
    for reference in legacy_references:
        snippet = clean_text(reference.get("snippet"))
        if not snippet:
            continue
        evidence_id = clean_text(reference.get("id")) or build_handoff_identifier("legacy", reference.get("title"), fallback="reference")
        if evidence_id in evidence_ids:
            continue
        legacy_evidence.append(
            {
                "evidenceId": evidence_id,
                "kind": clean_text(reference.get("kind")) or "legacyReference",
                "origin": clean_text(reference.get("origin") or reference.get("title")) or "legacy_reference",
                "authority": clean_text(reference.get("authority")) or "legacy_reference",
                "symbols": [symbol for symbol in list(reference.get("symbols") or []) if clean_text(symbol)],
                "summary": clean_text(reference.get("summary")) or summarize_snippet(snippet, fallback="legacy reference"),
                "snippet": snippet,
            }
        )
        evidence_ids.add(evidence_id)

    project_legacy_evidence, project_constraints, project_unresolved = build_project_policy_handoff(
        project_root=project_root,
        agent_dir=agent_dir,
        function_code=function_code,
        api_category=api_category,
        api_name=api_name,
        request_fields=request_fields,
        response_fields=response_fields,
        business_logic_payload=business_logic_payload,
    )
    for evidence in project_legacy_evidence:
        evidence_id = clean_text(evidence.get("evidenceId"))
        if not evidence_id or evidence_id in evidence_ids:
            continue
        legacy_evidence.append(evidence)
        evidence_ids.add(evidence_id)

    for step in business_steps:
        title = clean_text(step.get("title"))
        details = clean_text(step.get("details"))
        if "舊代碼" not in title and "legacy" not in title.casefold() and "舊代碼" not in details:
            continue
        snippet = details
        if not snippet:
            continue
        evidence_id = build_handoff_identifier("legacy", title or step.get("step"), fallback="step")
        if evidence_id in evidence_ids:
            continue
        legacy_evidence.append(
            {
                "evidenceId": evidence_id,
                "kind": "legacyStep",
                "origin": title or f"step-{clean_text(step.get('step'))}",
                "authority": "legacy_step_detail",
                "symbols": extract_symbol_candidates(snippet),
                "summary": summarize_snippet(snippet, fallback=title or "legacy step"),
                "snippet": snippet,
            }
        )
        evidence_ids.add(evidence_id)

    query_contracts: list[dict[str, Any]] = []
    for sql_spec in sql_specs:
        contract_id = clean_text(sql_spec.get("id")) or build_handoff_identifier("query", sql_spec.get("title"), fallback="contract")
        sql_text = clean_text(sql_spec.get("sqlText") or sql_spec.get("queryText"))
        if not contract_id or not sql_text:
            continue
        query_contracts.append(
            {
                "contractId": contract_id,
                "purpose": clean_text(sql_spec.get("title")) or contract_id,
                "dataSources": [clean_text(value) for value in list(sql_spec.get("dataSources") or []) if clean_text(value)],
                "mustContain": [clean_text(value) for value in list(sql_spec.get("mustContain") or []) if clean_text(value)],
                "sqlText": sql_text,
                "parameterHints": extract_parameter_hints(sql_text),
                "resultShape": response_paths,
                "evidenceIds": [],
            }
        )

    mapping_rules: list[dict[str, Any]] = []
    for mapping in field_mappings:
        fields = [field for field in list(mapping.get("fields") or []) if isinstance(field, dict)]
        if fields:
            for field in fields:
                target_field = clean_text(field.get("field"))
                source_field = clean_text(field.get("source") or mapping.get("source"))
                if not target_field or not source_field:
                    continue
                mapping_rules.append(
                    {
                        "ruleId": build_handoff_identifier("map", f"{source_field}_{target_field}", fallback="field"),
                        "sourceField": source_field,
                        "targetField": target_field,
                        "mappingType": "field_mapping",
                        "mappingTable": None,
                        "defaultValue": normalize_mapping_rule_value(field.get("rule") or mapping.get("rule")),
                        "evidenceIds": [],
                    }
                )
        else:
            target_field = clean_text(mapping.get("target"))
            source_field = clean_text(mapping.get("source"))
            if not target_field or not source_field:
                continue
            mapping_rules.append(
                {
                    "ruleId": build_handoff_identifier("map", f"{source_field}_{target_field}", fallback="field"),
                    "sourceField": source_field,
                    "targetField": target_field,
                    "mappingType": "field_mapping",
                    "mappingTable": None,
                    "defaultValue": normalize_mapping_rule_value(mapping.get("rule")),
                    "evidenceIds": [],
                }
            )

    for lookup_table in lookup_tables:
        source_field = clean_text(lookup_table.get("sourceField"))
        entries = [entry for entry in list(lookup_table.get("entries") or []) if isinstance(entry, dict)]
        target_fields = sorted(
            {
                clean_text(target_field)
                for entry in entries
                for target_field in (entry.get("mappedValues") or {}).keys()
                if clean_text(target_field)
            }
        )
        for target_field in target_fields:
            table_payload = {
                clean_text(entry.get("key")): normalize_mapping_rule_value((entry.get("mappedValues") or {}).get(target_field))
                for entry in entries
                if clean_text(entry.get("key")) and normalize_mapping_rule_value((entry.get("mappedValues") or {}).get(target_field)) is not None
            }
            if not source_field or not target_field or not table_payload:
                continue
            mapping_rules.append(
                {
                    "ruleId": build_handoff_identifier("lookup", f"{lookup_table.get('id')}_{target_field}", fallback="table"),
                    "sourceField": source_field,
                    "targetField": target_field,
                    "mappingType": "lookup_table",
                    "mappingTable": table_payload,
                    "defaultValue": None,
                    "evidenceIds": [],
                }
            )

    dependency_hints: list[dict[str, Any]] = []
    for dependency in runtime_dependencies:
        dependency_id = clean_text(dependency.get("id"))
        preferred: list[str] = []
        lowered = dependency_id.casefold()
        if "sql" in lowered:
            preferred.append("ISqlQueryExecutor")
        if "runtime" in lowered or "context" in lowered or "cust" in lowered:
            preferred.append("ICurrentRuntimeContextAccessor")
            preferred.append("IRedisService")
        if "api" in lowered:
            preferred.append("IApiRequestService")
        if not preferred:
            preferred.append("FrameworkProvidedDependency")
        dependency_hints.append(
            {
                "dependencyType": clean_text(dependency.get("type")) or dependency_id or "dependency",
                "preferredAbstractions": preferred,
                "purpose": clean_text(dependency.get("description")) or dependency_id or "dependency",
                "evidenceIds": [],
            }
        )
    for source in data_sources:
        source_name = clean_text(source.get("name"))
        source_type = clean_text(source.get("type")) or "data_source"
        if not source_name:
            continue
        preferred = ["ISqlQueryExecutor"] if "sql" in source_type.casefold() else ["FrameworkProvidedDependency"]
        dependency_hints.append(
            {
                "dependencyType": source_type,
                "preferredAbstractions": preferred,
                "purpose": f"Read from {source_name}",
                "evidenceIds": [],
            }
        )

    constraints: list[dict[str, Any]] = []
    identity_contract_available = False
    for shortcut in prohibited_shortcuts:
        rule = clean_text(shortcut if isinstance(shortcut, str) else json.dumps(shortcut, ensure_ascii=False, sort_keys=True))
        if not rule:
            continue
        constraints.append(
            {
                "constraintType": "prohibited_shortcut",
                "rule": rule,
                "severity": "error",
                "evidenceIds": [],
            }
        )
    for error_rule in error_code_rules:
        code = clean_text(error_rule.get("code"))
        scenario = clean_text(error_rule.get("scenario"))
        message = clean_text(error_rule.get("message"))
        if not code:
            continue
        constraints.append(
            {
                "constraintType": "error_code_rule",
                "rule": f"{code} | {scenario or 'n/a'} | {message or 'n/a'}",
                "severity": "warning",
                "evidenceIds": [],
            }
        )

    validation_constraints, validation_unresolved = build_request_validation_constraints(
        request_fields=request_fields,
        business_steps=business_steps,
        error_code_rules=error_code_rules,
    )
    constraints.extend(validation_constraints)
    constraints.extend(project_constraints)

    unresolved: list[dict[str, Any]] = list(validation_unresolved)
    unresolved.extend(
        build_runtime_dependency_unresolved(
            runtime_dependencies=runtime_dependencies,
            business_steps=business_steps,
            request_paths=request_paths,
            identity_contract_available=identity_contract_available,
        )
    )
    unresolved.extend(project_unresolved)
    unresolved.extend(item for item in list(additional_unresolved or []) if isinstance(item, dict))
    if any(("舊代碼" in clean_text(step.get("title")) or "legacy" in clean_text(step.get("title")).casefold()) for step in business_steps) and not legacy_evidence:
        unresolved.append(
            {
                "topic": "legacy_logic",
                "reason": "Detected legacy-reference steps, but no structured legacy evidence snippet could be extracted.",
                "blocking": True,
            }
        )

    logic_flow: list[dict[str, Any]] = []
    legacy_evidence_ids = [entry["evidenceId"] for entry in legacy_evidence]
    request_preview = request_paths[:4]
    response_preview = response_paths[:4]
    for index, step in enumerate(business_steps, start=1):
        title = clean_text(step.get("title")) or f"step-{index}"
        details = clean_text(step.get("details"))
        logic_flow.append(
            {
                "stepId": f"step_{clean_text(step.get('step')) or index}",
                "title": title,
                "actionType": guess_logic_action_type(title, details),
                "inputs": request_preview,
                "outputs": response_preview,
                "evidenceIds": legacy_evidence_ids if ("舊代碼" in title or "legacy" in title.casefold()) else [],
            }
        )

    return {
        "schemaVersion": "1.0.0",
        "logicSummary": {
            "stepCount": len(logic_flow),
            "queryContractCount": len(query_contracts),
            "mappingRuleCount": len(mapping_rules),
            "legacyEvidenceCount": len(legacy_evidence),
            "dependencyHintCount": len(dependency_hints),
            "constraintCount": len(constraints),
            "unresolvedCount": len(unresolved),
            "primarySource": "businessLogic",
        },
        "logicFlow": logic_flow,
        "legacyEvidence": legacy_evidence,
        "queryContracts": query_contracts,
        "mappingRules": mapping_rules,
        "dependencyHints": dependency_hints,
        "constraints": constraints,
        "unresolved": unresolved,
    }


def normalize_step_title(title: str) -> str:
    normalized = clean_text(title)
    normalized = re.sub(r"^[.\-、;；\s]+", "", normalized)
    return normalized or clean_text(title)


def build_legacy_business_logic_steps(logic_entries: list[dict[str, str]]) -> list[dict[str, object]]:
    steps: list[dict[str, object]] = []
    next_fallback_step = 0
    for entry in logic_entries:
        key = clean_text(entry.get("key"))
        details = clean_text(entry.get("value"))
        if not key or not details:
            continue
        if normalize_token("backendapi") in normalize_token(key):
            continue
        match = re.match(r"^(?P<step>\d+)\s*(?P<title>.+)$", key.split("\n", 1)[0])
        if match:
            step = int(match.group("step"))
            title = normalize_step_title(match.group("title")) or key
            next_fallback_step = max(next_fallback_step, step + 1)
        else:
            step = next_fallback_step
            title = normalize_step_title(key)
            next_fallback_step += 1
        steps.append({"step": step, "title": title, "details": details})
    return steps


def walk_field_nodes(fields: list[dict[str, Any]]) -> Iterable[dict[str, Any]]:
    for field in fields:
        yield field
        yield from walk_field_nodes(field.get("properties") or [])


def build_response_field_index(fields: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    return {clean_text(field.get("rowKey")): field for field in walk_field_nodes(fields) if clean_text(field.get("rowKey"))}


def derive_mapping_group_row_key(row_key: str, response_index: dict[str, dict[str, Any]]) -> str:
    cleaned = clean_text(row_key)
    if not cleaned:
        return cleaned
    parts = cleaned.split(".")
    if len(parts) >= 2:
        candidate = ".".join(parts[:2])
        if candidate in response_index:
            return candidate
    parent = cleaned
    while "." in parent:
        parent = parent.rsplit(".", 1)[0]
        if parent in response_index:
            return parent
    return cleaned


def extract_call_prefix(source: str) -> str:
    cleaned = clean_text(source)
    if not cleaned:
        return ""
    match = re.match(r"^(?P<prefix>.*?\(\))", cleaned)
    return clean_text(match.group("prefix")) if match else cleaned


def normalize_mapping_source_text(source: str, *, call_prefix_only: bool = False) -> str:
    cleaned = clean_text(source)
    if not cleaned:
        return ""
    cleaned = re.sub(r"\s*\.\s*", ".", cleaned)
    reference_match = re.match(r"^(?P<ref>[A-Za-z0-9_]+(?:\.[A-Za-z0-9_]+)*(?:\(\))?(?:\.[A-Za-z0-9_]+)*)", cleaned)
    if reference_match and reference_match.end() < len(cleaned) and cleaned[reference_match.end()] in {' ', '\t', ',', '，', '；', ';', '：', ':', '"', '「', '（', '('}:
        cleaned = reference_match.group("ref")
    if call_prefix_only:
        return extract_call_prefix(cleaned)
    return cleaned


def derive_group_source(primary_source: str, field_sources: list[str]) -> str:
    if clean_text(primary_source):
        return normalize_mapping_source_text(primary_source)
    prefixes = [normalize_mapping_source_text(source, call_prefix_only=True) for source in field_sources if clean_text(source)]
    unique = list(dict.fromkeys(prefix for prefix in prefixes if prefix))
    if len(unique) == 1:
        return unique[0]
    return unique[0] if unique else ""


def join_non_empty_unique(values: Iterable[str], *, separator: str) -> str:
    ordered: list[str] = []
    seen: set[str] = set()
    for value in values:
        cleaned = clean_text(value)
        if not cleaned or cleaned in seen:
            continue
        ordered.append(cleaned)
        seen.add(cleaned)
    return separator.join(ordered)


def aggregate_field_mappings(field_mappings: list[dict[str, Any]], response_fields: list[dict[str, Any]]) -> list[dict[str, Any]]:
    response_index = build_response_field_index(response_fields)
    response_groups: dict[str, dict[str, Any]] = {}
    response_order: list[str] = []

    for mapping in field_mappings:
        kind = clean_text(mapping.get("kind"))
        if kind and kind != "response_field":
            continue

        row_key = clean_text(mapping.get("rowKey"))
        if not row_key:
            continue
        group_key = derive_mapping_group_row_key(row_key, response_index)
        group_field = response_index.get(group_key, {})
        group = response_groups.setdefault(
            group_key,
            {
                "target": clean_text(group_field.get("fieldName")) or clean_text(mapping.get("field")) or group_key,
                "source": "",
                "rule": clean_text(mapping.get("rule")) if row_key == group_key else "",
                "fields": [],
            },
        )
        if group_key not in response_order:
            response_order.append(group_key)

        field_name = clean_text(mapping.get("field"))
        field_entry = {
            "field": field_name or clean_text(group_field.get("fieldName")) or group["target"],
            "source": normalize_mapping_source_text(clean_text(mapping.get("source"))),
            "rule": clean_text(mapping.get("rule")),
        }

        if row_key == group_key:
            group["source"] = normalize_mapping_source_text(clean_text(mapping.get("source"))) or group["source"]
            if not (group_field.get("properties") or []):
                group["fields"].append(field_entry)
            continue

        group["fields"].append(field_entry)

    grouped_response_mappings: list[dict[str, Any]] = []
    for group_key in response_order:
        group = response_groups[group_key]
        if not group["fields"]:
            group["fields"].append(
                {
                    "field": group["target"],
                    "source": group["source"],
                    "rule": group["rule"],
                }
            )
        group["source"] = derive_group_source(group["source"], [field.get("source") or "" for field in group["fields"]])
        if not group["rule"]:
            group["rule"] = join_non_empty_unique((field.get("rule") for field in group["fields"]), separator="；")
        grouped_response_mappings.append(group)

    merged_groups: list[dict[str, Any]] = []
    for group in grouped_response_mappings:
        merge_key = extract_call_prefix(group.get("source", ""))
        is_scalar_group = len(group.get("fields") or []) == 1
        if (
            merged_groups
            and merge_key
            and is_scalar_group
            and len(merged_groups[-1].get("fields") or []) == 1
            and extract_call_prefix(merged_groups[-1].get("source", "")) == merge_key
        ):
            merged_groups[-1]["target"] = f"{merged_groups[-1]['target']} / {group['target']}"
            merged_groups[-1]["source"] = normalize_mapping_source_text(merge_key)
            merged_groups[-1]["rule"] = join_non_empty_unique([merged_groups[-1].get("rule"), group.get("rule")], separator="；")
            merged_groups[-1]["fields"].extend(group.get("fields") or [])
            continue
        merged_groups.append(group)

    return merged_groups


def row_mentions_api_name(text: str, api_name: str) -> bool:
    return normalize_token(api_name) in normalize_token(text)


def filter_logic_entries_by_api_name(logic_entries: list[dict[str, str]], api_name: str) -> list[dict[str, str]]:
    filtered: list[dict[str, str]] = []
    for entry in logic_entries:
        merged_text = "\n".join([clean_text(entry.get("key")), clean_text(entry.get("value"))])
        if row_mentions_api_name(merged_text, api_name):
            filtered.append(entry)
    return filtered


def filter_field_mapping_rules_by_api_name(rules: list[dict[str, str]], api_name: str) -> list[dict[str, str]]:
    return [rule for rule in rules if row_mentions_api_name(clean_text(rule.get("source")), api_name)]


def infer_field_mappings_from_logic_entries(logic_entries: list[dict[str, str]], api_name: str) -> list[dict[str, str]]:
    pattern = re.compile(
        rf"(?P<target>[A-Za-z_][A-Za-z0-9_]*)\s*=\s*(?P<source>{re.escape(api_name)}(?:\s*\(\s*\))?(?:\s*\.\s*[A-Za-z_][A-Za-z0-9_]*)+)",
        re.IGNORECASE,
    )
    inferred: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()
    counter = 1
    for entry in logic_entries:
        details = clean_text(entry.get("value"))
        if not details:
            continue
        for match in pattern.finditer(details):
            field = clean_text(match.group("target"))
            source = normalize_mapping_source_text(clean_text(match.group("source")))
            if not field or not source:
                continue
            dedupe_key = (field, source)
            if dedupe_key in seen:
                continue
            seen.add(dedupe_key)
            inferred.append(
                {
                    "rowKey": f"inferred.{counter}",
                    "field": field,
                    "source": source,
                    "rule": f"从调用方业务逻辑 {clean_text(entry.get('key'))} 推断",
                    "kind": "response_field",
                    "authority": "caller_sheet_inferred",
                }
            )
            counter += 1
    return inferred


def filter_field_tree_by_row_keys(fields: list[dict[str, Any]], row_keys: set[str]) -> list[dict[str, Any]]:
    filtered: list[dict[str, Any]] = []
    for field in fields:
        child_fields = filter_field_tree_by_row_keys(field.get("properties") or [], row_keys)
        row_key = clean_text(field.get("rowKey"))
        if row_key not in row_keys and not child_fields:
            continue
        node = dict(field)
        if child_fields:
            node["properties"] = child_fields
        else:
            node.pop("properties", None)
        filtered.append(node)
    return filtered


def build_synthetic_response_fields_from_mappings(
    field_mappings: list[dict[str, str]],
    *,
    api_name: str,
    existing_row_keys: set[str] | None = None,
) -> list[dict[str, Any]]:
    fields: list[dict[str, Any]] = []
    seen_rows = set(existing_row_keys or set())
    for mapping in field_mappings:
        row_key = clean_text(mapping.get("rowKey"))
        field_name = clean_text(mapping.get("field"))
        if not row_key or not field_name or row_key in seen_rows:
            continue
        seen_rows.add(row_key)
        fields.append(
            {
                "rowKey": row_key,
                "fieldName": field_name,
                "dataType": "string",
                "required": False,
                "description": f"从调用方 sheet 推断的 {api_name} 输出字段",
                "notes": clean_text(mapping.get("source")),
            }
        )
    return fields


def compact_raw_appendix(raw_appendix: dict[str, object]) -> dict[str, object] | None:
    compacted = {key: value for key, value in raw_appendix.items() if value not in (None, "", [], {})}
    return compacted or None


def sheet_declares_response_list(rows: list[list[str]]) -> bool:
    for row in rows[:5]:
        for cell in row:
            token = normalize_token(cell)
            if token in {normalize_token("返回List"), normalize_token("returnList")}:
                return True
    return False


def apply_response_shape_hints(
    response_fields: list[dict[str, object]],
    *,
    mock_examples: list[dict[str, object]],
    returns_list: bool,
) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    if not returns_list:
        return response_fields, mock_examples

    adjusted_fields: list[dict[str, object]] = []
    for field in response_fields:
        node = dict(field)
        if clean_text(node.get("rowKey")) == "5" or clean_text(node.get("fieldName")) == "responseData":
            node["dataType"] = "array"
            if clean_text(node.get("example")) in {'"responseData":{}', '"responseData": {}', ""}:
                node["example"] = '"responseData":[]'
        adjusted_fields.append(node)

    adjusted_examples: list[dict[str, object]] = []
    for item in mock_examples:
        updated = dict(item)
        payload = updated.get("responsePayload")
        if isinstance(payload, dict) and "responseData" in payload:
            updated_payload = dict(payload)
            response_data = updated_payload.get("responseData")
            if isinstance(response_data, dict):
                updated_payload["responseData"] = [response_data] if response_data else []
            elif response_data is None:
                updated_payload["responseData"] = []
            updated["responsePayload"] = updated_payload
        adjusted_examples.append(updated)
    return adjusted_fields, adjusted_examples


def extract_api_spec_sections(rows: list[list[str]]) -> dict[str, object]:
    request_fields, request_notes = extract_structured_fields(rows, section_title="Request", stop_titles=("Response",))
    response_fields, response_notes = extract_structured_fields(
        rows,
        section_title="Response",
        stop_titles=("範例", "For中台開發人員", "API 內部業務邏輯"),
    )
    mock_examples, examples_external = extract_mock_examples(rows)
    logic_entries = extract_business_logic_section(rows)
    backend_apis, backend_api_lines = extract_backend_apis(logic_entries)
    merge_backend_apis_from_business_text(backend_apis, logic_entries)
    returns_list = sheet_declares_response_list(rows)
    response_fields, mock_examples = apply_response_shape_hints(
        response_fields,
        mock_examples=mock_examples,
        returns_list=returns_list,
    )
    mock_examples = normalize_mock_examples_against_contract(
        mock_examples,
        request_fields=request_fields,
        response_fields=response_fields,
    )
    other_notes: list[str] = []
    other_notes.extend(examples_external.get("mockExampleNotes") or [])
    raw_appendix = compact_raw_appendix(
        {
            "requestNotes": request_notes,
            "responseNotes": response_notes,
            "backendApiLines": backend_api_lines,
            "unparsedMockExamples": examples_external.get("unparsedMockExamples") or [],
            "otherNotes": other_notes,
        }
    )
    return {
        "request": request_fields,
        "response": response_fields,
        "mockExamples": select_representative_mock_examples(mock_examples),
        "backendApis": backend_apis,
        "rawAppendix": raw_appendix,
        "errorCodeRules": build_error_code_rules_from_examples(mock_examples),
    }


def format_validation_path(path: Iterable[object]) -> str:
    parts = [str(part) for part in path]
    return ".".join(parts) if parts else "$"


def validate_api_spec_payload(payload: dict[str, object]) -> None:
    schema = load_schema("api-spec.schema.json")
    validator = Draft202012Validator(schema)
    errors = sorted(validator.iter_errors(payload), key=lambda item: list(item.absolute_path))
    if errors:
        first = errors[0]
        raise SkillError(f"API_Spec JSON 未通过 schema 校验：{format_validation_path(first.absolute_path)} - {first.message}")


def flatten_contract_fields(fields: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    items: dict[str, dict[str, Any]] = {}
    for field in walk_field_nodes(fields):
        row_key = clean_text(field.get("rowKey"))
        if not row_key:
            continue
        items[row_key] = {
            "fieldName": field.get("fieldName"),
            "dataType": field.get("dataType"),
            "required": field.get("required"),
            "description": field.get("description"),
            "example": field.get("example"),
            "notes": field.get("notes"),
        }
    return items


def build_code_handoff_review_findings(actual_api_spec: dict[str, Any]) -> tuple[list[bool], list[dict[str, Any]]]:
    checks: list[bool] = []
    findings: list[dict[str, Any]] = []
    code_handoff = actual_api_spec.get("codeHandoff") if isinstance(actual_api_spec.get("codeHandoff"), dict) else {}
    business_logic = actual_api_spec.get("businessLogic") if isinstance(actual_api_spec.get("businessLogic"), dict) else {}

    unresolved = [item for item in list(code_handoff.get("unresolved") or []) if isinstance(item, dict)]
    for index, item in enumerate(unresolved):
        is_blocking = bool(item.get("blocking"))
        checks.append(not is_blocking)
        if not is_blocking:
            continue
        findings.append(
            {
                "severity": "blocking",
                "kind": "code_handoff_unresolved",
                "path": f"codeHandoff.unresolved.{index}",
                "expected": "no blocking unresolved items",
                "actual": {
                    "topic": clean_text(item.get("topic")),
                    "reason": clean_text(item.get("reason")),
                },
            }
        )

    runtime_dependencies = [item for item in list(business_logic.get("runtimeDependencies") or []) if isinstance(item, dict)]
    dependency_ids = {normalize_token(clean_text(item.get("id"))) for item in runtime_dependencies}
    query_contracts = [item for item in list(code_handoff.get("queryContracts") or []) if isinstance(item, dict)]
    for index, contract in enumerate(query_contracts):
        sql_text = clean_text(contract.get("sqlText"))
        if not sql_text:
            checks.append(False)
            findings.append(
                {
                    "severity": "blocking",
                    "kind": "query_contract_sql_missing",
                    "path": f"codeHandoff.queryContracts.{index}.sqlText",
                    "expected": "non-empty sqlText",
                    "actual": clean_text(contract.get("sqlText")),
                }
            )
            continue

        expected_must_contain = build_sql_must_contain(sql_text)
        actual_must_contain = [clean_text(value) for value in list(contract.get("mustContain") or []) if clean_text(value)]
        missing_must_contain = [value for value in expected_must_contain if value not in actual_must_contain]
        unexpected_must_contain = [value for value in actual_must_contain if value not in expected_must_contain]
        must_contain_ok = not missing_must_contain and not unexpected_must_contain
        checks.append(must_contain_ok)
        if not must_contain_ok:
            findings.append(
                {
                    "severity": "blocking",
                    "kind": "query_contract_must_contain_mismatch",
                    "path": f"codeHandoff.queryContracts.{index}.mustContain",
                    "expected": expected_must_contain,
                    "actual": actual_must_contain,
                }
            )

        expected_data_sources = extract_table_names(sql_text)
        actual_data_sources = [clean_text(value) for value in list(contract.get("dataSources") or []) if clean_text(value)]
        missing_data_sources = [value for value in expected_data_sources if value not in actual_data_sources]
        unexpected_data_sources = [value for value in actual_data_sources if value not in expected_data_sources]
        data_sources_ok = not missing_data_sources and not unexpected_data_sources
        checks.append(data_sources_ok)
        if not data_sources_ok:
            findings.append(
                {
                    "severity": "blocking",
                    "kind": "query_contract_data_sources_mismatch",
                    "path": f"codeHandoff.queryContracts.{index}.dataSources",
                    "expected": expected_data_sources,
                    "actual": actual_data_sources,
                }
            )
    return checks, findings


def build_spec_review_findings(
    *,
    expected_request_fields: list[dict[str, Any]],
    expected_response_fields: list[dict[str, Any]],
    expected_mock_examples: list[dict[str, Any]],
    actual_api_spec: dict[str, Any],
) -> tuple[list[bool], list[dict[str, Any]]]:
    expected_request_map = flatten_contract_fields(expected_request_fields)
    expected_response_map = flatten_contract_fields(expected_response_fields)
    actual_request_map = flatten_contract_fields(list(actual_api_spec.get("request") or []))
    actual_response_map = flatten_contract_fields(list(actual_api_spec.get("response") or []))

    checks: list[bool] = []
    findings: list[dict[str, Any]] = []

    for scope, expected_map, actual_map in (
        ("request", expected_request_map, actual_request_map),
        ("response", expected_response_map, actual_response_map),
    ):
        for row_key, expected_field in expected_map.items():
            actual_field = actual_map.get(row_key)
            if actual_field is None:
                checks.append(False)
                findings.append(
                    {
                        "severity": "blocking",
                        "kind": "missing_field",
                        "path": f"{scope}.{row_key}",
                        "expected": expected_field,
                        "actual": None,
                    }
                )
                continue

            for field_name in ("fieldName", "dataType", "required", "description", "example", "notes"):
                expected_value = expected_field.get(field_name)
                actual_value = actual_field.get(field_name)
                if expected_value in (None, "") and actual_value in (None, ""):
                    continue
                same = expected_value == actual_value
                checks.append(same)
                if same:
                    continue

                severity = "warning"
                if field_name == "dataType":
                    severity = "blocking"
                if field_name == "required" and expected_value is True and actual_value is False:
                    severity = "blocking"
                if field_name == "notes" and expected_value and not actual_value:
                    severity = "blocking"

                findings.append(
                    {
                        "severity": severity,
                        "kind": f"{field_name}_mismatch",
                        "path": f"{scope}.{row_key}.{field_name}",
                        "expected": expected_value,
                        "actual": actual_value,
                    }
                )

    expected_response_data = expected_response_map.get("5")
    actual_response_data = actual_response_map.get("5")
    if (
        expected_response_data
        and actual_response_data
        and expected_response_data.get("dataType") != actual_response_data.get("dataType")
    ):
        findings.append(
            {
                "severity": "blocking",
                "kind": "response_shape",
                "path": "response.5.dataType",
                "expected": expected_response_data.get("dataType"),
                "actual": actual_response_data.get("dataType"),
            }
        )

    success_examples = [
        item
        for item in list(actual_api_spec.get("mockExamples") or [])
        if isinstance(item, dict)
        and isinstance(item.get("responsePayload"), dict)
        and (
            item["responsePayload"].get("isSuccess") is True
            or clean_text(item["responsePayload"].get("responseCode")) == "0000"
        )
    ]
    if success_examples and expected_response_data and clean_text(expected_response_data.get("dataType")) == "array":
        response_data = success_examples[0]["responsePayload"].get("responseData")
        ok = isinstance(response_data, list)
        checks.append(ok)
        if not ok:
            findings.append(
                {
                    "severity": "blocking",
                    "kind": "mock_response_shape",
                    "path": "mockExamples.success.responsePayload.responseData",
                    "expected": "list",
                    "actual": type(response_data).__name__,
                }
            )

    if expected_mock_examples and not list(actual_api_spec.get("mockExamples") or []):
        checks.append(False)
        findings.append(
            {
                "severity": "warning",
                "kind": "missing_mock_examples",
                "path": "mockExamples",
                "expected": len(expected_mock_examples),
                "actual": 0,
            }
        )

    handoff_checks, handoff_findings = build_code_handoff_review_findings(actual_api_spec)
    checks.extend(handoff_checks)
    findings.extend(handoff_findings)
    return checks, findings


def build_spec_review_artifact(
    *,
    api_spec: dict[str, Any],
    checks: list[bool],
    findings: list[dict[str, Any]],
) -> dict[str, Any]:
    blocking_findings = [item for item in findings if item.get("severity") == "blocking"]
    warning_findings = [item for item in findings if item.get("severity") == "warning"]
    total_checks = len(checks)
    matched_checks = sum(1 for item in checks if item)
    return {
        "reviewStatus": "blocked" if blocking_findings else "passed",
        "checkedAt": now_iso(),
        "apiId": clean_text(api_spec.get("apiId")),
        "source": api_spec.get("source") or {},
        "summary": {
            "checks": total_checks,
            "matched": matched_checks,
            "matchPct": round((matched_checks / total_checks) * 100, 1) if total_checks else 100.0,
            "blockingCount": len(blocking_findings),
            "warningCount": len(warning_findings),
        },
        "findings": findings,
    }


def write_spec_review_artifact(context: ExecutionContext, api_id: str, artifact: dict[str, Any]) -> None:
    dump_json(context.paths.api_dir(api_id) / "spec-review.json", artifact)


def run_spec_review_gate(
    *,
    context: ExecutionContext,
    api_id: str,
    api_spec: dict[str, Any],
    expected_request_fields: list[dict[str, Any]],
    expected_response_fields: list[dict[str, Any]],
    expected_mock_examples: list[dict[str, Any]],
) -> dict[str, Any]:
    checks, findings = build_spec_review_findings(
        expected_request_fields=expected_request_fields,
        expected_response_fields=expected_response_fields,
        expected_mock_examples=expected_mock_examples,
        actual_api_spec=api_spec,
    )
    artifact = build_spec_review_artifact(api_spec=api_spec, checks=checks, findings=findings)
    write_spec_review_artifact(context, api_id, artifact)

    blocking_findings = [item for item in findings if item.get("severity") == "blocking"]
    if blocking_findings:
        first = blocking_findings[0]
        raise SkillError(
            f"AI review gate found {len(blocking_findings)} blocking drift(s): "
            f"{first.get('path')} ({first.get('kind')})",
            status="blocked",
        )
    return api_spec


def build_indirect_common_util_payload_parts(
    context: ExecutionContext,
    *,
    api_id: str,
    entry: ApiEntry,
    workbook_path: Path,
    sheet_names: list[str],
    reference_library: dict[str, list[dict[str, Any]]] | None = None,
) -> dict[str, Any]:
    primary_sheet_name = sheet_names[0]
    visual_notes, visual_unresolved = collect_excel_visual_warnings(workbook_path, sheet_names)
    rows = load_sheet_rows(SheetMatch(api=entry, workbook_path=workbook_path, sheet_name=primary_sheet_name))
    logic_entries = extract_business_logic_section(rows)
    matched_logic_entries = filter_logic_entries_by_api_name(logic_entries, entry.name)
    extracted_sections = extract_api_spec_sections(rows)
    full_business_logic = extract_business_logic_structure_from_rows(rows)
    direct_field_mappings = filter_field_mapping_rules_by_api_name(extract_field_mapping_rules(rows), entry.name)
    inferred_field_mappings = infer_field_mappings_from_logic_entries(matched_logic_entries, entry.name)
    all_field_mappings = direct_field_mappings + inferred_field_mappings

    if not matched_logic_entries and not all_field_mappings:
        raise SkillError(f"无法从 {workbook_path.name} 的调用方 sheet 推断 {entry.name} 逻辑", status="blocked")

    selected_row_keys = {clean_text(item.get("rowKey")) for item in direct_field_mappings if clean_text(item.get("rowKey"))}
    response_fields = filter_field_tree_by_row_keys(extracted_sections["response"], selected_row_keys)
    response_fields.extend(
        build_synthetic_response_fields_from_mappings(
            inferred_field_mappings,
            api_name=entry.name,
            existing_row_keys={clean_text(field.get("rowKey")) for field in walk_field_nodes(response_fields)},
        )
    )
    if not response_fields:
        response_fields = build_synthetic_response_fields_from_mappings(all_field_mappings, api_name=entry.name)

    field_mappings = aggregate_field_mappings(all_field_mappings, response_fields)
    indirect_notes = [f"无独立 {entry.category} sheet；本规格依据 `{primary_sheet_name}` 中引用 `{entry.name}` 的业务逻辑推断。"]
    if len(sheet_names) > 1:
        indirect_notes.append(f"同 workbook 另命中引用 sheet：{', '.join(sheet_names[1:])}")
    indirect_notes.extend(visual_notes)

    raw_appendix = compact_raw_appendix(
        {
            "backendApiLines": ((extracted_sections.get("rawAppendix") or {}).get("backendApiLines") or []),
            "otherNotes": indirect_notes,
        }
    )
    business_logic_payload = {
        "steps": build_legacy_business_logic_steps(matched_logic_entries),
        "fieldMappings": field_mappings,
        "lookupTables": full_business_logic.get("lookupTables", []),
        "errorCodeRules": [],
        "runtimeDependencies": full_business_logic.get("runtimeDependencies", []),
        "dataSources": full_business_logic.get("dataSources", []),
        "sqlSpecs": full_business_logic.get("sqlSpecs", []),
        "legacyReferences": full_business_logic.get("legacyReferences", []),
        "prohibitedShortcuts": full_business_logic.get("prohibitedShortcuts", []),
    }
    business_logic_payload = {
        key: value
        for key, value in business_logic_payload.items()
        if key in {"steps", "fieldMappings", "lookupTables", "errorCodeRules", "runtimeDependencies", "dataSources", "sqlSpecs", "legacyReferences", "prohibitedShortcuts"}
        and not (key == "lookupTables" and not value)
    }
    business_logic_payload = apply_reference_hints_to_business_logic(
        business_logic_payload,
        reference_library=reference_library or {},
        backend_apis=extracted_sections["backendApis"],
        raw_appendix=raw_appendix,
    )
    source = build_api_source(context, workbook_path=workbook_path, sheet_names=sheet_names)
    api_spec = build_api_spec_payload(
        context=context,
        api_id=api_id,
        entry=entry,
        workbook_path=workbook_path,
        sheet_names=sheet_names,
        business_logic_payload=business_logic_payload,
        request_fields=[],
        response_fields=response_fields,
        mock_examples=[],
        backend_apis=extracted_sections["backendApis"],
        raw_appendix=raw_appendix,
        additional_unresolved=visual_unresolved + list(full_business_logic.get("dependencyUnresolved", [])),
    )
    api_spec = run_spec_review_gate(
        context=context,
        api_id=api_id,
        api_spec=api_spec,
        expected_request_fields=[],
        expected_response_fields=response_fields,
        expected_mock_examples=[],
    )
    validate_api_spec_payload(api_spec)
    return {
        "source": source,
        "apiSpec": api_spec,
        "sourceFingerprint": build_source_fingerprint(
            context=context,
            api_category=entry.category,
            api_name=entry.name,
            source=source,
            request_fields=[],
            response_fields=response_fields,
            mock_examples=[],
            backend_apis=extracted_sections["backendApis"],
            business_logic_payload=business_logic_payload,
            raw_appendix=raw_appendix,
            additional_unresolved=visual_unresolved + list(full_business_logic.get("dependencyUnresolved", [])),
        ),
    }


def build_api_spec_payload(
    *,
    context: ExecutionContext,
    api_id: str,
    entry: ApiEntry,
    workbook_path: Path,
    sheet_names: list[str],
    business_logic_payload: dict[str, object],
    request_fields: list[dict[str, object]],
    response_fields: list[dict[str, object]],
    mock_examples: list[dict[str, object]],
    backend_apis: dict[str, list[str]],
    raw_appendix: dict[str, object] | None,
    additional_unresolved: list[dict[str, Any]] | None = None,
) -> dict[str, object]:
    version = normalize_display_version(extract_version_token_from_tsd_path(context.docx_path))
    code_handoff = build_code_handoff(
        project_root=context.project_root,
        agent_dir=context.agent_dir,
        function_code=context.function_code,
        api_category=entry.category,
        api_name=entry.name,
        request_fields=request_fields,
        response_fields=response_fields,
        business_logic_payload=business_logic_payload,
        additional_unresolved=additional_unresolved,
    )
    payload = {
        "schemaVersion": API_SPEC_SCHEMA_VERSION,
        "apiId": api_id,
        "newAuthor": context.new_author,
        "functionCode": context.function_code,
        "version": version,
        "apiCategory": entry.category,
        "apiName": entry.name,
        "source": {
            "tsdFile": context.docx_path.name,
            "workbookFile": workbook_path.name,
            "sheetNames": sheet_names,
        },
        "request": request_fields,
        "response": response_fields,
        "mockExamples": mock_examples,
        "backendApis": backend_apis,
        "businessLogic": business_logic_payload,
        "codeHandoff": code_handoff,
    }
    if raw_appendix:
        payload["rawAppendix"] = raw_appendix
    return payload


def build_source_fingerprint(
    *,
    context: ExecutionContext,
    api_category: str,
    api_name: str,
    source: dict[str, Any],
    request_fields: list[dict[str, object]],
    response_fields: list[dict[str, object]],
    mock_examples: list[dict[str, object]],
    backend_apis: dict[str, list[str]],
    business_logic_payload: dict[str, Any],
    raw_appendix: dict[str, object] | None,
    additional_unresolved: list[dict[str, Any]] | None = None,
) -> str:
    return stable_payload_hash(
        {
            "schemaVersion": API_SPEC_SCHEMA_VERSION,
            "newAuthor": context.new_author,
            "tsdFile": source.get("tsdFile"),
            "workbookFile": source.get("workbookFile"),
            "sheetNames": source.get("sheetNames") or [],
            "request": request_fields,
            "response": response_fields,
            "mockExamples": mock_examples,
            "backendApis": backend_apis,
            "businessLogic": business_logic_payload,
            "codeHandoff": build_code_handoff(
                project_root=context.project_root,
                agent_dir=context.agent_dir,
                function_code=context.function_code,
                api_category=api_category,
                api_name=api_name,
                request_fields=request_fields,
                response_fields=response_fields,
                business_logic_payload=business_logic_payload,
                additional_unresolved=additional_unresolved,
            ),
            "rawAppendix": raw_appendix or {},
        }
    )


def has_legacy_lookup_field_mappings(payload: dict[str, Any]) -> bool:
    business_logic = payload.get("businessLogic") or {}
    if business_logic.get("lookupTables"):
        return False
    contract_field_names = {
        clean_text(field.get("fieldName"))
        for field in walk_field_nodes((payload.get("request") or []) + (payload.get("response") or []))
        if clean_text(field.get("fieldName"))
    }
    for mapping in business_logic.get("fieldMappings") or []:
        target = clean_text(mapping.get("target"))
        if not target:
            continue
        target_parts = [clean_text(part) for part in re.split(r"\s*/\s*", target) if clean_text(part)]
        if target_parts and not all(part in contract_field_names for part in target_parts):
            return True
    return False


def resolve_existing_api_spec_path(
    context: ExecutionContext,
    api_id: str,
    *,
    require_exists: bool = False,
) -> Path | None:
    preferred = context.paths.api_spec_path(api_id, function_code=context.function_code, tsd_path=context.docx_path)
    if preferred.exists():
        return preferred
    api_dir = context.paths.api_dir(api_id)
    candidates = sorted(path.resolve() for path in api_dir.glob("*_API_Spec.json") if path.is_file())
    if len(candidates) == 1:
        return candidates[0]
    if require_exists:
        if not candidates:
            raise SkillError(f"API_Spec.json not found for {api_id}")
        raise SkillError(f"multiple API_Spec.json files found for {api_id}")
    return None


def needs_api_spec_refresh(context: ExecutionContext, previous: dict[str, Any] | None) -> tuple[bool, str]:
    if previous is None:
        return False, ""
    api_id = clean_text(previous.get("apiId"))
    api_spec_path = resolve_existing_api_spec_path(context, api_id) if api_id else None
    if api_spec_path is None:
        return True, "API_Spec.json 不存在"
    try:
        payload = load_json(api_spec_path)
    except Exception:
        return True, "API_Spec.json 无法读取"
    if clean_text(payload.get("schemaVersion")) != API_SPEC_SCHEMA_VERSION:
        return True, "API_Spec schemaVersion 需要刷新"
    if has_legacy_lookup_field_mappings(payload):
        return True, "API_Spec fieldMappings 仍包含旧 lookup 映射"
    return False, ""


def create_checklist_item(context: ExecutionContext, entry: ApiEntry) -> dict[str, Any]:
    api_id = build_api_id(context.function_code, entry.category, entry.name)
    return {
        "apiId": api_id,
        "apiCategory": entry.category,
        "apiName": entry.name,
        "status": "pending",
        "blockReason": None,
        "source": build_api_source(context, workbook_path=None, sheet_names=[]),
        "sourceFingerprint": None,
        "fixtureStatus": "waiting_spec",
        "fixturePhase": "waiting_spec",
        "fixtureBlockReason": None,
        "fixtureSourceFingerprint": None,
        "fixtureUpdatedAt": None,
        "fixtureArtifacts": {
            "dbFixtureReport": None,
            "tableChecks": None,
            "seedPlan": None,
            "seedExecuted": None,
            "seedManifest": None,
        },
        "codeStatus": "waiting_spec",
        "codePhase": "waiting_spec",
        "codeBlockReason": None,
        "codeReset": False,
    }


def read_fixture_report_payload(context: ExecutionContext, api_id: str) -> dict[str, Any] | None:
    report_path = context.paths.api_dir(api_id) / "db-fixture-report.json"
    if not report_path.exists():
        return None
    try:
        payload = load_json(report_path)
    except Exception:
        return None
    return payload if isinstance(payload, dict) else None


def build_fixture_artifacts_payload(
    context: ExecutionContext,
    api_id: str,
    existing_payload: dict[str, Any] | None = None,
) -> dict[str, Any]:
    existing_payload = existing_payload if isinstance(existing_payload, dict) else {}
    api_dir = context.paths.api_dir(api_id)
    candidates = {
        "dbFixtureReport": api_dir / "db-fixture-report.json",
        "tableChecks": api_dir / "table-checks.json",
        "seedPlan": api_dir / "seed-plan.sql",
        "seedExecuted": api_dir / "seed-executed.sql",
        "seedManifest": api_dir / "seed-manifest.json",
    }
    return {
        key: normalize_persisted_path(path, project_root=context.project_root) if path.exists() else existing_payload.get(key)
        for key, path in candidates.items()
    }


def resolve_fixture_state(
    context: ExecutionContext,
    api_id: str,
    raw_item: dict[str, Any],
    manifest: dict[str, Any],
    *,
    spec_status: str,
) -> dict[str, Any]:
    report_payload = read_fixture_report_payload(context, api_id) or {}
    raw_fixture_status = clean_text(raw_item.get("fixtureStatus"))
    manifest_fixture_status = clean_text(manifest.get("fixtureStatus"))
    report_fixture_status = clean_text(report_payload.get("status"))
    fixture_status = raw_fixture_status or manifest_fixture_status or report_fixture_status
    if report_fixture_status and fixture_status in {"", "pending", "waiting_spec"}:
        fixture_status = report_fixture_status
    if not fixture_status:
        fixture_status = "pending" if spec_status == "done" else "waiting_spec"
    default_phase = "pending" if fixture_status == "pending" else "waiting_spec"
    raw_fixture_phase = clean_text(raw_item.get("fixturePhase"))
    manifest_fixture_phase = clean_text(manifest.get("fixturePhase"))
    report_fixture_phase = clean_text(report_payload.get("phase"))
    fixture_phase = raw_fixture_phase or manifest_fixture_phase or report_fixture_phase or default_phase
    if report_fixture_phase and fixture_phase in {"", "pending", "waiting_spec"} and fixture_status == report_fixture_status:
        fixture_phase = report_fixture_phase
    fixture_block_reason = (
        clean_text(raw_item.get("fixtureBlockReason"))
        or clean_text(manifest.get("fixtureBlockReason"))
        or None
    )
    return {
        "fixtureStatus": fixture_status,
        "fixturePhase": fixture_phase,
        "fixtureBlockReason": fixture_block_reason,
        "fixtureSourceFingerprint": clean_text(raw_item.get("fixtureSourceFingerprint") or manifest.get("fixtureSourceFingerprint")) or None,
        "fixtureUpdatedAt": clean_text(manifest.get("fixtureUpdatedAt") or report_payload.get("updatedAt")) or None,
        "fixtureArtifacts": build_fixture_artifacts_payload(
            context,
            api_id,
            manifest.get("fixtureArtifacts") if isinstance(manifest.get("fixtureArtifacts"), dict) else None,
        ),
    }


def normalize_existing_item(
    context: ExecutionContext,
    raw_item: dict[str, Any],
    manifest: dict[str, Any] | None,
) -> dict[str, Any]:
    manifest = manifest or {}
    spec_source = manifest.get("specSource") if isinstance(manifest.get("specSource"), dict) else manifest.get("source")
    spec_status = clean_text(raw_item.get("specStatus") or raw_item.get("status") or manifest.get("specStatus") or manifest.get("status")) or "pending"
    fixture_state = resolve_fixture_state(context, raw_item["apiId"], raw_item, manifest, spec_status=spec_status)
    item = {
        "apiId": raw_item["apiId"],
        "apiCategory": raw_item["apiCategory"],
        "apiName": raw_item["apiName"],
        "status": spec_status,
        "blockReason": raw_item.get("specBlockReason", raw_item.get("blockReason")),
        "source": spec_source or build_api_source(context, workbook_path=None, sheet_names=[]),
        "sourceFingerprint": manifest.get("specSourceFingerprint", manifest.get("sourceFingerprint")),
        "fixtureStatus": fixture_state["fixtureStatus"],
        "fixturePhase": fixture_state["fixturePhase"],
        "fixtureBlockReason": fixture_state["fixtureBlockReason"],
        "fixtureSourceFingerprint": fixture_state["fixtureSourceFingerprint"],
        "fixtureUpdatedAt": fixture_state["fixtureUpdatedAt"],
        "fixtureArtifacts": fixture_state["fixtureArtifacts"],
        "codeStatus": clean_text(raw_item.get("codeStatus") or manifest.get("codeStatus")) or default_code_status(spec_status),
        "codePhase": clean_text(raw_item.get("codePhase") or manifest.get("codePhase")) or default_code_phase(spec_status),
        "codeBlockReason": clean_text(raw_item.get("codeBlockReason") or manifest.get("codeBlockReason")) or None,
        "codeReset": False,
    }
    return item


def load_manifest_map(context: ExecutionContext, checklist_items: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    manifest_map: dict[str, dict[str, Any]] = {}
    for item in checklist_items:
        manifest_path = context.paths.manifest_path(item["apiId"])
        if manifest_path.exists():
            manifest_map[item["apiId"]] = load_json(manifest_path)
    return manifest_map


def cleanup_code_artifacts(context: ExecutionContext, api_id: str) -> None:
    api_dir = context.paths.api_dir(api_id)
    for name in ("change-plan.json", "implementation-report.md", "diagnosis-report.json"):
        path = api_dir / name
        try:
            path.unlink()
        except FileNotFoundError:
            pass


def dump_checklist(context: ExecutionContext, items: list[dict[str, Any]]) -> None:
    payload = {
        "schemaVersion": STATE_SCHEMA_VERSION,
        "executionId": context.execution_id,
        "updatedAt": now_iso(),
        "items": [
            {
                "apiId": item["apiId"],
                "apiCategory": item["apiCategory"],
                "apiName": item["apiName"],
                "specStatus": item["status"],
                "specBlockReason": item["blockReason"],
                "specSourceFingerprint": item.get("sourceFingerprint"),
                "fixtureStatus": item.get("fixtureStatus"),
                "fixturePhase": item.get("fixturePhase"),
                "fixtureBlockReason": item.get("fixtureBlockReason"),
                "fixtureSourceFingerprint": item.get("fixtureSourceFingerprint"),
                "codeStatus": item.get("codeStatus") or default_code_status(item["status"]),
                "codePhase": item.get("codePhase") or default_code_phase(item["status"]),
                "codeBlockReason": item.get("codeBlockReason"),
            }
            for item in items
        ],
    }
    dump_json(context.paths.checklist_path, payload)


def build_manifest(context: ExecutionContext, item: dict[str, Any]) -> dict[str, Any]:
    manifest_path = context.paths.manifest_path(item["apiId"])
    existing_payload = load_json(manifest_path) if manifest_path.exists() else {}
    source = item.get("source") or build_api_source(context, workbook_path=None, sheet_names=[])
    api_spec_path = resolve_existing_api_spec_path(context, item["apiId"])
    spec_updated_at = now_iso()
    code_status = clean_text(item.get("codeStatus")) or clean_text(existing_payload.get("codeStatus")) or default_code_status(item["status"])
    code_phase = clean_text(item.get("codePhase")) or clean_text(existing_payload.get("codePhase")) or default_code_phase(item["status"])
    fixture_artifacts = build_fixture_artifacts_payload(
        context,
        item["apiId"],
        item.get("fixtureArtifacts") if isinstance(item.get("fixtureArtifacts"), dict) else existing_payload.get("fixtureArtifacts"),
    )
    reset_code_history = bool(item.get("codeReset")) or (
        code_status in {"pending", "waiting_spec"}
        and clean_text(existing_payload.get("codeStatus"))
        and clean_text(existing_payload.get("codeStatus")) != code_status
    )
    payload = {
        "schemaVersion": STATE_SCHEMA_VERSION,
        "manifestType": "api",
        "executionId": context.execution_id,
        "apiId": item["apiId"],
        "apiCategory": item["apiCategory"],
        "apiName": item["apiName"],
        "status": item["status"] if code_status == "waiting_spec" else code_status,
        "phase": code_phase or item["status"],
        "updatedAt": latest_timestamp(spec_updated_at, existing_payload.get("codeUpdatedAt")),
        "newAuthor": existing_payload.get("newAuthor") or context.new_author,
        "specStatus": item["status"],
        "specUpdatedAt": spec_updated_at,
        "specBlockReason": item["blockReason"],
        "specSourceFingerprint": item.get("sourceFingerprint"),
        "specSource": {
            "tsdFile": source.get("tsdFile"),
            "workbookFile": source.get("workbookFile"),
            "sheetNames": source.get("sheetNames") or [],
        },
        "specArtifacts": {
            "apiSpec": normalize_persisted_path(api_spec_path, project_root=context.project_root) if api_spec_path else None,
        },
        "fixtureStatus": clean_text(item.get("fixtureStatus")) or clean_text(existing_payload.get("fixtureStatus")) or ("pending" if item["status"] == "done" else "waiting_spec"),
        "fixturePhase": clean_text(item.get("fixturePhase")) or clean_text(existing_payload.get("fixturePhase")) or ("pending" if item["status"] == "done" else "waiting_spec"),
        "fixtureUpdatedAt": item.get("fixtureUpdatedAt") or existing_payload.get("fixtureUpdatedAt"),
        "fixtureBlockReason": item.get("fixtureBlockReason") if item.get("fixtureBlockReason") is not None else existing_payload.get("fixtureBlockReason"),
        "fixtureSourceFingerprint": item.get("fixtureSourceFingerprint") or existing_payload.get("fixtureSourceFingerprint"),
        "fixtureArtifacts": fixture_artifacts,
        "codeStatus": code_status,
        "codePhase": code_phase,
        "codeUpdatedAt": existing_payload.get("codeUpdatedAt"),
        "codeBlockReason": None if reset_code_history else existing_payload.get("codeBlockReason"),
        "codeProjectRoot": existing_payload.get("codeProjectRoot"),
        "codeSolutionPath": existing_payload.get("codeSolutionPath"),
        "inputHashes": {} if reset_code_history else existing_payload.get("inputHashes") or {},
        "modifiedFiles": [] if reset_code_history else existing_payload.get("modifiedFiles") or [],
        "validationChecks": [] if reset_code_history else existing_payload.get("validationChecks") or [],
        "validationResults": [] if reset_code_history else existing_payload.get("validationResults") or [],
        "repoDriftFiles": [] if reset_code_history else existing_payload.get("repoDriftFiles") or [],
        "codeArtifacts": build_code_artifacts_payload(context, item["apiId"]),
        "lastMessage": None if reset_code_history else existing_payload.get("lastMessage"),
    }
    validate_manifest_payload(payload)
    return payload


def manifest_payload_without_updated_at(payload: dict[str, Any]) -> dict[str, Any]:
    return {key: value for key, value in payload.items() if key not in {"updatedAt", "specUpdatedAt", "codeUpdatedAt"}}


def dump_manifest_if_changed(context: ExecutionContext, item: dict[str, Any]) -> None:
    manifest_path = context.paths.manifest_path(item["apiId"])
    manifest = build_manifest(context, item)
    if manifest_path.exists():
        try:
            existing = load_json(manifest_path)
        except Exception:
            existing = None
        if isinstance(existing, dict) and manifest_payload_without_updated_at(existing) == manifest_payload_without_updated_at(manifest):
            return
    dump_json(manifest_path, manifest)


def dump_manifests(context: ExecutionContext, items: list[dict[str, Any]]) -> None:
    for item in items:
        dump_manifest_if_changed(context, item)


def build_api_payload_parts(
    context: ExecutionContext,
    *,
    api_id: str,
    entry: ApiEntry,
    workbook_path: Path,
    sheet_names: list[str],
    reference_library: dict[str, list[dict[str, Any]]] | None = None,
) -> dict[str, Any]:
    if sheet_names and not sheet_name_is_direct_match(sheet_names[0], entry):
        return build_indirect_common_util_payload_parts(
            context,
            api_id=api_id,
            entry=entry,
            workbook_path=workbook_path,
            sheet_names=sheet_names,
            reference_library=reference_library,
        )

    visual_notes, visual_unresolved = collect_excel_visual_warnings(workbook_path, sheet_names)
    rows = load_sheet_rows(SheetMatch(api=entry, workbook_path=workbook_path, sheet_name=sheet_names[0]))
    logic_entries = extract_business_logic_section(rows)
    sections = extract_api_spec_sections(rows)
    raw_appendix = dict(sections["rawAppendix"] or {})
    other_notes = list(raw_appendix.get("otherNotes") or [])
    other_notes.extend(visual_notes)
    raw_appendix["otherNotes"] = other_notes
    sections["rawAppendix"] = compact_raw_appendix(raw_appendix)
    business_logic = extract_business_logic_structure_from_rows(rows)
    business_logic_payload = {
        "steps": build_legacy_business_logic_steps(logic_entries),
        "fieldMappings": aggregate_field_mappings(business_logic.get("fieldMappings", []), sections["response"]),
        "lookupTables": business_logic.get("lookupTables", []),
        "errorCodeRules": merge_error_code_rules(sections.get("errorCodeRules", []), business_logic.get("errorCodeRules", [])),
        "runtimeDependencies": business_logic.get("runtimeDependencies", []),
        "dataSources": business_logic.get("dataSources", []),
        "sqlSpecs": business_logic.get("sqlSpecs", []),
        "legacyReferences": business_logic.get("legacyReferences", []),
        "prohibitedShortcuts": business_logic.get("prohibitedShortcuts", []),
    }
    business_logic_payload = {
        key: value
        for key, value in business_logic_payload.items()
        if key in {"steps", "fieldMappings", "lookupTables", "errorCodeRules", "runtimeDependencies", "dataSources", "sqlSpecs", "legacyReferences", "prohibitedShortcuts"}
        and not (key == "lookupTables" and not value)
    }
    business_logic_payload = apply_reference_hints_to_business_logic(
        business_logic_payload,
        reference_library=reference_library or {},
        backend_apis=sections["backendApis"],
        raw_appendix=sections["rawAppendix"],
    )
    source = build_api_source(context, workbook_path=workbook_path, sheet_names=sheet_names)
    api_spec = build_api_spec_payload(
        context=context,
        api_id=api_id,
        entry=entry,
        workbook_path=workbook_path,
        sheet_names=sheet_names,
        business_logic_payload=business_logic_payload,
        request_fields=sections["request"],
        response_fields=sections["response"],
        mock_examples=sections["mockExamples"],
        backend_apis=sections["backendApis"],
        raw_appendix=sections["rawAppendix"],
        additional_unresolved=visual_unresolved + list(business_logic.get("dependencyUnresolved", [])),
    )
    api_spec = run_spec_review_gate(
        context=context,
        api_id=api_id,
        api_spec=api_spec,
        expected_request_fields=list(sections["request"]),
        expected_response_fields=list(sections["response"]),
        expected_mock_examples=list(sections["mockExamples"]),
    )
    validate_api_spec_payload(api_spec)
    return {
        "source": source,
        "apiSpec": api_spec,
        "sourceFingerprint": build_source_fingerprint(
            context=context,
            api_category=entry.category,
            api_name=entry.name,
            source=source,
            request_fields=sections["request"],
            response_fields=sections["response"],
            mock_examples=sections["mockExamples"],
            backend_apis=sections["backendApis"],
            business_logic_payload=business_logic_payload,
            raw_appendix=sections["rawAppendix"],
            additional_unresolved=visual_unresolved + list(business_logic.get("dependencyUnresolved", [])),
        ),
    }


def reconcile_execution(
    context: ExecutionContext,
    *,
    initial: bool,
    existing_items: list[dict[str, Any]] | None = None,
) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    api_entries = extract_api_entries(context.docx_path)
    reference_library = load_reference_library(context.agent_dir, context.function_code)
    existing_items = existing_items or []
    existing_manifest_map = load_manifest_map(context, existing_items)
    existing_map = {
        item["apiId"]: normalize_existing_item(context, item, existing_manifest_map.get(item["apiId"])) for item in existing_items
    }

    reconciled_items: list[dict[str, Any]] = []
    current_api_ids: set[str] = set()
    progress_messages: list[str] = []
    if initial:
        progress_messages.append(
            f"已初始化功能目录 `{context.execution_id}`，来源文档 `{context.docx_path.name}`，共识别 {len(api_entries)} 支 API。"
        )

    for entry in api_entries:
        base_item = create_checklist_item(context, entry)
        api_id = base_item["apiId"]
        current_api_ids.add(api_id)
        previous = existing_map.get(api_id)
        item = dict(previous) if previous else base_item
        item["codeReset"] = False
        item["apiCategory"] = entry.category
        item["apiName"] = entry.name

        workbook_path: Path | None = None
        sheet_names: list[str] = []
        try:
            workbook_path = choose_workbook_for_category(context.agent_dir, entry.category, context.function_code)
            sheet_names = find_sheet_names(workbook_path, entry)
            payload_parts = build_api_payload_parts(
                context,
                api_id=api_id,
                entry=entry,
                workbook_path=workbook_path,
                sheet_names=sheet_names,
                reference_library=reference_library,
            )
            item["source"] = payload_parts["source"]
            item["sourceFingerprint"] = payload_parts["sourceFingerprint"]
            item["blockReason"] = None

            if previous is None:
                item["status"] = "pending"
                item["codeStatus"] = "pending"
                item["codePhase"] = "pending"
                item["codeBlockReason"] = None
            else:
                previous_fingerprint = clean_text(previous.get("sourceFingerprint"))
                if previous.get("status") == "retired":
                    item["status"] = "pending"
                    item["codeStatus"] = "pending"
                    item["codePhase"] = "pending"
                    item["codeBlockReason"] = None
                    item["codeReset"] = True
                    cleanup_code_artifacts(context, api_id)
                elif previous_fingerprint != item["sourceFingerprint"]:
                    item["status"] = "pending"
                    item["codeStatus"] = "pending"
                    item["codePhase"] = "pending"
                    item["codeBlockReason"] = None
                    item["codeReset"] = True
                    cleanup_code_artifacts(context, api_id)
                else:
                    refresh_needed, refresh_reason = needs_api_spec_refresh(context, previous)
                    if refresh_needed:
                        item["status"] = "pending"
                        item["codeStatus"] = "pending"
                        item["codePhase"] = "pending"
                        item["codeBlockReason"] = None
                        item["codeReset"] = True
                        cleanup_code_artifacts(context, api_id)
                    else:
                        preserved_status = previous.get("status") or "done"
                        item["status"] = preserved_status
                        item["blockReason"] = previous.get("blockReason")
                        item["codeStatus"] = previous.get("codeStatus") or default_code_status(item["status"])
                        item["codePhase"] = previous.get("codePhase") or default_code_phase(item["status"])
                        item["codeBlockReason"] = previous.get("codeBlockReason")
        except SkillError as exc:
            item["source"] = build_api_source(context, workbook_path=workbook_path, sheet_names=sheet_names)
            item["sourceFingerprint"] = None
            item["status"] = "blocked" if exc.status == "blocked" else "error"
            item["blockReason"] = str(exc)
            item["codeStatus"] = "waiting_spec"
            item["codePhase"] = "waiting_spec"
            progress_messages.append(f"`{api_id}` 处理为 `{item['status']}`：{exc}")
        reconciled_items.append(item)

    for api_id, previous in existing_map.items():
        if api_id in current_api_ids:
            continue
        retired_item = dict(previous)
        retired_item["status"] = "retired"
        retired_item["blockReason"] = "当前 TSD 已不再包含该 API。"
        retired_item["codeStatus"] = "waiting_spec"
        retired_item["codePhase"] = "waiting_spec"
        retired_item["codeReset"] = True
        reconciled_items.append(retired_item)

    dump_checklist(context, reconciled_items)
    dump_manifests(context, reconciled_items)

    for message in progress_messages:
        append_progress(context.paths.progress_path, message)

    execution_state = write_execution_state(
        context,
        reconciled_items,
        status=derive_execution_status(reconciled_items),
    )
    return execution_state, reconciled_items


def initialize_execution(context: ExecutionContext) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    return reconcile_execution(context, initial=True)


def load_or_initialize_execution(context: ExecutionContext) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    if not context.paths.execution_state_path.exists():
        return initialize_execution(context)

    execution_state = load_json(context.paths.execution_state_path)
    checklist_payload = load_json(context.paths.checklist_path)
    if execution_state.get("executionId") != context.execution_id:
        raise SkillError("execution-state.json 与当前功能编号不一致。")
    return reconcile_execution(
        context,
        initial=False,
        existing_items=list(checklist_payload.get("items") or []),
    )


def choose_target_item(checklist_items: list[dict[str, Any]], requested_api_id: str | None) -> dict[str, Any] | None:
    if requested_api_id:
        for item in checklist_items:
            if item["apiId"] == requested_api_id:
                if item["status"] == "retired":
                    raise SkillError(f"api-id 已标记 retired：{requested_api_id}")
                return item
        raise SkillError(f"api-id not found in checklist: {requested_api_id}")

    for status in ("in_progress", "pending", "error"):
        for item in checklist_items:
            if item["status"] == status:
                return item
    return None


def update_checklist_item(checklist_items: list[dict[str, Any]], updated_item: dict[str, Any]) -> None:
    for index, item in enumerate(checklist_items):
        if item["apiId"] == updated_item["apiId"]:
            checklist_items[index] = updated_item
            return
    raise SkillError(f"cannot update unknown apiId: {updated_item['apiId']}")


def finalize_item(context: ExecutionContext, checklist_items: list[dict[str, Any]], item: dict[str, Any]) -> None:
    update_checklist_item(checklist_items, item)
    dump_checklist(context, checklist_items)
    dump_manifest_if_changed(context, item)


def mark_item_in_progress(context: ExecutionContext, checklist_items: list[dict[str, Any]], item: dict[str, Any]) -> None:
    item["status"] = "in_progress"
    item["blockReason"] = None
    item["codeStatus"] = "waiting_spec"
    item["codePhase"] = "waiting_spec"
    item["codeBlockReason"] = None
    finalize_item(context, checklist_items, item)
    write_execution_state(
        context,
        checklist_items,
        status="running",
    )


def remove_previous_api_spec_if_needed(current_output: Path) -> None:
    current_dir = current_output.resolve().parent
    for candidate in current_dir.glob("*_API_Spec.json"):
        if candidate.resolve() == current_output.resolve():
            continue
        candidate.unlink()


def complete_item(
    context: ExecutionContext,
    checklist_items: list[dict[str, Any]],
    item: dict[str, Any],
    *,
    source: dict[str, Any],
    source_fingerprint: str,
    api_spec_path: Path,
) -> None:
    item["status"] = "done"
    item["source"] = source
    item["sourceFingerprint"] = source_fingerprint
    item["blockReason"] = None
    if clean_text(item.get("codeStatus")) in {"waiting_spec", ""}:
        item["codeStatus"] = "pending"
        item["codePhase"] = "pending"
        item["codeBlockReason"] = None
    finalize_item(context, checklist_items, item)
    append_progress(
        context.paths.progress_path,
        f"已完成 `{item['apiId']}`，并写出 `{api_spec_path.name}`。",
    )
    write_execution_state(
        context,
        checklist_items,
        status=derive_execution_status(checklist_items),
    )


def fail_item(
    context: ExecutionContext,
    checklist_items: list[dict[str, Any]],
    item: dict[str, Any],
    *,
    error: SkillError,
) -> None:
    item["status"] = "blocked" if error.status == "blocked" else "error"
    item["blockReason"] = str(error)
    item["codeStatus"] = "waiting_spec"
    item["codePhase"] = "waiting_spec"
    item["codeBlockReason"] = None
    finalize_item(context, checklist_items, item)
    append_progress(context.paths.progress_path, f"处理 `{item['apiId']}` 时停止：{error}")
    write_execution_state(
        context,
        checklist_items,
        status=item["status"],
    )


def print_status(context: ExecutionContext, checklist_items: list[dict[str, Any]]) -> None:
    summary = summarize_checklist(checklist_items)
    print(f"executionId: {context.execution_id}")
    print(f"functionCode: {context.function_code or 'null'}")
    print(f"projectRoot: {context.project_root.as_posix()}")
    print(f"docxPath: {context.docx_path.as_posix()}")
    print(f"stateRoot: {context.state_root.as_posix()}")
    print(
        "summary: "
        f"total={summary['total']}, pending={summary['pending']}, in_progress={summary['in_progress']}, "
        f"done={summary['done']}, blocked={summary['blocked']}, error={summary['error']}, retired={summary['retired']}"
    )


def process_api(context: ExecutionContext, checklist_items: list[dict[str, Any]], item: dict[str, Any]) -> None:
    mark_item_in_progress(context, checklist_items, item)
    entry = ApiEntry(category=item["apiCategory"], name=item["apiName"], description="")
    workbook_path = choose_workbook_for_category(context.agent_dir, entry.category, context.function_code)
    sheet_names = find_sheet_names(workbook_path, entry)
    reference_library = load_reference_library(context.agent_dir, context.function_code)
    payload_parts = build_api_payload_parts(
        context,
        api_id=item["apiId"],
        entry=entry,
        workbook_path=workbook_path,
        sheet_names=sheet_names,
        reference_library=reference_library,
    )
    api_spec_path = context.paths.api_spec_path(item["apiId"], function_code=context.function_code, tsd_path=context.docx_path)
    api_spec = payload_parts["apiSpec"]

    remove_previous_api_spec_if_needed(api_spec_path)
    dump_json(api_spec_path, api_spec)
    complete_item(
        context,
        checklist_items,
        item,
        source=payload_parts["source"],
        source_fingerprint=payload_parts["sourceFingerprint"],
        api_spec_path=api_spec_path,
    )


def main() -> int:
    configure_stdio()
    args = parse_args()
    try:
        context = build_context(args)
        _, checklist_items = load_or_initialize_execution(context)
        reference_warning = build_reference_library_warning(context.agent_dir, context.function_code)
        if reference_warning:
            append_progress(context.paths.progress_path, reference_warning)
            print(f"[WARN] {reference_warning}", file=sys.stderr)
        target_item = choose_target_item(checklist_items, args.api_id)
        if target_item is None:
            final_status = derive_execution_status(checklist_items)
            write_execution_state(
                context,
                checklist_items,
                status=final_status,
            )
            refresh_batch_pointer(context, preferred_function_code=context.execution_id)
            append_progress(
                context.paths.progress_path,
                "所有 API 都已处理完成。" if final_status == "done" else "执行已停止，当前仅剩被阻塞的 API。",
            )
            print_status(context, checklist_items)
            return 0 if final_status == "done" else 1

        if target_item["status"] == "done":
            refresh_batch_pointer(context, preferred_function_code=context.execution_id)
            print(f"api-id already completed: {target_item['apiId']}")
            api_spec_path = resolve_existing_api_spec_path(context, target_item["apiId"], require_exists=True)
            print(f"api-spec: {normalize_persisted_path(api_spec_path, project_root=context.project_root)}")
            return 0

        try:
            process_api(context, checklist_items, target_item)
        except SkillError as exc:
            fail_item(context, checklist_items, target_item, error=exc)
            refresh_batch_pointer(context, preferred_function_code=context.execution_id)
            print_status(context, checklist_items)
            print(f"[ERROR] {exc}", file=sys.stderr)
            return 1

        refresh_batch_pointer(context, preferred_function_code=context.execution_id)
        print_status(context, checklist_items)
        api_spec_path = resolve_existing_api_spec_path(context, target_item["apiId"], require_exists=True)
        print(f"api-spec: {normalize_persisted_path(api_spec_path, project_root=context.project_root)}")
        return 0
    except SkillError as exc:
        print(f"[ERROR] {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
