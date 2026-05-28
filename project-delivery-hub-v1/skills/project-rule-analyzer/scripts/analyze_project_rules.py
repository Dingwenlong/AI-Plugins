#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import sys
from pathlib import Path
from typing import Any
from zipfile import BadZipFile

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
    from openpyxl import load_workbook
except ImportError as exc:
    missing_module = (getattr(exc, "name", "") or "").split(".", 1)[0]
    package_name = {
        "docx": "python-docx",
        "openpyxl": "openpyxl",
    }.get(missing_module, missing_module or "required package")
    raise SystemExit(
        "project-rule-analyzer 缺少 Python 依赖："
        f"{package_name}。请在当前解释器安装后重试，例如："
        f"python -m pip install {package_name}"
    ) from exc


OLE_MAGIC = bytes.fromhex("D0CF11E0A1B11AE1")
CONFIG_FILENAME = "local-workspaces.json"
WORKSPACE_ROOT_ENV_KEYS = ("PROJECT_WORKSPACE_ROOT",)
WORKSPACE_KEY_ENV_KEYS = ("PROJECT_WORKSPACE_KEY",)
RULES_ROOT_ENV_KEYS = ("PROJECT_RULES_ROOT",)
RULE_CATEGORIES = {
    "api-contract",
    "api-detail-workbook",
    "delivery-format",
    "sequence-diagram",
    "code-guidelines",
    "sql-fixture",
    "field-kb",
    "ut-report",
}
CODE_GUIDELINE_CATALOG_RELATIVE_PATH = "rules/code-guidelines/catalog.json"
CODE_GUIDELINE_AUDIENCE_SCOPES = ["frontstage", "midBackoffice", "shared", "unknown"]


def clean_text(value: object) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    return "\n".join(line.strip() for line in text.split("\n") if line.strip())


def first_env(keys: tuple[str, ...]) -> str | None:
    for key in keys:
        value = clean_text(os.environ.get(key))
        if value:
            return value
    return None


def resolve_config_path(value: str | None, *, base: Path | None = None) -> Path | None:
    text = clean_text(value)
    if not text:
        return None
    path = Path(text).expanduser()
    if not path.is_absolute() and base is not None:
        path = base / path
    return path.resolve()


def slugify(value: str) -> str:
    text = re.sub(r"\s+", "-", clean_text(value).casefold())
    text = re.sub(r"[^a-z0-9._-]+", "-", text)
    text = re.sub(r"-{2,}", "-", text)
    return text.strip("-") or "project-rule"


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return f"sha256:{digest.hexdigest()}"


def find_plugin_root(start: Path | None = None) -> Path | None:
    current = (start or Path(__file__).resolve()).resolve()
    if current.is_file():
        current = current.parent
    for candidate in [current, *current.parents]:
        if (candidate / ".codex-plugin" / "plugin.json").exists():
            return candidate
    return None


def load_workspace_config() -> dict[str, Any]:
    plugin_root = find_plugin_root()
    if plugin_root is None:
        return {}
    path = plugin_root / "references" / CONFIG_FILENAME
    if not path.exists():
        return {}
    payload = json.loads(path.read_text(encoding="utf-8-sig"))
    return payload if isinstance(payload, dict) else {}


def resolve_rules_root(args: argparse.Namespace) -> Path:
    if clean_text(args.rules_root):
        return Path(args.rules_root).expanduser().resolve()
    env_root = first_env(RULES_ROOT_ENV_KEYS)
    if env_root:
        return Path(env_root).expanduser().resolve()
    config = load_workspace_config()
    workspaces = config.get("workspaces") if isinstance(config.get("workspaces"), dict) else {}
    workspace_key = clean_text(args.workspace_key) or first_env(WORKSPACE_KEY_ENV_KEYS) or clean_text(config.get("defaultWorkspace"))
    workspace = workspaces.get(workspace_key) if workspace_key else None
    if isinstance(workspace, dict):
        workspace_root = resolve_config_path(args.workspace_root or first_env(WORKSPACE_ROOT_ENV_KEYS) or clean_text(workspace.get("workspaceRoot")))
        raw_rules_root = clean_text(workspace.get("rulesRoot"))
        if raw_rules_root:
            resolved_rules_root = resolve_config_path(raw_rules_root, base=workspace_root)
            if resolved_rules_root is not None:
                return resolved_rules_root
        raw_agent_root = clean_text(workspace.get("agentRoot"))
        if raw_agent_root and workspace_key:
            agent_root = resolve_config_path(raw_agent_root, base=workspace_root)
            if agent_root is not None:
                return agent_root / "project-rules" / workspace_key
    project_root = Path(args.project_root).expanduser().resolve()
    return project_root / ".agent" / "project-rules" / (workspace_key or "default")


def read_docx(path: Path) -> tuple[str, list[dict[str, str]]]:
    document = Document(path)
    chunks: list[dict[str, str]] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = clean_text(paragraph.text)
        if text:
            chunks.append({"locator": f"paragraph:{index}", "text": text})
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [clean_text(cell.text) for cell in row.cells if clean_text(cell.text)]
            if cells:
                chunks.append({"locator": f"table:{table_index}:row:{row_index}", "text": " | ".join(cells)})
    return "\n".join(chunk["text"] for chunk in chunks), chunks


def read_xlsx(path: Path) -> tuple[str, list[dict[str, str]]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    chunks: list[dict[str, str]] = []
    for sheet in workbook.worksheets:
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            cells = [clean_text(value) for value in row if clean_text(value)]
            if cells:
                chunks.append({"locator": f"sheet:{sheet.title}:row:{row_index}", "text": " | ".join(cells)})
    return "\n".join(chunk["text"] for chunk in chunks), chunks


def read_source(path: Path) -> tuple[str, list[dict[str, str]], str]:
    suffix = path.suffix.casefold()
    if path.read_bytes()[:8] == OLE_MAGIC:
        return "", [], "ole_or_encrypted"
    if suffix == ".docx":
        try:
            text, chunks = read_docx(path)
            return text, chunks, "docx"
        except (PackageNotFoundError, BadZipFile, ValueError):
            return "", [], "docx_unreadable"
    if suffix in {".xlsx", ".xlsm"}:
        text, chunks = read_xlsx(path)
        return text, chunks, "xlsx"
    if suffix in {".md", ".markdown", ".txt"}:
        text = path.read_text(encoding="utf-8-sig")
        chunks = [{"locator": f"line:{index}", "text": clean_text(line)} for index, line in enumerate(text.splitlines(), start=1) if clean_text(line)]
        return clean_text(text), chunks, "text"
    if suffix == ".json":
        payload = json.loads(path.read_text(encoding="utf-8-sig"))
        rendered = json.dumps(payload, ensure_ascii=False, indent=2)
        return rendered, [{"locator": "$", "text": rendered}], "json"
    return "", [], "unsupported"


def load_catalog(path: Path, workspace_key: str | None) -> dict[str, Any]:
    if path.exists():
        payload = json.loads(path.read_text(encoding="utf-8-sig"))
        if isinstance(payload, dict):
            return payload
    return {
        "schemaVersion": "1.0.0",
        "workspaceKey": workspace_key,
        "projectName": workspace_key or "project",
        "activeReviewStatuses": ["approved", "active"],
        "defaultNewRuleStatus": "draft",
        "defaults": {},
        "assets": {},
        "rules": [],
    }


def load_code_guideline_catalog(path: Path) -> dict[str, Any]:
    if path.exists():
        payload = json.loads(path.read_text(encoding="utf-8-sig"))
        if isinstance(payload, dict):
            payload.setdefault("rules", [])
            return payload
    return {
        "schemaVersion": "1.0.0",
        "sourceName": "Project code guidelines",
        "version": "external",
        "sourceStatus": "project_rule_analyzer",
        "rules": [],
    }


def sorted_unique_text(values: object) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    if isinstance(values, list):
        candidates = values
    else:
        candidates = []
    for value in candidates:
        text = clean_text(value)
        if not text or text in seen:
            continue
        seen.add(text)
        result.append(text)
    return sorted(result)


def update_code_guideline_integration(catalog: dict[str, Any], rules_root: Path, entry: dict[str, Any], args: argparse.Namespace) -> None:
    if entry.get("category") != "code-guidelines":
        return

    active_statuses = set(sorted_unique_text(catalog.get("activeReviewStatuses"))) or {"approved", "active"}
    is_active = clean_text(entry.get("reviewStatus")) in active_statuses
    defaults = catalog.setdefault("defaults", {})
    if not isinstance(defaults, dict):
        defaults = {}
        catalog["defaults"] = defaults

    catalog_rel = clean_text(defaults.get("codeGuidelineCatalog")) or CODE_GUIDELINE_CATALOG_RELATIVE_PATH
    defaults["codeGuidelineCatalog"] = catalog_rel.replace("\\", "/")
    code_catalog_path = rules_root / defaults["codeGuidelineCatalog"]
    code_catalog_path.parent.mkdir(parents=True, exist_ok=True)
    code_catalog = load_code_guideline_catalog(code_catalog_path)

    existing_rules = [
        rule
        for rule in list(code_catalog.get("rules") or [])
        if isinstance(rule, dict) and clean_text(rule.get("ruleId")) != clean_text(entry.get("ruleId"))
    ]
    if is_active:
        action = "select_and_check_gaps" if clean_text(args.priority) in {"high", "critical"} else "load_on_demand"
        existing_rules.append(
            {
                "ruleId": clean_text(entry.get("ruleId")),
                "title": clean_text(entry.get("title")),
                "category": "project-code-guideline",
                "direction": "production-code",
                "ruleType": "blocking_gap" if action == "select_and_check_gaps" else "style_only",
                "audienceScopes": CODE_GUIDELINE_AUDIENCE_SCOPES,
                "featureTriggers": ["always"],
                "loadPath": clean_text(entry.get("markdownPath")) or clean_text(entry.get("path")),
                "loadReason": "Project-rule-analyzer approved code guideline.",
                "action": action,
            }
        )
        code_catalog["sourceStatus"] = "active"
    elif not existing_rules:
        code_catalog["sourceStatus"] = "draft_pending_approval"
    code_catalog["rules"] = sorted(existing_rules, key=lambda item: clean_text(item.get("ruleId")))
    code_catalog_path.write_text(json.dumps(code_catalog, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    packs = catalog.setdefault("rulePacks", {})
    if not isinstance(packs, dict):
        packs = {}
        catalog["rulePacks"] = packs
    pack = packs.get("apiCodeWriter") if isinstance(packs.get("apiCodeWriter"), dict) else None
    created_pack = pack is None
    if pack is None:
        pack = {
            "strict": True,
            "requiredRuleIds": [],
            "optionalRuleIds": [],
            "requiredAssets": [],
            "optionalAssets": [],
            "purpose": "Code writer must resolve project code guideline rules before selecting per-API guideline load hints.",
        }
        packs["apiCodeWriter"] = pack

    rule_id = clean_text(entry.get("ruleId"))
    for key in ("requiredRuleIds", "optionalRuleIds"):
        values = [value for value in sorted_unique_text(pack.get(key)) if value != rule_id]
        pack[key] = values
    if is_active and rule_id:
        target_key = "requiredRuleIds" if created_pack else "optionalRuleIds"
        pack[target_key] = sorted_unique_text([*list(pack.get(target_key) or []), rule_id])


def write_catalog(path: Path, catalog: dict[str, Any], entry: dict[str, Any]) -> None:
    rules = [rule for rule in list(catalog.get("rules") or []) if isinstance(rule, dict) and rule.get("ruleId") != entry["ruleId"]]
    rules.append(entry)
    catalog["rules"] = sorted(rules, key=lambda item: (str(item.get("category")), str(item.get("ruleId"))))
    path.write_text(json.dumps(catalog, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def write_rule_outputs(args: argparse.Namespace, rules_root: Path, source: Path) -> dict[str, Any]:
    category = clean_text(args.category)
    if category not in RULE_CATEGORIES:
        raise SystemExit(f"unsupported category: {category}")
    text, chunks, source_kind = read_source(source)
    source_hash = sha256_file(source)
    raw_target = rules_root / "sources" / "raw" / source.name
    raw_target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, raw_target)

    converted_dir = rules_root / "sources" / "converted"
    converted_dir.mkdir(parents=True, exist_ok=True)
    status_path = converted_dir / f"{source.stem}.source-status.json"
    status = {
        "schemaVersion": "1.0.0",
        "sourceFile": source.name,
        "sourceHash": source_hash,
        "sourceKind": source_kind,
        "status": "readable" if text else "source_unreadable",
    }
    status_path.write_text(json.dumps(status, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    if not text:
        return status

    rule_id = slugify(args.rule_id or source.stem)
    title = clean_text(args.title) or source.stem
    review_status = "approved" if args.approve else "draft"
    output_root = rules_root / "rules" / category
    output_root.mkdir(parents=True, exist_ok=True)
    json_path = output_root / f"{rule_id}.json"
    md_path = output_root / f"{rule_id}.md"
    source_locator = chunks[0]["locator"] if chunks else ""
    rule_payload = {
        "schemaVersion": "1.0.0",
        "ruleId": rule_id,
        "title": title,
        "category": category,
        "priority": args.priority,
        "reviewStatus": review_status,
        "sourceFile": source.name,
        "sourceHash": source_hash,
        "sourceLocator": source_locator,
        "sourceKind": source_kind,
        "extractedText": text,
        "chunks": chunks[:200],
    }
    json_path.write_text(json.dumps(rule_payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    md_lines = [
        f"# {title}",
        "",
        f"- ruleId: `{rule_id}`",
        f"- category: `{category}`",
        f"- reviewStatus: `{review_status}`",
        f"- sourceFile: `{source.name}`",
        f"- sourceHash: `{source_hash}`",
        f"- sourceLocator: `{source_locator}`",
        "",
        "## Extracted Rule Text",
        "",
        text,
        "",
    ]
    md_path.write_text("\n".join(md_lines), encoding="utf-8")
    catalog_path = rules_root / "catalog.json"
    catalog = load_catalog(catalog_path, args.workspace_key)
    rel_json = json_path.relative_to(rules_root).as_posix()
    entry = {
        "ruleId": rule_id,
        "category": category,
        "title": title,
        "path": rel_json,
        "format": "json",
        "markdownPath": md_path.relative_to(rules_root).as_posix(),
        "reviewStatus": review_status,
        "priority": args.priority,
        "sourceFile": source.name,
        "sourceHash": source_hash,
        "sourceLocator": source_locator,
    }
    update_code_guideline_integration(catalog, rules_root, entry, args)
    write_catalog(catalog_path, catalog, entry)
    return {"status": "rule_written", "rule": entry, "jsonPath": json_path.as_posix(), "markdownPath": md_path.as_posix()}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Analyze project rule documents into .agent/project-rules JSON and Markdown.")
    parser.add_argument("--project-root", default=str(Path.cwd()))
    parser.add_argument("--workspace-key")
    parser.add_argument("--workspace-root", help="共享工作区根目录；用于解析 local-workspaces.json 中的相对 agentRoot/rulesRoot。")
    parser.add_argument("--rules-root")
    parser.add_argument("--category", required=True, choices=sorted(RULE_CATEGORIES))
    parser.add_argument("--source", action="append", required=True)
    parser.add_argument("--rule-id")
    parser.add_argument("--title")
    parser.add_argument("--priority", default="normal", choices=["low", "normal", "high", "critical"])
    parser.add_argument("--approve", action="store_true")
    parser.add_argument("--format", choices=["text", "json"], default="text")
    return parser.parse_args()


def main() -> int:
    for stream in (sys.stdout, sys.stderr):
        reconfigure = getattr(stream, "reconfigure", None)
        if reconfigure:
            reconfigure(encoding="utf-8", errors="replace")
    args = parse_args()
    rules_root = resolve_rules_root(args)
    rules_root.mkdir(parents=True, exist_ok=True)
    results = []
    for raw_source in args.source:
        source = Path(raw_source).expanduser().resolve()
        if not source.exists() or not source.is_file():
            raise SystemExit(f"source not found: {source.as_posix()}")
        results.append(write_rule_outputs(args, rules_root, source))
    if args.format == "json":
        print(json.dumps(results, ensure_ascii=False, indent=2))
    else:
        for result in results:
            print(json.dumps(result, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
