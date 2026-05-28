#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import os
import shutil
import time
from collections import defaultdict
from pathlib import Path
from typing import Any
from zipfile import BadZipFile

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
    from openpyxl import load_workbook
except ImportError as exc:
    missing_module = (getattr(exc, "name", "") or "").split(".", 1)[0]
    package_name = {
        "docx": "python-docx",
        "openpyxl": "openpyxl",
    }.get(missing_module, missing_module or "required package")
    raise SystemExit(
        "reference-index-importer 缺少 Python 依赖："
        f"{package_name}。请在当前解释器安装后重试，例如："
        f"python -m pip install {package_name}"
    ) from exc

from reference_support import (
    REFERENCE_SCHEMA_VERSION,
    clean_text,
    extract_date_token,
    find_api_codes,
    find_external_keywords,
    find_field_names,
    find_framework_keywords,
    find_table_keys,
    normalize_match_key,
    project_relative_path,
    reference_catalog_path,
    reference_index_path,
    reference_indexes_dir,
    reference_raw_dir,
    reference_root,
    slugify,
    version_sort_key,
)
from chain_workspace import update_chain_status
from runtime import configure_stdio, dump_json, now_iso, resolve_agent_dir, resolve_project_root, sha256_file


class ZhArgumentParser(argparse.ArgumentParser):
    def __init__(self, *args: object, **kwargs: object) -> None:
        kwargs.setdefault("add_help", False)
        super().__init__(*args, **kwargs)
        self._positionals.title = "位置参数"
        self._optionals.title = "可选参数"

    def format_usage(self) -> str:
        return super().format_usage().replace("usage:", "用法：", 1)

    def format_help(self) -> str:
        text = super().format_help()
        return (
            text.replace("usage:", "用法：", 1)
            .replace("positional arguments:", "位置参数：")
            .replace("optional arguments:", "可选参数：")
            .replace("options:", "可选参数：")
            .replace("show this help message and exit", "显示此帮助并退出")
        )

    def error(self, message: str) -> None:
        self.print_usage()
        self.exit(2, f"{self.prog}: 错误：{message}\n")


def parse_args() -> argparse.Namespace:
    parser = ZhArgumentParser(description="导入精选外部参考资料到集中 .agent/reference/global，并生成可迁移索引。")
    parser.add_argument("-h", "--help", action="help", help="显示此帮助并退出")
    parser.add_argument("--project-root", required=True)
    parser.add_argument("--agent-dir", default=".agent")
    parser.add_argument("--agent-root", help="集中 .agent 根目录；优先级高于环境变量与插件本地配置。")
    parser.add_argument("--workspace-root", help="共享工作区根目录，例如 D:\\Repo\\Project。")
    parser.add_argument("--workspace-key", help="插件 local-workspaces.json 中的工作区 key，例如 PROJECT。")
    parser.add_argument("--rules-root", help="专案规则库根目录；优先级高于环境变量与 workspace 配置。")
    parser.add_argument("--function-code", help="可选：若本次导入只服务某一功能，则同步写入 chain-status。")
    parser.add_argument("--external-api-dir", required=True)
    parser.add_argument("--db-schema-dir", required=True)
    return parser.parse_args()


def ensure_existing_directory(path_value: str, *, label: str) -> Path:
    path = Path(path_value).expanduser().resolve()
    if not path.exists() or not path.is_dir():
        raise SystemExit(f"{label} 不存在或不是目录：{path.as_posix()}")
    return path


def logical_name_key(path: Path) -> str:
    stem = path.stem
    date_token = extract_date_token(path.name)
    if date_token:
        stem = stem.replace(date_token, "")
    stem = stem.replace("_", "-")
    stem = stem.replace(" ", "-")
    stem = stem.replace("--", "-")
    return slugify(stem)


def select_latest_by_group(paths: list[Path], group_key_fn) -> list[Path]:
    grouped: dict[str, list[Path]] = defaultdict(list)
    for path in paths:
        grouped[group_key_fn(path)].append(path)
    selected: list[Path] = []
    for group_paths in grouped.values():
        selected.append(sorted(group_paths, key=version_sort_key, reverse=True)[0])
    return sorted(selected, key=lambda item: item.name.casefold())


def select_external_api_files(root: Path) -> list[dict[str, Any]]:
    selections: list[dict[str, Any]] = []
    iris_dir = root / "IRIS"
    if iris_dir.exists():
        iris_files = [path for path in iris_dir.rglob("*.xlsx") if path.is_file()]
        iris_api_files = [path for path in iris_files if find_api_codes(path.stem)]
        for path in select_latest_by_group(iris_api_files, lambda item: sorted(find_api_codes(item.stem))[0]):
            selections.append({"path": path, "sourceRoot": root, "sourceLabel": "iris-api"})

    summary_files = [path for path in root.rglob("*.xlsx") if path.is_file() and "IRIS_OPENAPI_SUMMARY" in path.name.upper()]
    if summary_files:
        selections.append(
            {"path": sorted(summary_files, key=version_sort_key, reverse=True)[0], "sourceRoot": root, "sourceLabel": "iris-summary"}
        )

    bsp_files = [path for path in root.rglob("*.xlsx") if path.is_file() and path.name.upper().startswith("(ALL) BSP")]
    if bsp_files:
        selections.append(
            {"path": sorted(bsp_files, key=version_sort_key, reverse=True)[0], "sourceRoot": root, "sourceLabel": "bsp-api-list"}
        )

    if not selections:
        generic_doc_files = [
            path
            for path in root.rglob("*")
            if path.is_file() and path.suffix.lower() in {".xlsx", ".docx", ".pdf", ".pptx"}
        ]
        for path in sorted(generic_doc_files, key=lambda item: item.as_posix().casefold()):
            selections.append({"path": path, "sourceRoot": root, "sourceLabel": "external-doc"})
    return selections


def select_db_schema_files(root: Path) -> list[dict[str, Any]]:
    selections: list[dict[str, Any]] = []

    table_schema_files = [path for path in root.rglob("*.xlsx") if path.is_file() and "TABLESCHEMA" in path.name.upper()]
    for path in select_latest_by_group(table_schema_files, logical_name_key):
        selections.append({"path": path, "sourceRoot": root, "sourceLabel": "db-tableschema"})

    relation_files = [path for path in root.rglob("*.xlsx") if path.is_file() and "關連資料庫文件" in path.name]
    for path in select_latest_by_group(relation_files, logical_name_key):
        selections.append({"path": path, "sourceRoot": root, "sourceLabel": "db-relation"})

    trace_files = [path for path in root.rglob("*.xlsx") if path.is_file() and "TRACEABILITY" in path.name.upper()]
    if trace_files:
        selections.append(
            {"path": sorted(trace_files, key=version_sort_key, reverse=True)[0], "sourceRoot": root, "sourceLabel": "db-traceability"}
        )

    dictionary_files = [path for path in root.rglob("*.xlsx") if path.is_file() and "字典" in path.name]
    for path in select_latest_by_group(dictionary_files, logical_name_key):
        selections.append({"path": path, "sourceRoot": root, "sourceLabel": "db-dictionary"})

    return selections


def choose_section_title(text: str, current_section: str, *, fallback: str) -> str:
    normalized = clean_text(text)
    if not normalized:
        return current_section
    if len(normalized) <= 40 and not any(marker in normalized for marker in ("。", "，", ":", "：", ";", "；")):
        return normalized
    return current_section or fallback


def merge_match_key_map(target: dict[str, set[str]], source: dict[str, set[str]]) -> None:
    for key, values in source.items():
        target.setdefault(key, set()).update(values)


def collect_text_signals(
    text: str,
    *,
    locator_value: str,
    match_key_map: dict[str, set[str]],
    api_codes: set[str],
    table_names: set[str],
    field_names: set[str],
    keywords: set[str],
) -> None:
    rendered = clean_text(text)
    if not rendered:
        return
    extracted_api_codes = find_api_codes(rendered)
    extracted_tables = find_table_keys(rendered)
    extracted_fields = find_field_names(rendered)
    extracted_keywords = find_external_keywords(rendered) | find_framework_keywords(rendered)

    api_codes.update(extracted_api_codes)
    table_names.update(extracted_tables)
    field_names.update(extracted_fields)
    keywords.update(extracted_keywords)

    for key in extracted_api_codes | extracted_tables | extracted_keywords:
        match_key_map.setdefault(normalize_match_key(key), set()).add(locator_value)


def extract_xlsx_metadata(path: Path) -> dict[str, Any]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    api_codes: set[str] = find_api_codes(path.name)
    table_names: set[str] = find_table_keys(path.name)
    field_names: set[str] = find_field_names(path.stem)
    keywords: set[str] = find_external_keywords(path.name) | find_framework_keywords(path.name)
    sheet_match_keys: dict[str, set[str]] = {}
    sheet_names = list(workbook.sheetnames)
    try:
        for sheet_name in workbook.sheetnames:
            collect_text_signals(
                sheet_name,
                locator_value=sheet_name,
                match_key_map=sheet_match_keys,
                api_codes=api_codes,
                table_names=table_names,
                field_names=field_names,
                keywords=keywords,
            )
            sheet = workbook[sheet_name]
            for row in sheet.iter_rows(values_only=True):
                for cell in row:
                    collect_text_signals(
                        cell,
                        locator_value=sheet_name,
                        match_key_map=sheet_match_keys,
                        api_codes=api_codes,
                        table_names=table_names,
                        field_names=field_names,
                        keywords=keywords,
                    )
    finally:
        workbook.close()
    return {
        "sheetNames": sheet_names,
        "sheetMatchKeys": {key: sorted(values) for key, values in sorted(sheet_match_keys.items())},
        "apiCodes": sorted(api_codes),
        "tableNames": sorted(table_names),
        "fieldNames": sorted(field_names),
        "keywords": sorted(keywords),
    }


def extract_docx_metadata(path: Path) -> dict[str, Any]:
    document = Document(path)
    api_codes: set[str] = find_api_codes(path.name)
    table_names: set[str] = find_table_keys(path.name)
    field_names: set[str] = find_field_names(path.stem)
    keywords: set[str] = find_external_keywords(path.name) | find_framework_keywords(path.name)
    section_match_keys: dict[str, set[str]] = {}
    current_section = clean_text(document.core_properties.title) or path.stem

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue
        current_section = choose_section_title(text, current_section, fallback=path.stem)
        collect_text_signals(
            text,
            locator_value=current_section,
            match_key_map=section_match_keys,
            api_codes=api_codes,
            table_names=table_names,
            field_names=field_names,
            keywords=keywords,
        )

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                collect_text_signals(
                    cell.text,
                    locator_value=current_section,
                    match_key_map=section_match_keys,
                    api_codes=api_codes,
                    table_names=table_names,
                    field_names=field_names,
                    keywords=keywords,
                )

    return {
        "sectionMatchKeys": {key: sorted(values) for key, values in sorted(section_match_keys.items())},
        "apiCodes": sorted(api_codes),
        "tableNames": sorted(table_names),
        "fieldNames": sorted(field_names),
        "keywords": sorted(keywords),
    }


def extract_pdf_metadata(path: Path) -> dict[str, Any]:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        return {
            "pageMatchKeys": {},
            "apiCodes": sorted(find_api_codes(path.name)),
            "tableNames": sorted(find_table_keys(path.name)),
            "fieldNames": [],
            "keywords": sorted(find_external_keywords(path.name) | find_framework_keywords(path.name)),
        }

    reader = PdfReader(str(path))
    api_codes: set[str] = find_api_codes(path.name)
    table_names: set[str] = find_table_keys(path.name)
    field_names: set[str] = find_field_names(path.stem)
    keywords: set[str] = find_external_keywords(path.name) | find_framework_keywords(path.name)
    page_match_keys: dict[str, set[str]] = {}
    for index, page in enumerate(reader.pages, start=1):
        collect_text_signals(
            page.extract_text() or "",
            locator_value=str(index),
            match_key_map=page_match_keys,
            api_codes=api_codes,
            table_names=table_names,
            field_names=field_names,
            keywords=keywords,
        )
    return {
        "pageMatchKeys": {key: sorted(values) for key, values in sorted(page_match_keys.items())},
        "apiCodes": sorted(api_codes),
        "tableNames": sorted(table_names),
        "fieldNames": sorted(field_names),
        "keywords": sorted(keywords),
    }


def extract_reference_metadata(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    try:
        if suffix == ".xlsx":
            return extract_xlsx_metadata(path)
        if suffix == ".docx":
            return extract_docx_metadata(path)
        if suffix == ".pdf":
            return extract_pdf_metadata(path)
    except (BadZipFile, PackageNotFoundError, OSError, ValueError):
        pass
    return {
        "apiCodes": sorted(find_api_codes(path.name)),
        "tableNames": sorted(find_table_keys(path.name)),
        "fieldNames": sorted(find_field_names(path.stem)),
        "keywords": sorted(find_external_keywords(path.name) | find_framework_keywords(path.name)),
    }


def build_ref_id(category: str, path: Path, metadata: dict[str, Any]) -> str:
    category_dir_names = {
        "external_api": "external-api",
        "db_schema": "db-schema",
    }
    parent_suffix = ""
    path_identity_suffix = ""
    category_dir_name = category_dir_names.get(category)
    if category_dir_name and category_dir_name in path.parts:
        category_dir_index = path.parts.index(category_dir_name)
        relative_parent_parts = path.parts[category_dir_index + 1 : -1]
        relative_item_parts = path.parts[category_dir_index + 1 :]
        if relative_parent_parts:
            parent_suffix = f".{slugify('-'.join(relative_parent_parts))}"
        relative_item_key = "/".join(relative_item_parts)
        path_identity_suffix = f".{hashlib.sha1(relative_item_key.encode('utf-8')).hexdigest()[:8]}"

    if category == "external_api":
        api_codes = metadata.get("apiCodes") or []
        if api_codes:
            return f"external.{slugify(api_codes[0])}.{slugify(path.stem)}{parent_suffix}{path_identity_suffix}"
    if category == "db_schema":
        table_names = metadata.get("tableNames") or []
        if table_names:
            return f"db.{slugify(table_names[0])}.{slugify(path.stem)}{parent_suffix}{path_identity_suffix}"
    return f"{category.replace('_', '.')}.{slugify(path.stem)}{parent_suffix}{path_identity_suffix}"


def build_reference_record(
    *,
    project_root: Path,
    agent_dir: Path,
    copied_path: Path,
    category: str,
    source_label: str,
    metadata: dict[str, Any],
) -> dict[str, Any]:
    relative_path = project_relative_path(agent_dir, copied_path)
    match_keys = sorted(
        {
            normalize_match_key(item)
            for item in (
                list(metadata.get("apiCodes") or [])
                + list(metadata.get("tableNames") or [])
                + list(metadata.get("keywords") or [])
            )
            if clean_text(item)
        }
    )
    return {
        "refId": build_ref_id(category, copied_path, metadata),
        "category": category,
        "title": copied_path.stem,
        "relativePath": relative_path,
        "checksum": sha256_file(copied_path),
        "fileType": copied_path.suffix.lower().lstrip("."),
        "versionDateToken": extract_date_token(copied_path.name),
        "sourceLabel": source_label,
        "keywords": list(metadata.get("keywords") or []),
        "matchKeys": match_keys,
        "apiCodes": list(metadata.get("apiCodes") or []),
        "tableNames": list(metadata.get("tableNames") or []),
        "fieldNames": list(metadata.get("fieldNames") or []),
        "sheetNames": list(metadata.get("sheetNames") or []),
        "sheetMatchKeys": metadata.get("sheetMatchKeys") or {},
        "sectionMatchKeys": metadata.get("sectionMatchKeys") or {},
        "pageMatchKeys": metadata.get("pageMatchKeys") or {},
    }


def copy_selected_file(
    *,
    source_path: Path,
    source_root: Path,
    agent_dir: Path,
    category: str,
) -> Path:
    destination_root = reference_raw_dir(agent_dir, category)
    relative_source_path = source_path.resolve().relative_to(source_root.resolve())
    destination_path = destination_root / relative_source_path
    destination_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source_path, destination_path)
    destination_path.chmod(0o666)
    return destination_path


def write_reference_payload(path: Path, *, category: str | None, items: list[dict[str, Any]]) -> None:
    payload = {
        "schemaVersion": REFERENCE_SCHEMA_VERSION,
        "generatedAt": now_iso(),
        "category": category,
        "items": items,
    }
    dump_json(path, payload)


def import_reference_category(
    *,
    category: str,
    selections: list[dict[str, Any]],
    project_root: Path,
    agent_dir: Path,
) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for selection in selections:
        copied_path = copy_selected_file(
            source_path=selection["path"],
            source_root=selection["sourceRoot"],
            agent_dir=agent_dir,
            category=category,
        )
        metadata = extract_reference_metadata(copied_path)
        records.append(
            build_reference_record(
                project_root=project_root,
                agent_dir=agent_dir,
                copied_path=copied_path,
                category=category,
                source_label=selection["sourceLabel"],
                metadata=metadata,
            )
        )
    return sorted(records, key=lambda item: item["refId"])


def main() -> int:
    configure_stdio()
    args = parse_args()

    project_root = resolve_project_root(args.project_root)
    agent_dir = resolve_agent_dir(project_root, args.agent_dir, args.agent_root, args.workspace_root, args.workspace_key, args.rules_root)
    external_api_dir = ensure_existing_directory(args.external_api_dir, label="external-api-dir")
    db_schema_dir = ensure_existing_directory(args.db_schema_dir, label="db-schema-dir")

    target_root = reference_root(agent_dir)
    if target_root.exists():
        def on_remove_error(func, path, exc_info) -> None:
            last_error: Exception | None = None
            for _ in range(10):
                try:
                    os.chmod(path, 0o666)
                    func(path)
                    return
                except OSError as error:
                    last_error = error
                    time.sleep(0.5)
            if last_error is not None:
                raise last_error

        shutil.rmtree(target_root, onexc=on_remove_error)

    reference_indexes_dir(agent_dir).mkdir(parents=True, exist_ok=True)
    reference_raw_dir(agent_dir, "db_schema").mkdir(parents=True, exist_ok=True)
    reference_raw_dir(agent_dir, "external_api").mkdir(parents=True, exist_ok=True)

    external_records = import_reference_category(
        category="external_api",
        selections=select_external_api_files(external_api_dir),
        project_root=project_root,
        agent_dir=agent_dir,
    )
    db_records = import_reference_category(
        category="db_schema",
        selections=select_db_schema_files(db_schema_dir),
        project_root=project_root,
        agent_dir=agent_dir,
    )
    write_reference_payload(reference_index_path(agent_dir, "external_api"), category="external_api", items=external_records)
    write_reference_payload(reference_index_path(agent_dir, "db_schema"), category="db_schema", items=db_records)
    write_reference_payload(
        reference_catalog_path(agent_dir),
        category=None,
        items=sorted(external_records + db_records, key=lambda item: (item["category"], item["refId"])),
    )
    update_chain_status(
        agent_root=agent_dir,
        function_code=args.function_code,
        stage="reference",
        status="done",
        phase="indexed",
        message="reference indexes imported",
        project_root=project_root,
        workspace_key=args.workspace_key,
        artifacts={
            "referenceRoot": project_relative_path(agent_dir, reference_root(agent_dir)),
            "catalog": project_relative_path(agent_dir, reference_catalog_path(agent_dir)),
        },
    )

    print(f"referenceRoot: {reference_root(agent_dir).as_posix()}")
    print(f"externalApiImported: {len(external_records)}")
    print(f"dbSchemaImported: {len(db_records)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
