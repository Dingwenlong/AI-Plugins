from __future__ import annotations

import argparse
import copy
import html
import json
import os
import re
import sys
import urllib.parse
from collections import defaultdict
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    from docx.table import Table
    from PIL import Image, ImageDraw, ImageFont
except ImportError as exc:
    missing_module = (getattr(exc, "name", "") or "").split(".", 1)[0]
    package_name = {
        "docx": "python-docx",
        "PIL": "Pillow",
    }.get(missing_module, missing_module or "required package")
    raise SystemExit(
        "docx-unittest-report 缺少 Python 依赖："
        f"{package_name}。请在当前解释器安装后重试，例如："
        f"python -m pip install {package_name}"
    ) from exc

from chain_workspace import resolve_chain_workspace, update_chain_status, write_workspace_snapshot
from docx_report_utils import resolve_current_login_display_name, today_slash
from project_rules import resolve_asset_path
from trx_result_utils import build_test_lookup, find_latest_trx, parse_trx


DEFAULT_TEMPLATE_DOCX = (
    Path(__file__).resolve().parents[1] / "assets" / "API_UT自測報告模板_v3.5_20260512.docx"
)
FEATURE_TESTER_MAP_JSON = Path(__file__).resolve().parents[1] / "assets" / "feature-tester-map.json"
DEFAULT_REPORT_FONT_NAME = "微軟正黑體"


def normalize_text(value: Any) -> str:
    return " ".join(str(value or "").split()).strip()


def load_json(path: Path) -> dict[str, Any]:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except UnicodeDecodeError:
        return json.loads(path.read_text(encoding="utf-8-sig"))


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def set_run_report_font(run: Any, font_name: str = DEFAULT_REPORT_FONT_NAME) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)


def apply_report_font_to_document(document: Document, font_name: str = DEFAULT_REPORT_FONT_NAME) -> None:
    parts: list[Any] = [document]
    for section in document.sections:
        parts.extend([section.header, section.footer])

    for style in document.styles:
        if getattr(style, "type", None) is None:
            continue
        if hasattr(style, "font"):
            style.font.name = font_name
            if style.element is not None:
                r_pr = style.element.get_or_add_rPr()
                r_fonts = r_pr.rFonts
                if r_fonts is None:
                    r_fonts = OxmlElement("w:rFonts")
                    r_pr.append(r_fonts)
                for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                    r_fonts.set(qn(f"w:{attr}"), font_name)

    for part in parts:
        for paragraph in part.paragraphs:
            for run in paragraph.runs:
                set_run_report_font(run, font_name)
        for table in part.tables:
            for row in table.rows:
                for cell in unique_cells(row):
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_run_report_font(run, font_name)


def set_paragraph_after_auto(paragraph: Any) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:afterAutospacing"), "1")
    spacing.attrib.pop(qn("w:after"), None)


def resolve_context_root(raw_path: str) -> Path:
    path = Path(raw_path).expanduser().resolve()
    if not path.exists():
        raise SystemExit(f"context root not found: {path.as_posix()}")
    if not (path / "api-checklist.json").exists():
        raise SystemExit(f"api-checklist.json not found under context root: {path.as_posix()}")
    return path


def infer_repo_root(context_root: Path) -> Path:
    state_path = context_root / "execution-state.json"
    if state_path.exists():
        state = load_json(state_path)
        code_project_root = normalize_text(state.get("codeProjectRoot"))
        if code_project_root and code_project_root != ".":
            project_root = Path(code_project_root).expanduser()
            if project_root.is_absolute() and project_root.exists():
                return project_root.resolve()
    try:
        return context_root.parents[2]
    except IndexError as exc:
        raise SystemExit(f"unable to infer repo root from context root: {context_root.as_posix()}") from exc


def resolve_central_context_root(args: argparse.Namespace) -> Path:
    if args.context_root:
        return resolve_context_root(args.context_root)
    function_code = normalize_text(args.function_code)
    if not function_code:
        raise SystemExit("please provide --context-root or --function-code")
    project_root = Path(args.project_root).expanduser()
    if not project_root.is_absolute():
        project_root = (Path.cwd() / project_root).resolve()
    workspace = resolve_chain_workspace(
        project_root=project_root,
        agent_dir_arg=args.agent_dir,
        agent_root_arg=args.agent_root,
        rules_root_arg=args.rules_root,
        workspace_root_arg=args.workspace_root,
        workspace_key_arg=args.workspace_key,
        start_path=Path(__file__).resolve().parents[2],
    )
    write_workspace_snapshot(workspace)
    return resolve_context_root(str(workspace.context_root / function_code))


def extract_template_tail_token(path: Path) -> str:
    match = re.search(r"(\d+)(?!.*\d)", path.stem)
    return match.group(1) if match else ""


def extract_template_version(path: Path) -> int:
    matches = re.findall(r"\bv\s*(\d+)\b", path.stem, flags=re.IGNORECASE)
    if not matches:
        return -1
    return max(int(token) for token in matches)


def is_ut_report_template(path: Path) -> bool:
    if path.suffix.lower() != ".docx":
        return False
    stem = normalize_text(path.stem)
    stem_lower = stem.casefold()
    return "ut" in stem_lower and ("測試報告" in stem or "测试报告" in stem)


def rank_template_docx(path: Path) -> tuple[int, int, str, int, int, str]:
    tail_token = extract_template_tail_token(path)
    version = extract_template_version(path)
    return (
        1 if tail_token else 0,
        len(tail_token),
        tail_token,
        version,
        path.stat().st_mtime_ns,
        path.name.casefold(),
    )


def find_latest_ut_template_docx(repo_root: Path) -> Path:
    template_dir = repo_root / ".agent" / "Template"
    if not template_dir.exists():
        raise SystemExit(f"template directory not found: {template_dir.as_posix()}")

    docx_files = [path for path in template_dir.iterdir() if path.is_file() and path.suffix.lower() == ".docx"]
    if not docx_files:
        raise SystemExit(f"no DOCX templates found under: {template_dir.as_posix()}")

    preferred = [path for path in docx_files if is_ut_report_template(path)]
    fallback = [path for path in docx_files if "ut" in normalize_text(path.stem).casefold()]
    candidates = preferred or fallback or docx_files
    return max(candidates, key=rank_template_docx)


def resolve_template_docx(
    raw_path: str | None,
    repo_root: Path,
    *,
    rules_root_arg: str | None = None,
    workspace_key: str | None = None,
) -> Path:
    if normalize_text(raw_path):
        path = Path(raw_path).expanduser().resolve()
        if not path.exists():
            raise SystemExit(f"template DOCX not found: {path.as_posix()}")
        return path
    project_template = resolve_asset_path(
        "utReportTemplate",
        rules_root_arg=rules_root_arg,
        workspace_key=workspace_key,
        fallback=DEFAULT_TEMPLATE_DOCX,
    )
    if project_template is not None:
        return project_template
    if DEFAULT_TEMPLATE_DOCX.exists():
        return DEFAULT_TEMPLATE_DOCX
    return find_latest_ut_template_docx(repo_root)


def default_output_docx(context_root: Path, template_docx: Path) -> Path:
    report_date = date.today().strftime("%Y%m%d")
    return context_root / "ut-report" / f"{context_root.name}_API_UT 測試報告 {report_date}.docx"


def default_assets_dir(output_docx: Path) -> Path:
    return output_docx.parent / f"{output_docx.stem}_assets"


def render_status_summary(total: int, passed: int, failed: int, skipped: int, pending: int) -> str:
    return (
        f"測試運行已完成: 共 {total} 個測試 "
        f"({passed} 個已通過, {failed} 個失敗, {skipped} 個已跳過, {pending} 個待補)"
    )


def safe_file_name(value: str) -> str:
    token = normalize_text(value)
    if not token:
        return "unknown"
    return re.sub(r"[^A-Za-z0-9._-]+", "_", token)


def text_width(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont | ImageFont.FreeTypeFont) -> int:
    if not text:
        return 0
    left, _, right, _ = draw.textbbox((0, 0), text, font=font)
    return right - left


def truncate_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    font: ImageFont.ImageFont | ImageFont.FreeTypeFont,
    max_width: int,
) -> str:
    value = normalize_text(text)
    if not value:
        return ""
    if text_width(draw, value, font) <= max_width:
        return value

    ellipsis = "..."
    low = 0
    high = len(value)
    best = ellipsis
    while low <= high:
        middle = (low + high) // 2
        candidate = value[:middle].rstrip() + ellipsis
        if text_width(draw, candidate, font) <= max_width:
            best = candidate
            low = middle + 1
        else:
            high = middle - 1
    return best


def format_duration(seconds: float) -> str:
    if seconds >= 1:
        return f"{seconds:.1f} 秒"
    milliseconds = round(seconds * 1000)
    return f"{milliseconds} 毫秒"


def format_status_runtime(seconds: float) -> str:
    milliseconds = round(seconds * 1000)
    if milliseconds >= 1000:
        whole_seconds = milliseconds // 1000
        remain_ms = milliseconds % 1000
        if remain_ms:
            return f"{whole_seconds} 秒 {remain_ms} 毫秒"
        return f"{whole_seconds} 秒"
    return f"{milliseconds} 毫秒"


def summarize_tests(tests: list[dict[str, Any]]) -> dict[str, int]:
    return {
        "total": len(tests),
        "passed": sum(1 for test in tests if test.get("status") == "passed"),
        "failed": sum(1 for test in tests if test.get("status") == "failed"),
        "skipped": sum(1 for test in tests if test.get("status") == "skipped"),
        "pending": sum(1 for test in tests if test.get("status") == "pending"),
    }


def not_passed_count(counts: dict[str, int]) -> int:
    return max(0, int(counts.get("total", 0) or 0) - int(counts.get("passed", 0) or 0))


def format_pass_fail_text(counts: dict[str, int]) -> str:
    return f"通過 {counts.get('passed', 0)} 項 / 不通過 {not_passed_count(counts)} 項"


def format_report_result_text(counts: dict[str, int]) -> str:
    return (
        f"通過 {counts.get('passed', 0)} 項 / "
        f"不通過 {not_passed_count(counts)} 項 / "
        f"不適用 {counts.get('notApplicable', 0)} 項"
    )


def find_font_path() -> str | None:
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\msyhbd.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return None


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    font_path = find_font_path()
    if font_path:
        return ImageFont.truetype(font_path, size=size)
    return ImageFont.load_default()


@dataclass
class ApiVisualSummary:
    api_id: str
    api_name: str
    api_display_name: str
    mock_examples: list[str]
    not_applicable_count: int
    checklist_section_applicability: dict[str, str]
    checklist_item_applicability: dict[str, list[str]]
    unit_tests: list[dict[str, Any]]
    integration_tests: list[dict[str, Any]]
    image_path: str
    uses_sql: bool = False
    checklist_counts: dict[str, int] = field(default_factory=dict)
    image_display_width_inches: float = 5.0

    @property
    def all_tests(self) -> list[dict[str, Any]]:
        return self.unit_tests + self.integration_tests

    @property
    def counts(self) -> dict[str, int]:
        return summarize_tests(self.all_tests)

    @property
    def has_enterprise_config_runtime_validation(self) -> bool:
        for test in self.integration_tests:
            if is_enterprise_config_runtime_test(test):
                return True
        return False


def is_enterprise_config_runtime_test(test: dict[str, Any]) -> bool:
    test_name = normalize_text(test.get("testName", "")).casefold()
    return "enterpriseconfig" in test_name or "configured" in test_name or "設定連線" in test_name


def resolve_trx_dirs(repo_root: Path, context_root: Path, api_id: str, test_evidence: dict[str, Any]) -> tuple[str, str]:
    hints = test_evidence.get("trxHints") or {}
    default_base = repo_root / ".agent" / "report-results" / context_root.name / api_id
    unit_dir = normalize_text(hints.get("unit")) or (default_base / "unit").as_posix()
    integration_dir = normalize_text(hints.get("integration")) or (default_base / "integration").as_posix()
    return unit_dir, integration_dir


def load_existing_results(api_root: Path, layer_name: str) -> dict[str, Any] | None:
    results_path = api_root / "ut-report" / "report-job.results.json"
    if not results_path.exists():
        return None

    payload = load_json(results_path)
    source_results = payload.get("sourceResults") or {}
    layer_payload = source_results.get(layer_name) or {}
    tests = list(layer_payload.get("tests") or [])
    if not tests:
        return None

    summary = payload.get(f"{layer_name}Summary") or layer_payload.get("summary") or summarize_tests(tests)
    return {
        "trxPath": normalize_text(layer_payload.get("trxPath")) or normalize_text(payload.get("trxPath")),
        "tests": tests,
        "summary": summary,
    }


def load_api_mock_examples(api_root: Path) -> list[str]:
    spec_paths = sorted(api_root.glob("*_API_Spec.json"))
    for spec_path in spec_paths:
        payload = load_json(spec_path)
        examples = payload.get("mockExamples") or []
        scenarios: list[str] = []
        for index, example in enumerate(examples, start=1):
            scenario = normalize_text(
                example.get("scenario")
                or example.get("title")
                or example.get("description")
                or example.get("name")
            )
            scenarios.append(scenario or f"範例情境 {index}")
        if scenarios:
            return scenarios
    return []


def payload_has_sql_evidence(payload: Any) -> bool:
    if isinstance(payload, dict):
        for key, value in payload.items():
            normalized_key = normalize_text(key).casefold()
            if normalized_key in {
                "querycontracts",
                "sqlspecs",
                "querycontractsselected",
                "tables",
                "tablerefs",
                "dbtarget",
                "databasedependencies",
                "datadependencies",
                "serviceruntimevalidationplan",
            }:
                if payload_has_sql_evidence(value):
                    return True
            if normalized_key == "serviceruntimevalidationrequired" and str(value).strip().casefold() in {"true", "1", "yes"}:
                return True
            if "querycontract" in normalized_key or normalized_key in {"sqltext", "sql", "sqlspec"}:
                if value:
                    return True
            if normalized_key in {
                "serviceruntimevalidationplan",
                "runtimedependencies",
                "backendapis",
                "backendapi",
                "codehandoff",
                "businesslogic",
                "sqlfixture",
            }:
                if payload_has_sql_evidence(value):
                    return True
        return False
    if isinstance(payload, list):
        return any(payload_has_sql_evidence(item) for item in payload)
    if isinstance(payload, str):
        text = normalize_text(payload).casefold()
        return any(
            token in text
            for token in (
                "select ",
                "insert ",
                "update ",
                "delete ",
                " join ",
                ".dbo.",
                "sql",
                "connectionstrings.",
                "db ->",
                "mma ->",
                "dbtarget",
            )
        )
    return False


def load_api_uses_sql(api_root: Path) -> bool:
    for spec_path in sorted(api_root.glob("*_API_Spec.json")):
        try:
            if payload_has_sql_evidence(load_json(spec_path)):
                return True
        except Exception:
            continue
    for evidence_path in ("change-plan.json", "db-fixture-report.json", "seed-manifest.json", "table-checks.json"):
        path = api_root / evidence_path
        if not path.exists():
            continue
        try:
            if payload_has_sql_evidence(load_json(path)):
                return True
        except Exception:
            continue
    return False


def load_api_not_applicable_count(api_root: Path) -> int:
    classification_path = api_root / "ut-report" / "template-classification.json"
    if not classification_path.exists():
        return 0

    payload = load_json(classification_path)
    item_counts = payload.get("itemCounts") or (payload.get("summary") or {}).get("itemCounts") or {}
    raw_value = item_counts.get("not_applicable")
    if raw_value is not None:
        try:
            return max(0, int(raw_value))
        except (TypeError, ValueError):
            return 0

    count = 0
    for section in payload.get("sections") or []:
        for item in section.get("items") or []:
            if normalize_text(item.get("applicability")) == "not_applicable":
                count += 1
    return count


def extract_ut_section_id(value: str) -> str:
    match = re.search(r"\b(UT-\d{2})\b", normalize_text(value))
    return match.group(1) if match else ""


def load_api_checklist_applicability(api_root: Path) -> tuple[dict[str, str], dict[str, list[str]]]:
    classification_path = api_root / "ut-report" / "template-classification.json"
    if not classification_path.exists():
        return {}, {}

    payload = load_json(classification_path)
    section_applicability: dict[str, str] = {}
    item_applicability: dict[str, list[str]] = {}
    for section in payload.get("sections") or []:
        section_id = extract_ut_section_id(
            normalize_text(section.get("sectionId")) or normalize_text(section.get("title"))
        )
        if not section_id:
            continue
        section_value = normalize_text(section.get("applicability")) or "applicable"
        section_applicability[section_id] = section_value
        item_values: list[str] = []
        for item in section.get("items") or []:
            item_values.append(normalize_text(item.get("applicability")) or section_value)
        item_applicability[section_id] = item_values
    return section_applicability, item_applicability


def select_tests(
    trx_payload: dict[str, Any] | None,
    explicit_names: list[str],
    api_name: str,
    api_category: str,
) -> list[dict[str, Any]]:
    if not trx_payload:
        return []

    tests = trx_payload.get("tests") or []
    lookup = build_test_lookup(tests)
    selected: list[dict[str, Any]] = []
    for name in explicit_names:
        matched = lookup.get(normalize_text(name))
        if matched is None:
            normalized_name = normalize_text(name).casefold()
            normalized_method_name = normalized_name.rsplit(".", 1)[-1]
            matched = next(
                (
                    test
                    for test in tests
                    if normalize_text(test.get("testName", "")).casefold().endswith(normalized_name)
                    or normalize_text(test.get("testName", "")).casefold().rsplit(".", 1)[-1] == normalized_method_name
                ),
                None,
            )
        if matched is not None:
            selected.append(matched)
    if selected:
        return selected

    name_tokens = {
        normalize_text(api_name).replace("_", "").replace("-", "").casefold(),
        normalize_text(api_category).replace("_", "").replace("-", "").casefold(),
    }
    selected = [
        test
        for test in tests
        if any(token and token in normalize_text(test.get("testName", "")).replace("_", "").replace("-", "").casefold() for token in name_tokens)
    ]
    return selected


def load_api_visual_summary(
    repo_root: Path,
    context_root: Path,
    item: dict[str, Any],
    assets_dir: Path,
) -> ApiVisualSummary:
    api_id = item["apiId"]
    api_root = context_root / "apis" / api_id
    test_evidence_path = api_root / "test-evidence.json"
    test_evidence = load_json(test_evidence_path) if test_evidence_path.exists() else {}
    api_name = normalize_text(item.get("apiName", ""))
    api_category = normalize_text(item.get("apiCategory", ""))
    api_display_name = normalize_text(test_evidence.get("apiDisplayName")) or f"{api_category}/{api_name}".strip("/")

    unit_payload = load_existing_results(api_root, "unitTest")
    integration_payload = load_existing_results(api_root, "integrationTest")

    unit_dir, integration_dir = resolve_trx_dirs(repo_root, context_root, api_id, test_evidence)
    unit_trx = find_latest_trx(unit_dir)
    integration_trx = find_latest_trx(integration_dir)
    if unit_trx:
        unit_payload = parse_trx(unit_trx)
    if integration_trx:
        integration_payload = parse_trx(integration_trx)

    test_names = test_evidence.get("testNames") or {}
    unit_tests = select_tests(unit_payload, list(test_names.get("unit") or []), api_name, api_category)
    integration_tests = select_tests(integration_payload, list(test_names.get("integration") or []), api_name, api_category)

    image_path = assets_dir / f"{safe_file_name(api_id)}.png"
    checklist_section_applicability, checklist_item_applicability = load_api_checklist_applicability(api_root)
    return ApiVisualSummary(
        api_id=api_id,
        api_name=api_name,
        api_display_name=api_display_name,
        mock_examples=load_api_mock_examples(api_root),
        not_applicable_count=load_api_not_applicable_count(api_root),
        checklist_section_applicability=checklist_section_applicability,
        checklist_item_applicability=checklist_item_applicability,
        unit_tests=unit_tests,
        integration_tests=integration_tests,
        image_path=image_path.as_posix(),
        uses_sql=load_api_uses_sql(api_root),
    )


def build_overall_visual_summary(
    function_code: str,
    function_display_name: str,
    api_summaries: list[ApiVisualSummary],
    assets_dir: Path,
) -> ApiVisualSummary:
    unit_tests: list[dict[str, Any]] = []
    integration_tests: list[dict[str, Any]] = []
    for summary in api_summaries:
        # The overall screenshot must contain only tests collected for the current function APIs.
        unit_tests.extend(summary.unit_tests)
        integration_tests.extend(summary.integration_tests)

    image_path = assets_dir / f"{safe_file_name(function_code)}_UnitTest_Summary.png"
    return ApiVisualSummary(
        api_id=function_code,
        api_name=function_code,
        api_display_name=function_display_name,
        mock_examples=[],
        not_applicable_count=0,
        checklist_section_applicability={},
        checklist_item_applicability={},
        unit_tests=unit_tests,
        integration_tests=integration_tests,
        image_path=image_path.as_posix(),
        uses_sql=any(summary.uses_sql for summary in api_summaries),
    )


def group_tests(tests: list[dict[str, Any]]) -> list[tuple[str, list[dict[str, Any]]]]:
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for test in tests:
        full_name = normalize_text(test.get("testName", ""))
        parts = [part for part in full_name.split(".") if part]
        if len(parts) >= 2:
            class_name = parts[-2]
        else:
            class_name = "Tests"
        grouped[class_name].append(test)
    return sorted(grouped.items())


def draw_check_or_x(
    draw: ImageDraw.ImageDraw,
    center_x: int,
    center_y: int,
    color: str,
    *,
    filled: bool,
    mark: str,
    radius: int = 8,
) -> None:
    fill = color if filled else "#252526"
    outline = color
    draw.rounded_rectangle(
        (center_x - radius, center_y - radius, center_x + radius, center_y + radius),
        radius=radius,
        fill=fill,
        outline=outline,
        width=2,
    )
    mark_color = "#ffffff" if filled else color
    if mark == "check":
        draw.line(
            (
                center_x - 4,
                center_y,
                center_x - 1,
                center_y + 4,
                center_x + 5,
                center_y - 4,
            ),
            fill=mark_color,
            width=2,
            joint="curve",
        )
        return
    draw.line((center_x - 4, center_y - 4, center_x + 4, center_y + 4), fill=mark_color, width=2)
    draw.line((center_x + 4, center_y - 4, center_x - 4, center_y + 4), fill=mark_color, width=2)


def draw_warning_icon(draw: ImageDraw.ImageDraw, x: int, y: int, color: str) -> None:
    draw.polygon([(x, y + 10), (x + 8, y - 4), (x + 16, y + 10)], outline=color, fill="#252526")
    draw.text((x + 6, y), "!", fill=color, font=load_font(10))


def recommend_docx_image_width_inches(image_width_px: int) -> float:
    min_px = 1280
    max_px = 1900
    min_inches = 3.9
    max_inches = 5.0
    clamped = max(min_px, min(max_px, image_width_px))
    ratio = (clamped - min_px) / (max_px - min_px) if max_px > min_px else 1.0
    return round(min_inches + (max_inches - min_inches) * ratio, 2)


def draw_caret(draw: ImageDraw.ImageDraw, x: int, y: int, *, expanded: bool, color: str) -> None:
    if expanded:
        points = [(x, y), (x + 10, y), (x + 5, y + 6)]
    else:
        points = [(x + 2, y - 2), (x + 2, y + 8), (x + 8, y + 3)]
    draw.polygon(points, fill=color)


def css_px_for_label(value: str) -> int:
    ascii_count = sum(1 for char in value if ord(char) < 128)
    wide_count = len(value) - ascii_count
    return ascii_count * 7 + wide_count * 13


SOLID_CHECK_SVG = """<svg width="256" height="256" viewBox="0 0 256 256" xmlns="http://www.w3.org/2000/svg" fill="none"><circle cx="128" cy="128" r="104" fill="#83F380"/><path d="M78 132 L114 168 L182 94" stroke="#1F1F1F" stroke-width="20" stroke-linecap="round" stroke-linejoin="round"/></svg>"""
HOLLOW_CHECK_SVG = """<svg viewBox="0 0 1024 1024" xmlns="http://www.w3.org/2000/svg"><path d="M483.072 698.208c-5.728 0-11.264-2.016-15.616-5.792l-143.68-123.232c-10.048-8.64-11.232-23.776-2.592-33.856 8.64-10.016 23.776-11.2 33.856-2.592l101.632 87.136c13.152 11.296 32.992 9.76 44.288-3.392l203.936-237.824c8.672-10.016 23.808-11.2 33.856-2.592 10.048 8.64 11.232 23.776 2.592 33.856l-240 279.872c-4.16 4.832-10.016 7.808-16.384 8.32C484.288 698.176 483.68 698.208 483.072 698.208z" fill="#54c974"/><path d="M512 896.608c-212.096 0-384.64-172.544-384.64-384.64C127.392 299.936 299.936 127.392 512 127.392c212.064 0 384.608 172.544 384.608 384.608C896.608 724.064 724.064 896.608 512 896.608zM512 175.392C326.4 175.392 175.392 326.4 175.392 512c0 185.632 151.008 336.64 336.64 336.64 185.6 0 336.608-151.008 336.608-336.64C848.608 326.4 697.6 175.392 512 175.392z" fill="#54c974"/></svg>"""
BLUE_DIAMOND_EXCLAMATION_OUTLINE_SVG = """<svg width="256" height="256" viewBox="0 0 256 256" xmlns="http://www.w3.org/2000/svg" fill="none"><g transform="translate(128 128) rotate(45) translate(-128 -128)"><rect x="56" y="56" width="144" height="144" rx="18" fill="none" stroke="#51A2E3" stroke-width="16"/></g><rect x="118" y="76" width="20" height="78" rx="10" fill="#51A2E3"/><circle cx="128" cy="182" r="12" fill="#51A2E3"/></svg>"""
BLUE_DIAMOND_EXCLAMATION_FILLED_SVG = """<svg width="256" height="256" viewBox="0 0 256 256" xmlns="http://www.w3.org/2000/svg" fill="none"><g transform="translate(128 128) rotate(45) translate(-128 -128)"><rect x="56" y="56" width="144" height="144" rx="18" fill="#51A2E3"/></g><rect x="118" y="76" width="20" height="78" rx="10" fill="#FFFFFF"/><circle cx="128" cy="182" r="12" fill="#FFFFFF"/></svg>"""
RED_CLOSE_FILLED_SVG = """<svg width="256" height="256" viewBox="0 0 256 256" xmlns="http://www.w3.org/2000/svg" fill="none"><circle cx="128" cy="128" r="104" fill="#F36A77"/><path d="M92 92 L164 164" stroke="#FFFFFF" stroke-width="20" stroke-linecap="round"/><path d="M164 92 L92 164" stroke="#FFFFFF" stroke-width="20" stroke-linecap="round"/></svg>"""
RED_CLOSE_OUTLINE_SVG = """<svg width="256" height="256" viewBox="0 0 256 256" xmlns="http://www.w3.org/2000/svg" fill="none"><circle cx="128" cy="128" r="96" fill="none" stroke="#F36A77" stroke-width="18"/><path d="M96 96 L160 160" stroke="#F36A77" stroke-width="18" stroke-linecap="round"/><path d="M160 96 L96 160" stroke="#F36A77" stroke-width="18" stroke-linecap="round"/></svg>"""


def svg_data_uri(svg_text: str) -> str:
    encoded = urllib.parse.quote(svg_text, safe="(),-./:=;_")
    return f"data:image/svg+xml,{encoded}"


def build_test_explorer_rows(summary: ApiVisualSummary) -> tuple[list[dict[str, Any]], dict[str, Any] | None]:
    grouped_unit = group_tests(summary.unit_tests)
    grouped_integration = group_tests(summary.integration_tests)

    rows: list[dict[str, Any]] = []

    def append_group(
        title: str,
        tests: list[dict[str, Any]],
        source_label: str,
        class_groups: list[tuple[str, list[dict[str, Any]]]],
        group_key: str,
    ) -> None:
        if not tests:
            return
        total_duration = sum(float(test.get("durationSeconds", 0) or 0) for test in tests)
        rows.append(
            {
                "kind": "project",
                "indent": 0,
                "label": f"{title} ({len(tests)})",
                "duration": format_duration(total_duration),
                "source": source_label,
                "expanded": True,
                "filled": True,
                "tests": tests,
            }
        )
        assembly_name = (
            "Sinopac.DawhoEnterprise.Test.UnitTesting.EnterpriseAPI.EnterpriseApiUnit"
            if source_label == "UnitTest"
            else "Sinopac.DawhoEnterprise.Test.IntegrationTesting.EnterpriseAPI.EnterpriseApiIntegration"
        )
        rows.append(
            {
                "kind": "assembly",
                "indent": 1,
                "label": f"{assembly_name} ({len(tests)})",
                "duration": format_duration(total_duration),
                "source": "",
                "expanded": True,
                "filled": True,
                "tests": tests,
            }
        )
        for class_name, grouped_tests in class_groups:
            class_duration = sum(float(test.get("durationSeconds", 0) or 0) for test in grouped_tests)
            class_row = {
                "kind": "class",
                "indent": 2,
                "label": f"{class_name} ({len(grouped_tests)})",
                "duration": format_duration(class_duration),
                "source": "",
                "expanded": True,
                "selected": False,
                "filled": True,
                "tests": grouped_tests,
            }
            rows.append(class_row)
            for test in grouped_tests:
                rows.append(
                    {
                        "kind": "test",
                        "indent": 3,
                        "label": normalize_text(test.get("testName", "")).split(".")[-1],
                        "fullName": normalize_text(test.get("testName", "")),
                        "duration": format_duration(float(test.get("durationSeconds", 0) or 0)),
                        "source": "",
                        "expanded": False,
                        "selected": False,
                        "filled": True,
                        "tests": [test],
                    }
                )

    append_group("EnterpriseAPIIntegration", summary.integration_tests, "IntegrationTest", grouped_integration, "integration")
    append_group("EnterpriseAPIUnit", summary.unit_tests, "UnitTest", grouped_unit, "unit")

    selected_row = next(
        (row for row in rows if row.get("kind") == "test" and has_failed_tests(list(row.get("tests") or []))),
        None,
    ) or next((row for row in rows if row.get("kind") == "test"), None)
    if selected_row is not None:
        selected_row["selected"] = True
    return rows, selected_row


def has_failed_tests(tests: list[dict[str, Any]]) -> bool:
    return any(normalize_text(test.get("status")).casefold() == "failed" for test in tests)


def has_passed_tests(tests: list[dict[str, Any]]) -> bool:
    return any(normalize_text(test.get("status")).casefold() == "passed" for test in tests)


def render_test_explorer_html(summary: ApiVisualSummary, width: int, height: int) -> str:
    rows, _ = build_test_explorer_rows(summary)
    counts = summary.checklist_counts if (summary.checklist_counts or {}).get("total") else summary.counts
    solid_check_uri = svg_data_uri(SOLID_CHECK_SVG)
    hollow_check_uri = svg_data_uri(HOLLOW_CHECK_SVG)
    blue_diamond_uri = svg_data_uri(BLUE_DIAMOND_EXCLAMATION_OUTLINE_SVG)
    red_close_filled_uri = svg_data_uri(RED_CLOSE_FILLED_SVG)
    red_close_outline_uri = svg_data_uri(RED_CLOSE_OUTLINE_SVG)
    total_runtime = sum(float(test.get("durationSeconds", 0) or 0) for test in summary.all_tests)
    status_text = (
        f"測試回合已完成: 在 {format_status_runtime(total_runtime)} 完成驗證，"
        f"通過 {counts['passed']} 項 / 不通過 {counts['failed']} 項 / "
        f"不適用 {counts.get('notApplicable', 0)} 項"
    )

    row_html = []
    for row in rows:
        classes = ["test-row", f"kind-{row['kind']}"]
        if row.get("selected"):
            classes.append("selected")
        if row.get("filled"):
            classes.append("filled")
        indent = int(row["indent"])
        caret_class = "expanded" if row.get("expanded") else "collapsed"
        if has_failed_tests(list(row.get("tests") or [])):
            icon_class = "red-solid" if row.get("filled") else "red-outline"
        elif has_passed_tests(list(row.get("tests") or [])):
            icon_class = "solid" if row.get("filled") else "hollow"
        else:
            icon_class = "blue"
        row_html.append(
            f"""
            <div class="{' '.join(classes)}">
              <div class="test-cell" style="--indent:{indent}">
                <span class="caret {caret_class}"></span>
                <span class="status-icon {icon_class}"></span>
                <span class="label">{html.escape(str(row['label']))}</span>
              </div>
              <div class="duration-cell">{html.escape(str(row['duration']))}</div>
              <div class="feature-cell">{html.escape(str(row.get('source') or ''))}</div>
              <div class="error-cell"></div>
            </div>
            """
        )

    return f"""<!doctype html>
<html lang="zh-Hant">
<head>
<meta charset="utf-8" />
<style>
* {{ box-sizing: border-box; }}
html, body {{
  margin: 0;
  width: {width}px;
  height: {height}px;
  overflow: hidden;
  background: #1f1f1f;
  color: #dcdcdc;
  font-family: "微軟正黑體", "Microsoft JhengHei UI", "Segoe UI", sans-serif;
  font-size: 12px;
}}
.window {{
  width: {width}px;
  height: {height}px;
  background: #1f1f1f;
  overflow: hidden;
}}
.run-status {{
  height: 26px;
  display: flex;
  align-items: center;
  padding: 0 6px;
  background: #1f1f1f;
  border-bottom: 1px solid #333333;
  color: #f5f5f5;
  font-weight: 600;
}}
.main {{ height: calc(100% - 26px); background: #1f1f1f; }}
.grid-head {{
  height: 24px; display: grid; grid-template-columns: 1fr 80px 80px 112px;
  border-bottom: 1px solid #333333; border-top: 1px solid #333333; color:#f5f5f5;
  background: #1f1f1f;
}}
.grid-head div {{ padding: 4px 6px; border-right: 1px solid #2d2d2d; }}
.test-row {{
  height: 20px; display: grid; grid-template-columns: 1fr 80px 80px 112px;
  color: #b9b9b9;
}}
.test-row.selected {{ background: #3a3a3a; color: #ffffff; box-shadow: inset 3px 0 0 #9179ff, inset 0 1px 0 #4d4d4d, inset 0 -1px 0 #262626; outline: 1px dotted #4a4a4a; outline-offset: -2px; }}
.test-row.filled, .test-row.selected {{ font-weight: 600; }}
.test-row > div {{ overflow: hidden; white-space: nowrap; text-overflow: ellipsis; }}
.test-cell {{ padding-left: calc(6px + var(--indent) * 22px); display: flex; align-items: center; gap: 7px; }}
.duration-cell, .feature-cell, .error-cell {{ padding: 2px 6px; border-left: 1px solid #292929; color: #c7c7c7; }}
.caret {{ width: 0; height: 0; display: inline-block; flex: 0 0 auto; }}
.caret.expanded {{ border-left: 5px solid transparent; border-right: 5px solid transparent; border-top: 6px solid #d0d0d0; }}
.caret.collapsed {{ border-top: 5px solid transparent; border-bottom: 5px solid transparent; border-left: 6px solid #d0d0d0; }}
.status-icon {{
  position: relative;
  width: 14px;
  height: 14px;
  flex: 0 0 auto;
  background-repeat: no-repeat;
  background-position: center;
  background-size: contain;
}}
.status-icon.solid {{
  background-image: url("{solid_check_uri}");
}}
.status-icon.hollow {{
  background-image: url("{hollow_check_uri}");
}}
.status-icon.red-solid {{
  background-image: url("{red_close_filled_uri}");
}}
.status-icon.red-outline {{
  background-image: url("{red_close_outline_uri}");
}}
.status-icon.blue {{
  width: 14px;
  height: 14px;
  background-image: url("{blue_diamond_uri}");
}}
.kind-test .caret {{ visibility: hidden; }}
</style>
</head>
<body>
<div class="window">
  <div class="run-status">{html.escape(status_text)}</div>
  <div class="main">
    <div class="grid-head"><div>測試</div><div>持續時間</div><div>特性</div><div>錯誤訊息</div></div>
    {''.join(row_html)}
  </div>
</div>
</body>
</html>
"""


def render_html_screenshot(html_text: str, html_path: Path, output_path: Path, width: int, height: int) -> None:
    try:
        from playwright.sync_api import sync_playwright
    except Exception as exc:
        raise SystemExit("Python Playwright is required to render HTML evidence screenshots.") from exc

    html_path.parent.mkdir(parents=True, exist_ok=True)
    html_path.write_text(html_text, encoding="utf-8")
    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(headless=True)
        page = browser.new_page(viewport={"width": width, "height": height}, device_scale_factor=1)
        page.goto(html_path.resolve().as_uri(), wait_until="networkidle")
        page.screenshot(path=str(output_path), full_page=False)
        browser.close()


def draw_test_explorer_image(summary: ApiVisualSummary, output_path: Path) -> int:
    rows, _ = build_test_explorer_rows(summary)
    longest = max((css_px_for_label(str(row["label"])) + int(row["indent"]) * 22 for row in rows), default=700)
    width = max(860, min(1180, longest + 360))
    content_rows = max(len(rows), 8)
    height = 26 + 24 + content_rows * 20 + 16
    html_path = output_path.with_suffix(".html")
    html_text = render_test_explorer_html(summary, width, height)
    render_html_screenshot(html_text, html_path, output_path, width, height)
    return width


def set_cell_text(cell: Any, text: str, font_size_pt: int | None = None) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    set_paragraph_after_auto(paragraph)
    run = paragraph.add_run(text)
    set_run_report_font(run)
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)


def align_cell_paragraphs(cell: Any, alignment: Any) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment


def set_table_font_size(table: Any, font_size_pt: int) -> None:
    for row in table.rows:
        for cell in unique_cells(row):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text:
                        set_run_report_font(run)
                        run.font.size = Pt(font_size_pt)


def set_cell_picture(cell: Any, image_path: str, width_inches: float) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_after_auto(paragraph)
    paragraph.add_run().add_picture(image_path, width=Inches(width_inches))


def replace_text_in_paragraph(paragraph: Any, replacements: dict[str, str]) -> bool:
    original_text = paragraph.text
    updated_text = original_text
    for token, value in replacements.items():
        updated_text = updated_text.replace(token, value)
    if updated_text == original_text:
        return False

    if len(paragraph.runs) == 1:
        paragraph.runs[0].text = updated_text
        set_run_report_font(paragraph.runs[0])
        return True

    paragraph.text = ""
    set_paragraph_after_auto(paragraph)
    run = paragraph.add_run(updated_text)
    set_run_report_font(run)
    return True


def replace_document_placeholders(document: Document, replacements: dict[str, str]) -> int:
    updated = 0
    parts: list[Any] = [document]
    for section in document.sections:
        parts.extend([section.header, section.footer])

    for part in parts:
        for paragraph in part.paragraphs:
            if replace_text_in_paragraph(paragraph, replacements):
                updated += 1
        for table in part.tables:
            for row in table.rows:
                for cell in unique_cells(row):
                    for paragraph in cell.paragraphs:
                        if replace_text_in_paragraph(paragraph, replacements):
                            updated += 1
    return updated


def paragraph_has_border(paragraph: Any) -> bool:
    p_pr = paragraph._p.pPr
    return p_pr is not None and p_pr.find(qn("w:pBdr")) is not None


def remove_paragraph_border(paragraph: Any) -> bool:
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        return False
    p_pr.remove(p_bdr)
    return True


def add_first_page_page_border(section: Any) -> bool:
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    sect_pr = section._sectPr
    existing_borders = sect_pr.find(qn("w:pgBorders"))
    if existing_borders is not None:
        sect_pr.remove(existing_borders)

    page_borders = OxmlElement("w:pgBorders")
    page_borders.set(qn("w:offsetFrom"), "text")
    page_borders.set(qn("w:display"), "firstPage")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "double")
        border.set(qn("w:sz"), "20")
        border.set(qn("w:space"), "12")
        border.set(qn("w:color"), "000000")
        page_borders.append(border)
    sect_pr.append(page_borders)
    return True


def trim_overflowing_cover_borders(document: Document) -> int:
    """Replace paragraph-drawn cover frames with a first-page page border."""
    updated = 0
    for paragraph in document.paragraphs:
        if paragraph_has_border(paragraph) and remove_paragraph_border(paragraph):
            updated += 1
    if document.sections:
        add_first_page_page_border(document.sections[0])
        updated += 1
    return updated


def prepare_compact_v1_cover_border(document: Document) -> int:
    updated = 0
    for paragraph in document.paragraphs:
        if paragraph_has_border(paragraph):
            if remove_paragraph_border(paragraph):
                updated += 1

    body = document._element.body
    final_sect_pr = document.sections[-1]._sectPr
    existing_final_borders = final_sect_pr.find(qn("w:pgBorders"))
    if existing_final_borders is not None:
        final_sect_pr.remove(existing_final_borders)

    first_table_index = next(
        (index for index, child in enumerate(body) if child.tag == qn("w:tbl")),
        None,
    )
    if first_table_index is None:
        return updated

    section_break_paragraph = None
    for index in range(first_table_index - 1, -1, -1):
        if body[index].tag == qn("w:p"):
            section_break_paragraph = body[index]
            break
    if section_break_paragraph is None:
        return updated

    p_pr = section_break_paragraph.get_or_add_pPr()
    existing_section = p_pr.find(qn("w:sectPr"))
    if existing_section is not None:
        p_pr.remove(existing_section)

    first_section = OxmlElement("w:sectPr")
    section_type = OxmlElement("w:type")
    section_type.set(qn("w:val"), "nextPage")
    first_section.append(section_type)

    for child in final_sect_pr:
        child_name = child.tag.rsplit("}", 1)[-1]
        if child_name in {"pgSz", "pgMar", "cols", "docGrid"}:
            first_section.append(copy.deepcopy(child))

    page_borders = OxmlElement("w:pgBorders")
    page_borders.set(qn("w:offsetFrom"), "text")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "double")
        border.set(qn("w:sz"), "20")
        border.set(qn("w:space"), "12")
        border.set(qn("w:color"), "000000")
        page_borders.append(border)
    first_section.append(page_borders)
    p_pr.append(first_section)
    return updated + 1


def infer_function_display_name(function_code: str, execution_state: dict[str, Any]) -> str:
    spec_path = normalize_text(execution_state.get("specDocxPath"))
    if spec_path:
        stem = Path(spec_path).stem
        pattern = rf"^TSD\.{re.escape(function_code)}_(.+?)(?:_v\d|\s+v\d|$)"
        match = re.match(pattern, stem, flags=re.IGNORECASE)
        if match:
            return f"{function_code} {match.group(1).strip('_ ')}"
    return f"{function_code} 功能模組"


def load_feature_tester_map(path: Path = FEATURE_TESTER_MAP_JSON) -> dict[str, str]:
    project_path = resolve_asset_path("featureTesterMap", fallback=path)
    if project_path is not None:
        path = project_path
    if not path.exists():
        return {}
    payload = load_json(path)
    mapping = payload.get("mapping")
    if not isinstance(mapping, dict):
        return {}
    return {
        normalize_text(feature_id).casefold(): normalize_text(tester_name)
        for feature_id, tester_name in mapping.items()
        if normalize_text(feature_id) and normalize_text(tester_name)
    }


def resolve_feature_tester_name(function_code: str) -> str:
    tester_name = load_feature_tester_map().get(normalize_text(function_code).casefold())
    return tester_name or resolve_current_login_display_name()


def unique_cells(row: Any) -> list[Any]:
    result: list[Any] = []
    seen: set[int] = set()
    for cell in row.cells:
        identity = id(cell._tc)
        if identity in seen:
            continue
        seen.add(identity)
        result.append(cell)
    return result


def set_value_after_label(table: Any, label: str, value: str) -> bool:
    return set_value_after_label_occurrence(table, label, value, occurrence=1)


def update_standalone_date_paragraphs(document: Document, value: str) -> int:
    date_pattern = re.compile(r"^\s*\d{4}/\d{1,2}/\d{1,2}\s*$")
    updated = 0
    for paragraph in document.paragraphs:
        if not date_pattern.fullmatch(paragraph.text or ""):
            continue
        paragraph.text = ""
        set_paragraph_after_auto(paragraph)
        paragraph.add_run(value)
        updated += 1
    return updated


def set_value_after_label_occurrence(table: Any, label: str, value: str, occurrence: int = 1) -> bool:
    seen = 0
    for row in table.rows:
        cells = unique_cells(row)
        for index, cell in enumerate(cells):
            if normalize_text(cell.text) == normalize_text(label):
                seen += 1
                if seen != occurrence:
                    continue
                if index + 1 < len(cells):
                    set_cell_text(cells[index + 1], value)
                    return True
    return False


def replace_row_with_api_entries(document: Document, table_index: int, api_summaries: list[ApiVisualSummary]) -> None:
    table = document.tables[table_index]
    if len(table.rows) < 2:
        raise SystemExit("beautified template evidence table must contain a header row and one sample row")

    template_row = table.rows[1]
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    for summary in api_summaries:
        new_row_xml = copy.deepcopy(template_row._tr)
        table._tbl.append(new_row_xml)
        row = table.rows[-1]
        cells = unique_cells(row)
        if len(cells) < 2:
            raise SystemExit("beautified template evidence row must contain at least two cells")

        output_cell = cells[0]
        api_cell = cells[1]
        result_cell = cells[2] if len(cells) >= 3 else cells[1]
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        set_cell_picture(output_cell, summary.image_path, summary.image_display_width_inches)

        counts = summary.counts
        api_lines = [
            summary.api_display_name,
        ]
        result_lines = [
            f"關聯測試：{counts['total']} 項",
            format_pass_fail_text(counts),
        ]
        if len(cells) >= 3:
            set_cell_text(api_cell, "\n".join(api_lines))
            set_cell_text(result_cell, "\n".join(result_lines))
        else:
            set_cell_text(api_cell, "\n".join([*api_lines, *result_lines]))


def is_categorized_ut_checklist_table(table: Any) -> bool:
    if not table.rows:
        return False
    headers = [normalize_text(cell.text) for cell in unique_cells(table.rows[0])]
    return headers[:3] == ["編號", "測試內容", "測試結果"]


def find_categorized_ut_checklist_table_index(document: Document) -> int | None:
    for index, table in enumerate(document.tables):
        if is_categorized_ut_checklist_table(table):
            return index
    return None


def find_api_detail_table_index(document: Document) -> int | None:
    for index, table in enumerate(document.tables):
        if not table.rows:
            continue
        first_cell_text = normalize_text(unique_cells(table.rows[0])[0].text)
        if (
            first_cell_text == "API 單元測試執行明細"
            or first_cell_text.startswith("測試 API：")
            or first_cell_text.startswith("測試結果：")
        ):
            return index
    return None


def paragraph_text_from_xml(paragraph_xml: Any) -> str:
    return normalize_text("".join(node.text or "" for node in paragraph_xml.iter(qn("w:t"))))


def row_text_from_xml(row_xml: Any) -> str:
    return normalize_text(" ".join(node.text or "" for node in row_xml.iter(qn("w:t"))))


def find_previous_paragraph_xml(table: Any, text: str) -> Any | None:
    node = table._tbl.getprevious()
    while node is not None:
        if node.tag == qn("w:p") and paragraph_text_from_xml(node) == normalize_text(text):
            return node
        if node.tag == qn("w:tbl"):
            return None
        node = node.getprevious()
    return None


def is_category_row(row: Any) -> bool:
    cells = unique_cells(row)
    if len(cells) == 1:
        return normalize_text(cells[0].text).startswith("UT-")
    values = [normalize_text(cell.text) for cell in cells]
    return bool(values and values[0].startswith("UT-") and all(value == values[0] for value in values))


def get_style_id(document: Document, style_name: str) -> str | None:
    try:
        return document.styles[style_name].style_id
    except KeyError:
        return None


def text_paragraph_xml(document: Document, text: str, style_name: str | None = None) -> Any:
    paragraph = OxmlElement("w:p")
    style_id = get_style_id(document, style_name) if style_name else None
    p_pr = None
    if style_id:
        p_pr = OxmlElement("w:pPr")
        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), style_id)
        p_pr.append(p_style)
        paragraph.append(p_pr)
    if style_name in {"Heading 1", "Heading 2"}:
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            paragraph.insert(0, p_pr)
        num_pr = OxmlElement("w:numPr")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "0")
        num_pr.append(num_id)
        p_pr.append(num_pr)

    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    paragraph.append(run)
    return paragraph


def replace_paragraph_xml_text(paragraph_xml: Any, text: str) -> None:
    for node in list(paragraph_xml.iter(qn("w:r"))):
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    paragraph_xml.append(run)


def chinese_ordinal(index: int) -> str:
    tokens = [
        "",
        "一",
        "二",
        "三",
        "四",
        "五",
        "六",
        "七",
        "八",
        "九",
        "十",
        "十一",
        "十二",
        "十三",
        "十四",
        "十五",
        "十六",
        "十七",
        "十八",
        "十九",
        "二十",
    ]
    if 0 < index < len(tokens):
        return tokens[index]
    return str(index)


@dataclass
class ChecklistCategoryTemplate:
    title: str
    table_xml: Any
    row_xmls: list[Any]
    heading_xml: Any | None = None


def checklist_category_from_single_table(table: Any) -> list[ChecklistCategoryTemplate]:
    categories: list[ChecklistCategoryTemplate] = []
    current_title = ""
    current_rows: list[Any] = []
    for row in table.rows[1:]:
        if is_category_row(row):
            if current_title:
                categories.append(
                    ChecklistCategoryTemplate(
                        title=current_title,
                        table_xml=copy.deepcopy(table._tbl),
                        row_xmls=current_rows,
                    )
                )
            current_title = normalize_text(unique_cells(row)[0].text)
            current_rows = []
            continue
        current_rows.append(copy.deepcopy(row._tr))

    if current_title:
        categories.append(
            ChecklistCategoryTemplate(
                title=current_title,
                table_xml=copy.deepcopy(table._tbl),
                row_xmls=current_rows,
            )
        )
    return categories


def previous_ut_heading_xml(table: Any) -> Any | None:
    node = table._tbl.getprevious()
    while node is not None:
        if node.tag == qn("w:p"):
            text = paragraph_text_from_xml(node)
            if text.startswith("UT-"):
                return node
            if text:
                return None
        if node.tag == qn("w:tbl"):
            return None
        node = node.getprevious()
    return None


def collect_checklist_category_templates(
    document: Document,
    checklist_table_index: int,
) -> list[ChecklistCategoryTemplate]:
    checklist_template = document.tables[checklist_table_index]
    if any(is_category_row(row) for row in checklist_template.rows[1:]):
        return checklist_category_from_single_table(checklist_template)

    categories: list[ChecklistCategoryTemplate] = []
    for table in document.tables[checklist_table_index:]:
        if not is_categorized_ut_checklist_table(table):
            break
        heading_xml = previous_ut_heading_xml(table)
        title = paragraph_text_from_xml(heading_xml) if heading_xml is not None else ""
        if not title.startswith("UT-"):
            break
        categories.append(
            ChecklistCategoryTemplate(
                title=title,
                table_xml=copy.deepcopy(table._tbl),
                row_xmls=[copy.deepcopy(row._tr) for row in table.rows[1:]],
                heading_xml=heading_xml,
            )
        )

    return categories or checklist_category_from_single_table(checklist_template)


def empty_table_from_template(table_xml: Any) -> Any:
    table_copy = copy.deepcopy(table_xml)
    rows = list(table_copy.findall(qn("w:tr")))
    for row in rows[1:]:
        table_copy.remove(row)
    return table_copy


REPORT_TO_CLASSIFICATION_SECTION = {
    "UT-02": "UT-01",
    "UT-03": "UT-02",
    "UT-05": "UT-03",
    "UT-08": "UT-04",
    "UT-09": "UT-05",
    "UT-10": "UT-06",
    "UT-11": "UT-07",
    "UT-12": "UT-08",
    "UT-13": "UT-09",
    "UT-15": "UT-10",
}


def infer_not_applicable_from_category(category_title: str, summary: ApiVisualSummary) -> bool:
    category_text = normalize_text(category_title)
    api_text = f"{summary.api_id} {summary.api_name} {summary.api_display_name}".casefold()
    if "DB / SQL 執行環境驗證" in category_text and not summary.uses_sql:
        return True
    if "檔案" in category_text or "上傳" in category_text:
        return not any(token in api_text for token in ("upload", "file", "檔案", "上傳"))
    if "匯出" in category_text or "列印" in category_text:
        return not any(token in api_text for token in ("export", "download", "print", "匯出", "下載", "列印"))
    if "Mail" in category_text or "通知" in category_text:
        return not any(token in api_text for token in ("mail", "email", "notify", "notice", "message", "通知", "訊息"))
    return False


def infer_result_from_row_text(summary: ApiVisualSummary, category_title: str, row_text: str) -> str | None:
    normalized_row_text = normalize_text(row_text)
    row_text_casefold = normalized_row_text.casefold()
    if (
        "上線前環境驗證" in normalized_row_text
        or "Manual release checklist" in normalized_row_text
        or "UAT / 準生產" in normalized_row_text
    ):
        return "不適用"

    db_sql_category = "DB / SQL 執行環境驗證" in normalize_text(category_title)
    db_sql_row = db_sql_category or any(
        token in row_text_casefold
        for token in (
            "service sql",
            "正式 sql",
            "sql syntax",
            "schema",
            "table",
            "column",
            "sqldbfactory",
            "sqlexecutor",
            "enterpriseapi",
            "配置連線",
            "設定連線",
        )
    )
    if db_sql_row and not summary.uses_sql:
        return "不適用"
    if db_sql_category and summary.uses_sql:
        if summary.has_enterprise_config_runtime_validation and summarize_visual_result(summary) == "通過":
            return "通過"
        if "Service SQL 可在 EnterpriseAPI 設定連線下實際執行" in normalized_row_text:
            return "不通過\n原因：本次尚未取得以 EnterpriseAPI 設定連線執行 Service SQL 的驗證證據；目前僅有 UnitTest、Controller IntegrationTest 或測試資料庫 fixture 證據。"
        if "DB schema、table、column 與 SQL 使用一致" in normalized_row_text:
            return "不通過\n原因：本次尚未取得 EnterpriseAPI 設定連線所指資料庫的 schema、table、column 實際檢核證據。"
        if "DB 權限滿足最小執行需求" in normalized_row_text:
            return "不通過\n原因：本次尚未以 EnterpriseAPI 設定連線的資料庫帳號驗證 SELECT / INSERT / UPDATE / DELETE 等最小執行權限。"
        return "不通過\n原因：本次尚未取得 EnterpriseAPI 設定連線下的 Service runtime validation 證據。"
    return None


def checklist_result_for_row(summary: ApiVisualSummary, category_title: str, row_index: int, row_text: str = "") -> str:
    row_text_result = infer_result_from_row_text(summary, category_title, row_text)
    if row_text_result:
        return row_text_result

    report_section_id = extract_ut_section_id(category_title)
    classification_section_id = REPORT_TO_CLASSIFICATION_SECTION.get(report_section_id, report_section_id)

    section_value = summary.checklist_section_applicability.get(classification_section_id, "")
    if section_value == "not_applicable" or infer_not_applicable_from_category(category_title, summary):
        return "不適用"

    item_values = summary.checklist_item_applicability.get(classification_section_id) or []
    if row_index < len(item_values) and item_values[row_index] == "not_applicable":
        return "不適用"

    return summarize_visual_result(summary, include_enterprise_config=False)


def visible_not_applicable_count(category_templates: list[ChecklistCategoryTemplate], summary: ApiVisualSummary) -> int:
    count = 0
    for category in category_templates:
        if category.title.startswith("UT-01 "):
            continue
        for row_index, row_xml in enumerate(category.row_xmls):
            if checklist_result_for_row(summary, category.title, row_index, row_text_from_xml(row_xml)) == "不適用":
                count += 1
    return count


def visible_checklist_result_counts(category_templates: list[ChecklistCategoryTemplate], summary: ApiVisualSummary) -> dict[str, int]:
    passed = 0
    failed = 0
    not_applicable = 0

    for category in category_templates:
        if category.title.startswith("UT-01 "):
            examples = summary.mock_examples or [
                "規格未提供可解析範例情境，需補齊 API Spec Excel / mockExamples 後再驗證"
            ]
            result = summarize_visual_result(summary, include_enterprise_config=False) if summary.mock_examples else "失敗"
            if result == "通過":
                passed += len(examples)
            elif result == "不適用":
                not_applicable += len(examples)
            else:
                failed += len(examples)
            continue

        for row_index, row_xml in enumerate(category.row_xmls):
            result = checklist_result_for_row(summary, category.title, row_index, row_text_from_xml(row_xml))
            if result == "通過":
                passed += 1
            elif result == "不適用":
                not_applicable += 1
            else:
                failed += 1

    total = passed + failed
    return {
        "total": total,
        "passed": passed,
        "failed": failed,
        "skipped": 0,
        "pending": 0,
        "notApplicable": not_applicable,
    }


def format_api_example_check_content(item_index: int, scenario: str, has_examples: bool) -> str:
    if has_examples:
        return (
            f"{item_index}. {scenario}\n"
            "驗證方式：UnitTest"
        )
    return (
        f"{item_index}. {scenario}\n"
        "驗證重點：需補齊 API Spec Excel / mockExamples 後再驗證\n"
        "驗證方式：UnitTest"
    )


def populate_checklist_category_table(
    table: Any,
    category: ChecklistCategoryTemplate,
    summary: ApiVisualSummary,
) -> None:
    template_row_xml = category.row_xmls[0] if category.row_xmls else None
    if category.title.startswith("UT-01 "):
        examples = summary.mock_examples
        if not examples:
            examples = ["規格未提供可解析範例情境，需補齊 API Spec Excel / mockExamples 後再驗證"]
        for item_index, scenario in enumerate(examples, start=1):
            row_xml = copy.deepcopy(template_row_xml) if template_row_xml is not None else copy.deepcopy(table.rows[0]._tr)
            table._tbl.append(row_xml)
            row = table.rows[-1]
            cells = unique_cells(row)
            if len(cells) < 3:
                continue
            set_cell_text(cells[0], f"UT-01-{item_index:02d}", font_size_pt=9)
            set_cell_text(
                cells[1],
                format_api_example_check_content(item_index, scenario, bool(summary.mock_examples)),
                font_size_pt=9,
            )
            set_cell_text(
                cells[2],
                summarize_visual_result(summary, include_enterprise_config=False) if summary.mock_examples else "不通過",
                font_size_pt=9,
            )
        set_table_font_size(table, 9)
        return

    for row_index, row_xml in enumerate(category.row_xmls):
        table._tbl.append(copy.deepcopy(row_xml))
        row = table.rows[-1]
        cells = unique_cells(row)
        if len(cells) >= 3 and not normalize_text(cells[2].text):
            set_cell_text(
                cells[2],
                checklist_result_for_row(summary, category.title, row_index, row_text_from_xml(row_xml)),
                font_size_pt=9,
            )
    set_table_font_size(table, 9)


def summarize_visual_result(summary: ApiVisualSummary, *, include_enterprise_config: bool = True) -> str:
    tests = summary.all_tests
    if not include_enterprise_config:
        tests = [test for test in tests if not is_enterprise_config_runtime_test(test)]
    counts = summarize_tests(tests)
    return "通過" if not_passed_count(counts) == 0 and counts["total"] else "不通過"


def replace_placeholder_with_picture(
    document: Document,
    placeholder: str,
    image_path: str,
    width_inches: float,
) -> bool:
    placeholder_text = normalize_text(placeholder)
    for table in document.tables:
        for row in table.rows:
            for cell in unique_cells(row):
                if placeholder_text in normalize_text(cell.text):
                    set_cell_picture(cell, image_path, width_inches)
                    return True
    for paragraph in document.paragraphs:
        if placeholder_text in normalize_text(paragraph.text):
            paragraph.text = ""
            set_paragraph_after_auto(paragraph)
            paragraph.add_run().add_picture(image_path, width=Inches(width_inches))
            return True
    return False


def update_revision_history(document: Document, current_date: str, tester_name: str) -> bool:
    for table in document.tables:
        if len(table.rows) < 2:
            continue
        header_cells = [normalize_text(cell.text) for cell in unique_cells(table.rows[0])]
        if len(header_cells) < 3 or header_cells[:3] != ["版本", "日期", "說明"]:
            continue
        header_row_cells = unique_cells(table.rows[0])
        for cell in header_row_cells[:2]:
            align_cell_paragraphs(cell, WD_ALIGN_PARAGRAPH.CENTER)
        value_cells = unique_cells(table.rows[1])
        if len(value_cells) < 3:
            continue
        set_cell_text(value_cells[0], "1.0")
        set_cell_text(value_cells[1], current_date)
        set_cell_text(value_cells[2], f"初版，由 {tester_name} 測試。")
        for cell in value_cells[:2]:
            align_cell_paragraphs(cell, WD_ALIGN_PARAGRAPH.CENTER)
        return True
    return False


def table_grid_total_width_twips(table: Any) -> int | None:
    total = 0
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is not None:
        for grid_col in tbl_grid.gridCol_lst:
            raw_width = grid_col.get(qn("w:w"))
            if raw_width:
                try:
                    total += int(raw_width)
                except ValueError:
                    return None
    return total or None


def set_single_column_table_width(table: Any, width_twips: int) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(width_twips))

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, tbl_grid)
    for grid_col in list(tbl_grid.gridCol_lst):
        tbl_grid.remove(grid_col)
    grid_col = OxmlElement("w:gridCol")
    grid_col.set(qn("w:w"), str(width_twips))
    tbl_grid.append(grid_col)

    for row in table.rows:
        for cell in unique_cells(row):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(width_twips))


def page_break_paragraph_xml() -> Any:
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    run.append(page_break)
    paragraph.append(run)
    return paragraph


def paragraph_has_page_break(paragraph_xml: Any) -> bool:
    return any(br.get(qn("w:type")) == "page" for br in paragraph_xml.iter(qn("w:br")))


def ensure_page_break_before_heading(document: Document, heading_text: str) -> bool:
    target_text = normalize_text(heading_text)
    for paragraph in document.paragraphs:
        if normalize_text(paragraph.text) != target_text:
            continue

        node = paragraph._p.getprevious()
        while node is not None:
            if node.tag == qn("w:p"):
                if paragraph_has_page_break(node):
                    return False
                if paragraph_text_from_xml(node):
                    break
            elif node.tag == qn("w:tbl"):
                break
            node = node.getprevious()

        paragraph._p.addprevious(page_break_paragraph_xml())
        return True
    return False


def clear_cell_borders(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = tc_borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            tc_borders.append(border)
        border.set(qn("w:val"), "nil")


def clear_row_borders(row: Any) -> None:
    for cell in unique_cells(row):
        clear_cell_borders(cell)


def remove_table_row(row: Any) -> None:
    tr = row._tr
    parent = tr.getparent()
    if parent is not None:
        parent.remove(tr)


def remove_trailing_summary_screenshot_row(table: Table) -> bool:
    if not table.rows:
        return False
    row = table.rows[-1]
    row_text = normalize_text(" ".join(cell.text for cell in unique_cells(row)))
    if row_text and "{UnitTest VS執行總截圖}" not in row_text and "{總執行截圖}" not in row_text:
        return False
    remove_table_row(row)
    return True


def replace_api_detail_entries(table: Any, api_summaries: list[ApiVisualSummary]) -> None:
    if len(table.rows) < 3:
        raise SystemExit("API detail table must contain title, summary placeholder, and screenshot placeholder rows")

    summary_template_row_xml = copy.deepcopy(table.rows[1]._tr)
    image_template_row_xml = copy.deepcopy(table.rows[2]._tr)
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    for index, summary in enumerate(api_summaries, 1):
        summary_row_xml = copy.deepcopy(summary_template_row_xml)
        table._tbl.append(summary_row_xml)
        summary_row = table.rows[-1]
        summary_cells = unique_cells(summary_row)
        if summary_cells:
            summary_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(
                summary_cells[0],
                f"{index}. {summary.api_display_name}　{format_pass_fail_text(summary.counts)}",
            )

        image_row_xml = copy.deepcopy(image_template_row_xml)
        table._tbl.append(image_row_xml)
        image_row = table.rows[-1]
        image_cells = unique_cells(image_row)
        if image_cells:
            image_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_picture(image_cells[0], summary.image_path, summary.image_display_width_inches)


def fill_api_section_table(table: Any, summary: ApiVisualSummary) -> None:
    if table.rows and normalize_text(unique_cells(table.rows[0])[0].text).startswith("測試 API："):
        table._tbl.remove(table.rows[0]._tr)
    if len(table.rows) < 2:
        raise SystemExit("API section table must contain result and screenshot rows")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = table.rows
    result_counts = summary.checklist_counts or summary.counts
    if summary.checklist_counts:
        result_text = format_report_result_text(result_counts)
    else:
        result_text = format_pass_fail_text(result_counts)
    set_cell_text(unique_cells(rows[0])[0], f"測試結果：{result_text}")
    set_cell_picture(unique_cells(rows[1])[0], summary.image_path, summary.image_display_width_inches)
    for row in table.rows:
        for cell in unique_cells(row):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def find_previous_api_heading_xml(table: Any) -> Any | None:
    node = table._tbl.getprevious()
    while node is not None:
        if node.tag == qn("w:p"):
            text = paragraph_text_from_xml(node)
            if "{API名稱}" in text or "API名稱" in text:
                return node
            if text:
                return None
        if node.tag == qn("w:tbl"):
            return None
        node = node.getprevious()
    return None


def remove_xml_node_once(nodes: list[Any], node: Any | None) -> None:
    if node is None:
        return
    if any(existing is node for existing in nodes):
        return
    nodes.append(node)


def replace_categorized_api_sections(
    document: Document,
    api_section_table_index: int,
    checklist_table_index: int,
    api_summaries: list[ApiVisualSummary],
) -> None:
    api_section_template = document.tables[api_section_table_index]
    checklist_template = document.tables[checklist_table_index]
    api_heading_template_xml = find_previous_api_heading_xml(api_section_template)
    checklist_heading_xml = find_previous_paragraph_xml(checklist_template, "測試內容清單")
    category_templates = collect_checklist_category_templates(document, checklist_table_index)
    checklist_table_nodes_to_remove: list[Any] = []
    for table in document.tables[checklist_table_index:]:
        if not is_categorized_ut_checklist_table(table):
            break
        checklist_table_nodes_to_remove.append(table._tbl)

    api_section_template_xml = copy.deepcopy(api_section_template._tbl)
    checklist_width_twips = table_grid_total_width_twips(checklist_template)

    for index, summary in enumerate(api_summaries, start=1):
        api_heading_xml = text_paragraph_xml(
            document,
            f"{chinese_ordinal(index + 1)}、{summary.api_display_name}",
            "Heading 1",
        )
        api_section_template._tbl.addprevious(api_heading_xml)

        api_section_xml = copy.deepcopy(api_section_template_xml)
        api_section_template._tbl.addprevious(api_section_xml)
        anchor = api_section_xml
        api_section_table = Table(api_section_xml, document._body)
        if checklist_width_twips:
            set_single_column_table_width(api_section_table, checklist_width_twips)
        fill_api_section_table(api_section_table, summary)

        for category in category_templates:
            category_heading_xml = text_paragraph_xml(document, category.title, "Heading 2")
            anchor.addnext(category_heading_xml)
            anchor = category_heading_xml

            checklist_xml = empty_table_from_template(category.table_xml)
            anchor.addnext(checklist_xml)
            anchor = checklist_xml
            checklist_table = Table(checklist_xml, document._body)
            if checklist_width_twips:
                checklist_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            populate_checklist_category_table(checklist_table, category, summary)

    nodes_to_remove: list[Any] = []
    remove_xml_node_once(nodes_to_remove, api_heading_template_xml)
    remove_xml_node_once(nodes_to_remove, api_section_template._tbl)
    remove_xml_node_once(nodes_to_remove, checklist_heading_xml)
    for category in category_templates:
        remove_xml_node_once(nodes_to_remove, category.heading_xml)
    for table_node in checklist_table_nodes_to_remove:
        remove_xml_node_once(nodes_to_remove, table_node)
    for node in nodes_to_remove:
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)


def replace_ut01_with_api_example_entries(table: Any, api_summaries: list[ApiVisualSummary]) -> None:
    if not api_summaries:
        return

    rows = table.rows
    ut01_index = None
    for index, row in enumerate(rows):
        cells = unique_cells(row)
        if cells and normalize_text(cells[0].text).startswith("UT-01 "):
            ut01_index = index
            break
    if ut01_index is None or ut01_index + 1 >= len(rows):
        return

    next_section_index = len(rows)
    for index in range(ut01_index + 1, len(rows)):
        if is_category_row(rows[index]):
            next_section_index = index
            break

    template_row_xml = copy.deepcopy(rows[ut01_index + 1]._tr)
    for index in range(next_section_index - 1, ut01_index, -1):
        table._tbl.remove(rows[index]._tr)

    anchor = table.rows[ut01_index]._tr
    item_index = 0
    for summary in api_summaries:
        examples = summary.mock_examples or ["規格未提供可解析範例情境，需補齊 API Spec Excel / mockExamples 後再驗證"]
        for scenario in examples:
            item_index += 1
            row_xml = copy.deepcopy(template_row_xml)
            anchor.addnext(row_xml)
            anchor = row_xml
            row = table.rows[ut01_index + item_index]
            cells = unique_cells(row)
            if len(cells) < 3:
                continue
            set_cell_text(cells[0], f"UT-01-{item_index:02d}", font_size_pt=9)
            set_cell_text(
                cells[1],
                format_api_example_check_content(
                    item_index,
                    f"{summary.api_display_name}／{scenario}",
                    bool(summary.mock_examples),
                ),
                font_size_pt=9,
            )
            set_cell_text(
                cells[2],
                summarize_visual_result(summary, include_enterprise_config=False) if summary.mock_examples else "不通過",
                font_size_pt=9,
            )


def build_module_visual_report(
    context_root: Path,
    template_docx: Path,
    output_docx: Path,
    assets_dir: Path,
) -> dict[str, Any]:
    repo_root = infer_repo_root(context_root)
    checklist = load_json(context_root / "api-checklist.json")
    execution_state = load_json(context_root / "execution-state.json")
    function_code = normalize_text(execution_state.get("functionCode")) or context_root.name
    function_display_name = infer_function_display_name(function_code, execution_state)

    api_summaries = [
        load_api_visual_summary(repo_root, context_root, item, assets_dir)
        for item in checklist.get("items") or []
    ]

    document = Document(str(template_docx))
    categorized_checklist_table_index = find_categorized_ut_checklist_table_index(document)
    if categorized_checklist_table_index is not None:
        info_table = None
        header_table = document.tables[1] if len(document.tables) > 1 else document.tables[0]
        evidence_table_index = categorized_checklist_table_index
        compact_template = False
        categorized_template = True
        category_templates_for_summary = collect_checklist_category_templates(document, categorized_checklist_table_index)
        for summary in api_summaries:
            summary.checklist_counts = visible_checklist_result_counts(category_templates_for_summary, summary)
            summary.not_applicable_count = summary.checklist_counts.get("notApplicable", 0)
    elif len(document.tables) >= 4:
        info_table = document.tables[1]
        header_table = document.tables[2]
        evidence_table_index = 3
        compact_template = False
        categorized_template = False
    elif len(document.tables) >= 2:
        info_table = None
        header_table = document.tables[0]
        evidence_table_index = 1
        compact_template = True
        categorized_template = False
    else:
        raise SystemExit("beautified template must contain a header table and an evidence table")

    for summary in api_summaries:
        image_width_px = draw_test_explorer_image(summary, Path(summary.image_path))
        summary.image_display_width_inches = recommend_docx_image_width_inches(image_width_px)

    overall_summary: ApiVisualSummary | None = None
    overall_image_path: str | None = None
    if not categorized_template:
        overall_summary = build_overall_visual_summary(function_code, function_display_name, api_summaries, assets_dir)
        overall_image_width_px = draw_test_explorer_image(overall_summary, Path(overall_summary.image_path))
        overall_summary.image_display_width_inches = recommend_docx_image_width_inches(overall_image_width_px)
        overall_image_path = overall_summary.image_path
    else:
        legacy_overall_image = assets_dir / f"{safe_file_name(function_code)}_UnitTest_Summary.png"
        if legacy_overall_image.exists():
            legacy_overall_image.unlink()

    current_date = today_slash()
    tester_name = resolve_feature_tester_name(function_code) if (compact_template or categorized_template) else resolve_current_login_display_name()
    update_standalone_date_paragraphs(document, current_date)
    replace_document_placeholders(
        document,
        {
            "{功能編號}_{功能名稱}": function_display_name,
            "{功能編號}": function_code,
            "{功能名稱}": "功能模組",
            "{Today:yyyy/MM/dd}": current_date,
            "{姓名}": tester_name,
            "UT自测报告模板": "UT自測報告",
            "UT自测报告": "UT自測報告",
            "UT自測報告模板": "UT自測報告",
        },
    )
    if compact_template:
        prepare_compact_v1_cover_border(document)
    else:
        trim_overflowing_cover_borders(document)
    if info_table is not None:
        set_value_after_label(info_table, "報告日期", current_date)
    update_revision_history(document, current_date, tester_name)
    set_value_after_label(header_table, "功能名稱", function_display_name)
    set_value_after_label(header_table, "API 名稱", function_display_name)
    set_value_after_label(header_table, "測試人員", tester_name)
    set_value_after_label(header_table, "測試日期", current_date)

    report_counts_by_api = [summary.checklist_counts or summary.counts for summary in api_summaries]
    total_tests = sum(counts.get("total", 0) for counts in report_counts_by_api)
    total_passed = sum(counts.get("passed", 0) for counts in report_counts_by_api)
    total_failed = sum(counts.get("failed", 0) for counts in report_counts_by_api)
    total_skipped = sum(counts.get("skipped", 0) for counts in report_counts_by_api)
    total_pending = sum(counts.get("pending", 0) for counts in report_counts_by_api)
    total_not_applicable = sum(counts.get("notApplicable", 0) for counts in report_counts_by_api)
    total_not_passed = max(0, total_tests - total_passed)
    all_passed = total_tests > 0 and total_not_passed == 0

    summary_text = (
        f"共 {len(api_summaries)} 支 API，關聯測試 {total_tests} 項，其中通過 {total_passed} 項、失敗 {total_failed} 項。"
    )
    conclusion_text = "☑ 符合需求 ☐ 不符合需求" if all_passed else "☐ 符合需求 ☑ 不符合需求"
    pass_fail_text = format_report_result_text(
        {
            "total": total_tests,
            "passed": total_passed,
            "failed": total_failed,
            "skipped": total_skipped,
            "pending": total_pending,
            "notApplicable": total_not_applicable,
        }
    )
    result_text = pass_fail_text if (compact_template or categorized_template) else summary_text
    set_value_after_label(header_table, "實際測試結果", result_text)
    if not set_value_after_label(header_table, "測試結論", conclusion_text) and compact_template:
        set_value_after_label_occurrence(header_table, "實際測試結果", conclusion_text, occurrence=2)
    if categorized_template:
        set_value_after_label_occurrence(header_table, "實際測試結果", conclusion_text, occurrence=2)
    if overall_summary is not None:
        replace_placeholder_with_picture(
            document,
            "{UnitTest VS執行總截圖}",
            overall_summary.image_path,
            overall_summary.image_display_width_inches,
        )
    if categorized_template:
        remove_trailing_summary_screenshot_row(header_table)
        ensure_page_break_before_heading(document, "彙總")

    if categorized_template:
        api_detail_table_index = find_api_detail_table_index(document)
        if api_detail_table_index is not None:
            replace_categorized_api_sections(document, api_detail_table_index, evidence_table_index, api_summaries)
        else:
            replace_ut01_with_api_example_entries(document.tables[evidence_table_index], api_summaries)
    else:
        replace_row_with_api_entries(document, evidence_table_index, api_summaries)

    apply_report_font_to_document(document)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))

    return {
        "functionCode": function_code,
        "contextRoot": context_root.as_posix(),
        "templateDocx": template_docx.as_posix(),
        "outputDocx": output_docx.as_posix(),
        "assetsDir": assets_dir.as_posix(),
        "apiCount": len(api_summaries),
        "overallImagePath": overall_image_path,
        "testSummary": {
            "total": total_tests,
            "passed": total_passed,
            "failed": total_failed,
            "skipped": total_skipped,
            "pending": total_pending,
            "notApplicable": total_not_applicable,
        },
        "apis": [
            {
                "apiId": summary.api_id,
                "apiDisplayName": summary.api_display_name,
                "mockExamples": summary.mock_examples,
                "notApplicable": summary.not_applicable_count,
                "imagePath": summary.image_path,
                "counts": summary.checklist_counts or summary.counts,
                "automationCounts": summary.counts,
                "tests": [test.get("testName", "") for test in summary.all_tests],
            }
            for summary in api_summaries
        ],
    }


def main() -> None:
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

    parser = argparse.ArgumentParser(
        description="Generate a module-level beautified UT report with one synthetic VS-style result image per API.",
    )
    parser.add_argument("--context-root", help="Path to .agent/context/<functionCode>.")
    parser.add_argument("--function-code", help="Function code used to resolve centralized .agent/context/<functionCode>.")
    parser.add_argument("--project-root", default=str(Path.cwd()), help="Current code branch/project root; used for centralized .agent resolution.")
    parser.add_argument("--agent-dir", default=".agent")
    parser.add_argument("--agent-root")
    parser.add_argument("--workspace-root")
    parser.add_argument("--workspace-key")
    parser.add_argument("--rules-root")
    parser.add_argument(
        "--template-docx",
        help="Optional path to the beautified module report template DOCX. Defaults to project-rules catalog asset utReportTemplate.",
    )
    parser.add_argument("--output-docx", help="Optional output DOCX path.")
    parser.add_argument("--assets-dir", help="Optional output directory for generated PNG evidence images.")
    parser.add_argument("--emit-summary-json", help="Optional summary JSON output path.")
    args = parser.parse_args()

    context_root = resolve_central_context_root(args)
    repo_root = infer_repo_root(context_root)
    template_docx = resolve_template_docx(
        args.template_docx,
        repo_root,
        rules_root_arg=args.rules_root,
        workspace_key=args.workspace_key,
    )
    output_docx = Path(args.output_docx).expanduser().resolve() if args.output_docx else default_output_docx(context_root, template_docx)
    assets_dir = Path(args.assets_dir).expanduser().resolve() if args.assets_dir else default_assets_dir(output_docx)

    payload = build_module_visual_report(context_root, template_docx, output_docx, assets_dir)
    update_chain_status(
        agent_root=context_root.parent.parent,
        function_code=context_root.name,
        stage="test",
        status="done",
        phase="report_generated",
        message="UT report generated",
        artifacts={
            "utReport": output_docx.as_posix(),
            "evidenceImages": assets_dir.as_posix(),
        },
    )
    if args.emit_summary_json:
        write_json(Path(args.emit_summary_json).expanduser().resolve(), payload)

    print(f"Module visual report written: {output_docx.as_posix()}")
    print(f"Evidence images: {assets_dir.as_posix()}")
    print(json.dumps(payload["testSummary"], ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
