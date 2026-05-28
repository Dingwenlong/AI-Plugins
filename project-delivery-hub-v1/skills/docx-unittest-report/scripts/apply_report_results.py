from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.table import Table
except ImportError as exc:
    raise SystemExit(
        "docx-unittest-report 缺少 Python 依赖：python-docx。"
        "请在当前解释器安装后重试，例如：python -m pip install python-docx"
    ) from exc

from docx_report_utils import (
    STATUS_LABELS,
    iter_block_items,
    is_header_table,
    load_json,
    load_report_outline,
    normalize_text,
    overall_status_from_summary,
    overall_status_label,
    result_summary,
    resolve_effective_metadata,
    strip_auto_status,
    write_json,
)


def set_cell_text(cell: Any, text: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.add_run(text)


def clear_paragraph(paragraph: Any) -> None:
    element = paragraph._p
    for child in list(element):
        if child.tag.endswith("}pPr"):
            continue
        element.remove(child)


def set_paragraph_text(paragraph: Any, text: str) -> None:
    clear_paragraph(paragraph)
    paragraph.add_run(text)


def total_cases(summary: dict[str, int]) -> int:
    return sum(
        int(summary.get(key, 0) or 0)
        for key in ("passed", "failed", "manual", "pending", "skipped", "not_in_run")
    )


def completed_cases(summary: dict[str, int]) -> int:
    return total_cases(summary) - int(summary.get("pending", 0) or 0) - int(summary.get("not_in_run", 0) or 0)


def render_summary_sentence(summary: dict[str, int]) -> str:
    total = total_cases(summary)
    completed = completed_cases(summary)
    segments = [
        f"共 {total} 項檢查",
        f"已完成 {completed} 項",
        (
            f"其中通過 {summary['passed']} 項、失敗 {summary['failed']} 項、"
            f"人工確認 {summary['manual']} 項、不適用 {summary['skipped']} 項"
        ),
    ]
    if summary.get("pending", 0):
        segments.append(f"另有 {summary['pending']} 項待補")
    if summary.get("not_in_run", 0):
        segments.append(f"另有 {summary['not_in_run']} 項接口未涉及")
    return "，".join(segments) + "。"


def automation_executed_cases(summary: dict[str, int]) -> int:
    return int(summary.get("passed", 0) or 0) + int(summary.get("failed", 0) or 0)


def render_automation_summary_sentence(summary: dict[str, int]) -> str:
    total = total_cases(summary)
    executed = automation_executed_cases(summary)
    passed = int(summary.get("passed", 0) or 0)
    failed = int(summary.get("failed", 0) or 0)
    return f"共 {total} 項檢查，已執行 {executed} 項，其中通過 {passed} 項、失敗 {failed} 項。"


def render_section_conclusion(summary: dict[str, int]) -> str:
    if summary.get("failed", 0):
        return "本節存在失敗項，需先修正問題並重新驗證。"
    if summary.get("pending", 0) and completed_cases(summary) == 0:
        return "本節尚未完成實質驗證，需先補齊測試綁定或人工回填。"
    if summary.get("pending", 0):
        return "本節已有部分驗證結果，但仍有待補項。"
    if summary.get("not_in_run", 0) and completed_cases(summary) == 0:
        return "本節項目接口未涉及。"
    if summary.get("not_in_run", 0):
        return "本節已完成部分驗證，其餘項目接口未涉及。"
    return "本節已完成回填，現有結果可作為交付依據。"


def source_kind_label(source_kind: str) -> str:
    if normalize_text(source_kind) == "integrationTest":
        return "IntegrationTest"
    if normalize_text(source_kind) == "unitTest":
        return "UnitTest"
    if normalize_text(source_kind) == "codeInspection":
        return "CodeInspection"
    if normalize_text(source_kind) == "apiRuntimeCall":
        return "Postman MCP / 真实接口调用"
    return "未指定"


def fallback_actual_result(case_result: dict[str, Any]) -> str:
    status = case_result.get("status", "pending")
    source_label = source_kind_label(case_result.get("sourceKind", ""))
    if status == "passed":
        return f"已由 {source_label} 驗證通過。"
    if status == "failed":
        return f"對應 {source_label} 驗證失敗，請參考下方失敗摘要。"
    if status == "manual":
        return "已完成人工確認。"
    if status == "skipped":
        return "此項目前不適用。"
    if status == "not_in_run":
        return "此項接口未涉及。"
    if case_result.get("missingTests"):
        return "已綁定測試，但尚未取得對應結果。"
    if case_result.get("boundTests"):
        return "對應測試尚未完成，待補執行結果。"
    if normalize_text(case_result.get("sourceKind", "")) == "codeInspection":
        return "尚未定位到充分代碼證據。"
    return "尚未綁定對應測試，待補驗證結果。"


def display_actual_result(case_result: dict[str, Any]) -> str:
    actual_result = normalize_text(case_result.get("actualResult", ""))
    status = case_result.get("status", "pending")
    if status == "manual" and actual_result:
        return actual_result
    generic_fragments = {
        "待補",
        "不適用",
        "人工確認",
        "未納入本次執行",
        "接口未涉及",
        "詳見 unittest 結果",
        "詳見 integrationtest 結果",
        "詳見unittest失敗摘要與附件",
        "詳見integrationtest失敗摘要與附件",
        "詳見 postman mcp / 真实接口调用結果",
    }
    normalized_casefold = actual_result.casefold()
    if actual_result and all(fragment not in normalized_casefold for fragment in generic_fragments):
        return actual_result
    return fallback_actual_result(case_result)


def update_header(document: Document, metadata: dict[str, Any], summary: dict[str, int]) -> None:
    table = next((candidate for candidate in document.tables if is_header_table(candidate)), document.tables[0])
    header_value = metadata.get("apiDisplayName") or ""
    tester = metadata.get("tester") or ""
    test_date = metadata.get("testDate") or ""

    if header_value and len(table.rows) > 0 and len(table.rows[0].cells) > 1:
        set_cell_text(table.rows[0].cells[1], header_value)
    if tester and len(table.rows) > 0 and len(table.rows[0].cells) > 4:
        set_cell_text(table.rows[0].cells[4], tester)
    if test_date and len(table.rows) > 0 and len(table.rows[0].cells) > 6:
        set_cell_text(table.rows[0].cells[6], test_date)

    actual_summary = metadata.get("actualSummary") or render_summary_sentence(summary)
    if len(table.rows) > 2 and len(table.rows[2].cells) > 2:
        set_cell_text(table.rows[2].cells[2], actual_summary)

    overall_status = metadata.get("overallStatus") or overall_status_from_summary(summary)
    if len(table.rows) > 3 and len(table.rows[3].cells) > 2:
        set_cell_text(table.rows[3].cells[2], overall_status_label(overall_status))


def update_cover_date(document: Document, test_date: str) -> bool:
    if not test_date:
        return False

    target_paragraph = None
    for block in iter_block_items(document):
        if isinstance(block, Table):
            break
        if normalize_text(block.text) and normalize_text(block.text).count("/") == 2:
            parts = normalize_text(block.text).split("/")
            if len(parts) == 3 and all(part.isdigit() for part in parts):
                target_paragraph = block

    if target_paragraph is None:
        return False

    set_paragraph_text(target_paragraph, test_date)
    return True


def status_for_case(case_result: dict[str, Any]) -> str:
    return STATUS_LABELS.get(case_result.get("status", "pending"), "待補")


def pending_case_result(item: dict[str, Any]) -> dict[str, Any]:
    return {
        "caseId": item["caseId"],
        "checkItem": item["checkItem"],
        "status": "pending",
        "actualResult": "待補",
        "sourceKind": (
            "codeInspection"
            if item.get("mode") == "code_inspection"
            else (
                "apiRuntimeCall"
                if item.get("mode") == "api_runtime_call"
                else ("integrationTest" if item.get("mode") == "integration_test" else "unitTest")
            )
        ),
        "attachmentPaths": [],
        "failureDetails": [],
        "missingTests": [],
        "boundTests": [],
        "trxPath": "",
    }


def update_checklist_row(table: Any, row_index: int, case_result: dict[str, Any]) -> None:
    row = table.rows[row_index]
    cells = row.cells
    status_label = status_for_case(case_result)
    actual_result = display_actual_result(case_result)

    if len(cells) >= 3:
        set_cell_text(cells[1], actual_result)
        return

    left_text = strip_auto_status(cells[0].text)
    rendered_status = status_label
    if (
        (
            case_result.get("status") == "manual"
            or normalize_text(case_result.get("sourceKind", "")) == "apiRuntimeCall"
        )
        and actual_result
        and normalize_text(actual_result) != normalize_text("已完成人工確認。")
    ):
        rendered_status = f"{status_label}\n{actual_result}"
    if left_text and "證據：" in left_text:
        set_cell_text(cells[0], f"{left_text}\n狀態：{rendered_status}")
    else:
        set_cell_text(cells[0], rendered_status)


def find_or_create_evidence_cell(table: Any) -> Any:
    evidence_rows = []
    for row in list(table.rows):
        marker = row.cells[0].text if row.cells else ""
        if marker.strip().startswith("自動化證據"):
            evidence_rows.append(row)

    if evidence_rows:
        keep_row = evidence_rows[0]
        for extra_row in evidence_rows[1:]:
            extra_row._tr.getparent().remove(extra_row._tr)
        return keep_row.cells[0]

    row = table.add_row()
    if len(row.cells) > 1:
        return row.cells[0].merge(row.cells[-1])
    return row.cells[0]


def remove_extra_paragraphs(cell: Any) -> None:
    for paragraph in list(cell.paragraphs)[1:]:
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def fill_evidence_cell(
    cell: Any,
    section: dict[str, Any],
    results_lookup: dict[str, dict[str, Any]],
) -> None:
    section_cases = [
        results_lookup.get(item["caseId"], pending_case_result(item))
        for item in section["items"]
    ]
    set_cell_text(cell, render_automation_summary_sentence(result_summary(section_cases)))
    remove_extra_paragraphs(cell)


def main() -> None:
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

    parser = argparse.ArgumentParser(
        description="Write UnitTest/TRX results back into a DOCX test report.",
    )
    parser.add_argument("manifest_path", help="Path to the JSON manifest.")
    parser.add_argument("results_path", help="Path to the collected results JSON.")
    parser.add_argument(
        "--output-docx",
        help="Optional override for the revised DOCX output path.",
    )
    parser.add_argument(
        "--emit-summary-json",
        help="Optional path to write the summary JSON that was applied to the DOCX.",
    )
    args = parser.parse_args()

    manifest = load_json(args.manifest_path)
    results = load_json(args.results_path)
    input_docx = manifest["document"]["inputPath"]
    output_docx = args.output_docx or manifest["document"].get("outputPath") or input_docx

    document = Document(input_docx)
    outline = load_report_outline(input_docx)
    metadata = resolve_effective_metadata(manifest.get("metadata", {}), outline.get("header", {}))
    results_lookup = {case["caseId"]: case for case in results.get("cases", [])}
    all_cases = [
        results_lookup.get(item["caseId"], pending_case_result(item))
        for section in outline["sections"]
        for item in section["items"]
    ]

    summary = result_summary(all_cases)
    update_header(document, metadata, summary)
    update_cover_date(document, metadata.get("testDate", ""))

    for section in outline["sections"]:
        table = document.tables[section["tableIndex"]]
        for item in section["items"]:
            case_result = results_lookup.get(item["caseId"], pending_case_result(item))
            update_checklist_row(table, item["rowIndex"], case_result)
        evidence_cell = find_or_create_evidence_cell(table)
        fill_evidence_cell(
            evidence_cell,
            section,
            results_lookup=results_lookup,
        )

    output_path = Path(output_docx).expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))

    applied_summary = {
        "manifestPath": Path(args.manifest_path).resolve().as_posix(),
        "resultsPath": Path(args.results_path).resolve().as_posix(),
        "outputDocx": output_path.as_posix(),
        "summary": summary,
    }
    if args.emit_summary_json:
        write_json(args.emit_summary_json, applied_summary)

    print(f"Updated DOCX written: {output_path.as_posix()}")


if __name__ == "__main__":
    main()
