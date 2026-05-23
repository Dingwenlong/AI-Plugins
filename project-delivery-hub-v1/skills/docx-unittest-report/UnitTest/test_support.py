from __future__ import annotations

import json
import sys
import tempfile
from pathlib import Path
from typing import Any

from docx import Document

SKILL_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = SKILL_ROOT / "scripts"
ASSETS_DIR = SKILL_ROOT / "assets"

if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def create_sample_docx(path: Path) -> Path:
    document = Document()
    document.add_paragraph("後台API 自測報告")
    document.add_paragraph("2026/04/13")

    header = document.add_table(rows=4, cols=7)
    header.rows[0].cells[0].text = "API 名稱"
    header.rows[0].cells[1].text = "Setting/UpdateUserAlias"
    header.rows[0].cells[3].text = "測試人員"
    header.rows[0].cells[4].text = "Codex"
    header.rows[0].cells[5].text = "測試日期"
    header.rows[0].cells[6].text = "yyyy/mm/dd"
    header.rows[2].cells[1].text = "實際測試結果"
    header.rows[2].cells[2].text = ""
    header.rows[3].cells[1].text = "實際測試結果"
    header.rows[3].cells[2].text = ""

    document.add_paragraph("UT-01 API 介面規格與一致性")
    checklist = document.add_table(rows=3, cols=2)
    checklist.cell(0, 0).merge(checklist.cell(0, 1)).text = "UT-01 API 介面規格與一致性"
    checklist.rows[1].cells[0].text = ""
    checklist.rows[1].cells[1].text = "HTTP 狀態碼使用正確(依規格)"
    checklist.rows[2].cells[0].text = ""
    checklist.rows[2].cells[1].text = "DB 資料已正確更新"

    document.add_paragraph("UT-10 需求/規格對照(逐條驗證)")
    matrix = document.add_table(rows=3, cols=3)
    matrix.cell(0, 0).merge(matrix.cell(0, 2)).text = "UT-10 需求/規格對照(逐條驗證)"
    matrix.rows[1].cells[0].text = "預期輸出/行為"
    matrix.rows[1].cells[1].text = "實際輸出/行為"
    matrix.rows[1].cells[2].text = "備註"
    matrix.rows[2].cells[0].text = "回應結構符合規格"
    matrix.rows[2].cells[1].text = ""
    matrix.rows[2].cells[2].text = ""

    document.save(str(path))
    return path


def build_temp_manifest(docx_path: Path, output_docx: Path, trx_path: Path | None = None) -> Path:
    from docx_report_utils import build_manifest_from_outline, load_report_outline

    outline = load_report_outline(docx_path)
    manifest = build_manifest_from_outline(outline, output_docx=output_docx.as_posix())
    if trx_path is not None:
        manifest["unitTest"]["trxPath"] = trx_path.as_posix()
    manifest_path = docx_path.with_suffix(".job.json")
    write_json(manifest_path, manifest)
    return manifest_path


def temp_dir() -> Path:
    return Path(tempfile.mkdtemp(prefix="docx-unittest-"))
