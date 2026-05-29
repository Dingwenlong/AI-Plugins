from __future__ import annotations

import json
import subprocess
import sys
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

UNIT_TEST_DIR = Path(__file__).resolve().parent
if str(UNIT_TEST_DIR) not in sys.path:
    sys.path.insert(0, str(UNIT_TEST_DIR))
SCRIPT_DIR = UNIT_TEST_DIR.parent / "scripts"
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from test_support import SKILL_ROOT, temp_dir, write_json
from build_module_visual_report import resolve_feature_tester_name
from project_rules import resolve_asset_path


def table_grid_width(table) -> int:
    return sum(int(grid_col.get(qn("w:w")) or 0) for grid_col in table._tbl.tblGrid.gridCol_lst)


def has_page_break_before_table(table) -> bool:
    node = table._tbl.getprevious()
    while node is not None:
        if node.tag == qn("w:tbl"):
            return False
        if node.tag == qn("w:p") and any(br.get(qn("w:type")) == "page" for br in node.iter(qn("w:br"))):
            return True
        node = node.getprevious()
    return False


def create_beautified_template(path: Path) -> Path:
    document = Document()
    document.add_paragraph("中台 API 自測報告")
    document.add_paragraph("執行證據與對應 API")

    document.add_table(rows=1, cols=1)

    info_table = document.add_table(rows=3, cols=2)
    info_table.rows[0].cells[0].text = "報告類型"
    info_table.rows[0].cells[1].text = "中台 API 自測報告"
    info_table.rows[1].cells[0].text = "適用範圍"
    info_table.rows[1].cells[1].text = "中台 API / Endpoint / Module"
    info_table.rows[2].cells[0].text = "報告日期"
    info_table.rows[2].cells[1].text = "2026/04/23"

    header = document.add_table(rows=4, cols=7)
    header.rows[0].cells[0].text = "功能名稱"
    header.rows[0].cells[1].text = "（API/Endpoint/Method）"
    header.rows[0].cells[3].text = "測試人員"
    header.rows[0].cells[4].text = "姓名"
    header.rows[0].cells[5].text = "測試日期"
    header.rows[0].cells[6].text = "YYYY/MM/DD"
    header.rows[1].cells[0].text = "預期測試結果"
    header.rows[1].cells[2].text = "《API 自測標準》"
    header.rows[2].cells[0].text = "實際測試結果"
    header.rows[2].cells[2].text = "如預期結果 / 依下方檢查項目"
    header.rows[3].cells[0].text = "測試結論"
    header.rows[3].cells[2].text = "☑ 符合需求 ☐ 不符合需求"

    evidence = document.add_table(rows=2, cols=2)
    evidence.rows[0].cells[0].text = "預期輸出/行為"
    evidence.rows[0].cells[1].text = "API"
    evidence.rows[1].cells[0].text = "UnitTest VS執行截圖"
    evidence.rows[1].cells[1].text = "測試API Name"

    document.save(str(path))
    return path


def write_trx(path: Path, test_names: list[str]) -> None:
    body = "\n".join(
        f'''
    <UnitTestResult executionId="{index}" testId="{index}" testName="{name}" computerName="LOCAL" duration="00:00:00.0100000" startTime="2026-04-23T10:00:00.0000000+08:00" endTime="2026-04-23T10:00:00.0100000+08:00" testType="13cdc9d9-ddb5-4fa4-a97d-d965ccfc6d4b" outcome="Passed" testListId="0" relativeResultsDirectory="{index}" />'''
        for index, name in enumerate(test_names, start=1)
    )
    content = f'''<?xml version="1.0" encoding="utf-8"?>
<TestRun id="1" name="sample" xmlns="http://microsoft.com/schemas/VisualStudio/TeamTest/2010">
  <Results>
{body}
  </Results>
</TestRun>
'''
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def build_results_payload(tests: list[str], prefix: str) -> dict[str, object]:
    entries = [
        {
            "testName": name,
            "outcome": "Passed",
            "status": "passed",
            "duration": "00:00:00.0100000",
            "durationSeconds": 0.01,
            "errorMessage": "",
            "errorSummary": "",
            "stackTrace": "",
            "stdOut": "",
            "attachments": [],
        }
        for name in tests
    ]
    summary = {
        "total": len(entries),
        "passed": len(entries),
        "failed": 0,
        "skipped": 0,
        "pending": 0,
    }
    return {
        "trxPath": prefix,
        "unitTestSummary": summary,
        "integrationTestSummary": {
            "total": 0,
            "passed": 0,
            "failed": 0,
            "skipped": 0,
            "pending": 0,
        },
        "sourceResults": {
            "unitTest": {
                "trxPath": prefix,
                "tests": entries,
            },
            "integrationTest": {
                "trxPath": "",
                "tests": [],
            },
        },
    }


def create_module_context(workspace: Path) -> Path:
    repo_root = workspace / "repo"
    context_root = repo_root / ".agent" / "context" / "M.001"
    context_root.mkdir(parents=True, exist_ok=True)
    write_json(repo_root / ".agent" / "config" / "feature-tester-map.json", {"mapping": {"M.001": "Kelly"}})
    write_json(
        context_root / "api-checklist.json",
        {
            "items": [
                {
                    "apiId": "M.001.setting.getuseralias",
                    "apiCategory": "Setting",
                    "apiName": "GetUserAlias",
                },
                {
                    "apiId": "M.001.setting.queryuserloginlog",
                    "apiCategory": "Setting",
                    "apiName": "QueryUserLoginLog",
                },
            ]
        },
    )
    write_json(
        context_root / "execution-state.json",
        {
            "functionCode": "M.001",
        },
    )

    for api_id, display_name, tests, scenarios in [
        (
            "M.001.setting.getuseralias",
            "Setting/GetUserAlias",
            [
                "Sinopac.DawhoEnterprise.Test.UnitTesting.EnterpriseAPI.EnterpriseApiUnit.SettingServiceTests.GetUserAliasAsync_ShouldReturnCachedAlias_WhenRedisHit",
                "Sinopac.DawhoEnterprise.Test.UnitTesting.EnterpriseAPI.EnterpriseApiUnit.SettingControllerTest.GetUserAlias_ShouldReturnOkObjectResult_WhenServiceReturnsResult",
            ],
            ["取得暱稱成功（有數據）", "連接數據庫查詢失敗"],
        ),
        (
            "M.001.setting.queryuserloginlog",
            "Setting/QueryUserLoginLog",
            [
                "Sinopac.DawhoEnterprise.Test.UnitTesting.EnterpriseAPI.EnterpriseApiUnit.SettingServiceTests.QueryUserLoginLogAsync_ShouldReturnNoDataFailure_WhenSqlReturnsEmptyCollection",
                "Sinopac.DawhoEnterprise.Test.UnitTesting.EnterpriseAPI.EnterpriseApiUnit.SettingControllerTest.QueryUserLoginLog_ShouldReturnOkObjectResult_WhenServiceReturnsResult",
            ],
            ["查詢登入紀錄成功", "查無登入紀錄"],
        ),
    ]:
        api_root = context_root / "apis" / api_id
        api_root.mkdir(parents=True, exist_ok=True)
        trx_dir = repo_root / ".agent" / "report-results" / "M.001" / api_id / "unit"
        write_json(
            api_root / "test-evidence.json",
            {
                "apiDisplayName": display_name,
                "trxHints": {
                    "unit": trx_dir.as_posix(),
                    "integration": "",
                },
                "testNames": {
                    "unit": tests,
                    "integration": [],
                },
            },
        )
        write_json(
            api_root / "M.001_API_Spec.json",
            {
                "apiId": api_id,
                "mockExamples": [{"scenario": scenario} for scenario in scenarios],
            },
        )
        write_trx(trx_dir / "EnterpriseAPIUnit.trx", tests)

    return context_root


class BuildModuleVisualReportTests(unittest.TestCase):
    def test_resolve_feature_tester_name_uses_project_plan_map(self) -> None:
        workspace = temp_dir()
        agent_map = workspace / ".agent" / "config" / "feature-tester-map.json"
        agent_map.parent.mkdir(parents=True, exist_ok=True)
        write_json(
            agent_map,
            {
                "mapping": {
                    "B.003": "Kelly",
                    "N.001.001": "Emily",
                    "N.006": "Emily",
                }
            },
        )
        config = {
            "defaultWorkspace": "LOCAL",
            "workspaces": {
                "LOCAL": {
                    "workspaceRoot": workspace.as_posix(),
                    "agentRoot": ".agent",
                }
            },
        }

        with patch("project_rules.load_workspace_config", return_value=(config, workspace / "local-workspaces.json")):
            self.assertEqual(resolve_feature_tester_name("B.003"), "Kelly")
            self.assertEqual(resolve_feature_tester_name("N.001.001"), "Emily")
            self.assertEqual(resolve_feature_tester_name("N.006"), "Emily")

    def test_resolve_feature_tester_name_blocks_when_agent_map_missing_feature(self) -> None:
        workspace = temp_dir()
        agent_map = workspace / ".agent" / "config" / "feature-tester-map.json"
        agent_map.parent.mkdir(parents=True, exist_ok=True)
        write_json(agent_map, {"mapping": {"B.003": "Kelly"}})
        config = {
            "defaultWorkspace": "LOCAL",
            "workspaces": {
                "LOCAL": {
                    "workspaceRoot": workspace.as_posix(),
                    "agentRoot": ".agent",
                }
            },
        }

        with patch("project_rules.load_workspace_config", return_value=(config, workspace / "local-workspaces.json")):
            with self.assertRaises(SystemExit):
                resolve_feature_tester_name("N.001.001")

    def test_build_module_visual_report_generates_docx_and_pngs(self) -> None:
        workspace = temp_dir()
        context_root = create_module_context(workspace)
        template_docx = create_beautified_template(workspace / "beautified-template.docx")
        output_docx = workspace / "module-report.docx"
        summary_json = workspace / "module-report.summary.json"

        subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "build_module_visual_report.py"),
                "--context-root",
                str(context_root),
                "--template-docx",
                str(template_docx),
                "--output-docx",
                str(output_docx),
                "--emit-summary-json",
                str(summary_json),
            ],
            cwd=str(SKILL_ROOT),
            check=True,
        )

        self.assertTrue(output_docx.exists())
        self.assertTrue(summary_json.exists())

        summary = json.loads(summary_json.read_text(encoding="utf-8"))
        self.assertEqual(summary["apiCount"], 2)
        self.assertEqual(summary["testSummary"]["passed"], 4)
        for api in summary["apis"]:
            self.assertTrue(Path(api["imagePath"]).exists())

        document = Document(str(output_docx))
        self.assertEqual(len(document.tables), 4)
        evidence_table = document.tables[3]
        self.assertEqual(len(evidence_table.rows), 3)
        self.assertIn("Setting/GetUserAlias", evidence_table.rows[1].cells[1].text)
        self.assertIn("Setting/QueryUserLoginLog", evidence_table.rows[2].cells[1].text)
        self.assertNotIn("相關測試", evidence_table.rows[1].cells[1].text)
        self.assertIn("通過 2 項 / 不通過 0 項", evidence_table.rows[1].cells[1].text)
        self.assertNotIn("跳過", evidence_table.rows[1].cells[1].text)
        self.assertNotIn("待補", evidence_table.rows[1].cells[1].text)
        self.assertEqual(evidence_table.rows[1].cells[0].vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
        self.assertEqual(evidence_table.rows[1].cells[1].vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
        spacing = evidence_table.rows[1].cells[1].paragraphs[0]._p.pPr.find(qn("w:spacing"))
        self.assertIsNotNone(spacing)
        self.assertEqual(spacing.get(qn("w:afterAutospacing")), "1")
        date_cell_borders = document.tables[1].rows[2].cells[1]._tc.tcPr.find(qn("w:tcBorders"))
        if date_cell_borders is not None:
            date_cell_bottom = date_cell_borders.find(qn("w:bottom"))
            if date_cell_bottom is not None:
                self.assertNotEqual(date_cell_bottom.get(qn("w:val")), "nil")

    def test_build_module_visual_report_uses_report_job_results_when_trx_is_missing(self) -> None:
        workspace = temp_dir()
        context_root = create_module_context(workspace)
        template_docx = create_beautified_template(workspace / "beautified-template.docx")
        output_docx = workspace / "module-report-results-only.docx"
        summary_json = workspace / "module-report-results-only.summary.json"

        for api_id, tests in [
            (
                "M.001.setting.getuseralias",
                [
                    "Sinopac.DawhoEnterprise.Test.UnitTesting.EnterpriseAPI.EnterpriseApiUnit.SettingServiceTests.GetUserAliasAsync_ShouldReturnCachedAlias_WhenRedisHit",
                    "Sinopac.DawhoEnterprise.Test.UnitTesting.EnterpriseAPI.EnterpriseApiUnit.SettingControllerTest.GetUserAlias_ShouldReturnOkObjectResult_WhenServiceReturnsResult",
                ],
            ),
            (
                "M.001.setting.queryuserloginlog",
                [
                    "Sinopac.DawhoEnterprise.Test.UnitTesting.EnterpriseAPI.EnterpriseApiUnit.SettingServiceTests.QueryUserLoginLogAsync_ShouldReturnNoDataFailure_WhenSqlReturnsEmptyCollection",
                    "Sinopac.DawhoEnterprise.Test.UnitTesting.EnterpriseAPI.EnterpriseApiUnit.SettingControllerTest.QueryUserLoginLog_ShouldReturnOkObjectResult_WhenServiceReturnsResult",
                ],
            ),
        ]:
            api_root = context_root / "apis" / api_id
            trx_dir = workspace / "repo" / ".agent" / "report-results" / "M.001" / api_id
            if trx_dir.exists():
                import shutil
                shutil.rmtree(trx_dir)
            (api_root / "ut-report").mkdir(parents=True, exist_ok=True)
            write_json(
                api_root / "ut-report" / "report-job.results.json",
                build_results_payload(tests, f"D:/fake/{api_id}/unit/EnterpriseAPIUnit.trx"),
            )

        subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "build_module_visual_report.py"),
                "--context-root",
                str(context_root),
                "--template-docx",
                str(template_docx),
                "--output-docx",
                str(output_docx),
                "--emit-summary-json",
                str(summary_json),
            ],
            cwd=str(SKILL_ROOT),
            check=True,
        )

        summary = json.loads(summary_json.read_text(encoding="utf-8"))
        self.assertEqual(summary["testSummary"]["passed"], 4)
        self.assertEqual(summary["testSummary"]["total"], 4)

    def test_build_module_visual_report_uses_skill_template_by_default(self) -> None:
        workspace = temp_dir()
        context_root = create_module_context(workspace)
        classified_api_root = context_root / "apis" / "M.001.setting.getuseralias" / "ut-report"
        classified_api_root.mkdir(parents=True, exist_ok=True)
        write_json(
            classified_api_root / "template-classification.json",
            {
                "summary": {
                    "itemCounts": {
                        "applicable": 6,
                        "not_applicable": 10,
                    }
                },
                "sections": [
                    {
                        "sectionId": "UT-05",
                        "title": "UT-05 新增/建立類 API",
                        "applicability": "not_applicable",
                        "items": [
                            {"caseId": "ut-05-001", "applicability": "not_applicable"},
                            {"caseId": "ut-05-002", "applicability": "not_applicable"},
                        ],
                    },
                    {
                        "sectionId": "UT-09",
                        "title": "UT-09 通知/訊息發送類 API",
                        "applicability": "not_applicable",
                        "items": [
                            {"caseId": "ut-09-001", "applicability": "not_applicable"},
                            {"caseId": "ut-09-002", "applicability": "not_applicable"},
                        ],
                    },
                ],
            },
        )
        template_dir = workspace / "repo" / ".agent" / "Template"
        template_dir.mkdir(parents=True, exist_ok=True)
        older_template = create_beautified_template(template_dir / "UT單元測試報告 v2 20260128.docx")
        latest_template = create_beautified_template(template_dir / "UT單元測試報告 v3 20260423.docx")
        output_docx = workspace / "module-report-default-template.docx"
        summary_json = workspace / "module-report-default-template.summary.json"

        subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "build_module_visual_report.py"),
                "--context-root",
                str(context_root),
                "--output-docx",
                str(output_docx),
                "--emit-summary-json",
                str(summary_json),
            ],
            cwd=str(SKILL_ROOT),
            check=True,
        )

        self.assertTrue(output_docx.exists())
        summary = json.loads(summary_json.read_text(encoding="utf-8"))
        expected_template = resolve_asset_path(
            "utReportTemplate",
            fallback=SKILL_ROOT / "assets" / "API_UT自測報告模板_v3.5_20260512.docx",
        )
        self.assertEqual(
            Path(summary["templateDocx"]),
            expected_template.resolve(),
        )
        self.assertIsNone(summary["overallImagePath"])
        document = Document(str(output_docx))
        self.assertEqual(len(document.inline_shapes), 2)
        self.assertEqual(
            sum(1 for paragraph in document.paragraphs if paragraph._p.pPr is not None and paragraph._p.pPr.find(qn("w:pBdr")) is not None),
            0,
        )
        cover_page_borders = document.sections[0]._sectPr.find(qn("w:pgBorders"))
        self.assertIsNotNone(cover_page_borders)
        self.assertEqual(cover_page_borders.get(qn("w:display")), "firstPage")
        self.assertEqual(cover_page_borders.get(qn("w:offsetFrom")), "text")
        self.assertEqual(cover_page_borders.find(qn("w:top")).get(qn("w:sz")), "20")
        self.assertEqual(document.sections[0].left_margin, document.sections[0].right_margin)
        self.assertEqual(document.tables[0].rows[1].cells[0].text.strip(), "1.0")
        self.assertEqual(document.tables[0].rows[1].cells[2].text.strip(), "初版，由 Emily 測試。")
        self.assertEqual(document.tables[0].rows[0].cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(document.tables[0].rows[0].cells[1].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(document.tables[0].rows[1].cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(document.tables[0].rows[1].cells[1].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertIn("通過 106 項 / 不通過 0 項 / 不適用 34 項", document.tables[1].rows[2].cells[2].text)
        self.assertNotIn("如預期結果", document.tables[1].rows[2].cells[2].text)
        self.assertIn("☑ 符合需求", document.tables[1].rows[3].cells[2].text)
        self.assertEqual(sum(1 for paragraph in document.paragraphs if paragraph.text == "測試內容清單"), 0)

        api_headings = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.style.name == "Heading 1" and "Setting/" in paragraph.text
        ]
        self.assertEqual(api_headings, ["二、Setting/GetUserAlias", "三、Setting/QueryUserLoginLog"])
        category_headings = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.style.name == "Heading 2" and paragraph.text.startswith("UT-")
        ]
        self.assertEqual(category_headings.count("UT-01 功能接口範例單元測試"), 2)
        self.assertEqual(category_headings.count("UT-02 API 契約與回應格式"), 2)
        self.assertEqual(category_headings.count("UT-07 DB / SQL 執行環境驗證"), 2)
        self.assertEqual(category_headings.count("UT-15 日誌與可觀測性"), 2)

        all_table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        detail_results = []
        for table in document.tables:
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 3 and cells[0].text.strip().startswith("UT-"):
                    detail_results.append(cells[-1].text.strip())
        self.assertEqual(detail_results.count("通過"), 106)
        self.assertEqual(detail_results.count("不適用"), 34)
        self.assertNotIn("測試 API：", all_table_text)
        self.assertNotIn("{UnitTest VS執行總截圖}", all_table_text)
        self.assertNotIn("依功能接口 API Spec Excel / mockExamples 的每个范例情境", all_table_text)
        self.assertIn("取得暱稱成功（有數據）", all_table_text)
        self.assertIn("連接數據庫查詢失敗", all_table_text)
        self.assertIn("查詢登入紀錄成功", all_table_text)
        self.assertIn("查無登入紀錄", all_table_text)
        self.assertIn("驗證方式：UnitTest", all_table_text)
        self.assertNotIn("驗證重點：依 API Spec 範例驗證回應契約、錯誤碼與主要資料映射", all_table_text)
        self.assertNotIn("適用判定：API 獨有", all_table_text)
        self.assertNotIn("建議驗證方式", all_table_text)
        self.assertNotIn("EnterpriseAPI 配置連線字串可正常開啟", all_table_text)
        self.assertIn("上線前環境驗證條件需保留確認紀錄", all_table_text)
        self.assertIn("測試結果：通過 51 項 / 不通過 0 項 / 不適用 19 項", all_table_text)
        self.assertIn("測試結果：通過 55 項 / 不通過 0 項 / 不適用 15 項", all_table_text)
        self.assertIn("不適用", all_table_text)
        ut01_content_sizes = []
        checklist_sizes = []
        for table in document.tables:
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2 and cells[0].text.strip().startswith("UT-01-"):
                    for paragraph in cells[1].paragraphs:
                        for run in paragraph.runs:
                            if run.text.strip() and run.font.size is not None:
                                ut01_content_sizes.append(run.font.size.pt)
                if len(cells) >= 3 and cells[0].text.strip().startswith("UT-"):
                    for cell in cells[:3]:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text.strip() and run.font.size is not None:
                                    checklist_sizes.append(run.font.size.pt)
        self.assertTrue(ut01_content_sizes)
        self.assertTrue(all(size == 9 for size in ut01_content_sizes))
        self.assertTrue(checklist_sizes)
        self.assertEqual(set(checklist_sizes), {9})
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                self.assertFalse(
                    cells
                    and cells[0].startswith("UT-")
                    and all(cell == cells[0] for cell in cells),
                    msg=f"category row should be rendered as heading, not table row: {cells[0]}",
                )
        self.assertNotEqual(Path(summary["templateDocx"]), latest_template.resolve())
        self.assertNotEqual(Path(summary["templateDocx"]), older_template.resolve())

    def test_build_module_visual_report_honors_explicit_template_override(self) -> None:
        workspace = temp_dir()
        context_root = create_module_context(workspace)
        template_dir = workspace / "repo" / ".agent" / "Template"
        template_dir.mkdir(parents=True, exist_ok=True)
        lower_version_template = create_beautified_template(template_dir / "API_UT 測試報告 20260424.docx")
        higher_version_template = create_beautified_template(template_dir / "API_UT 測試報告 v3 20260424.docx")
        output_docx = workspace / "module-report-version-priority.docx"
        summary_json = workspace / "module-report-version-priority.summary.json"

        subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "build_module_visual_report.py"),
                "--context-root",
                str(context_root),
                "--output-docx",
                str(output_docx),
                "--template-docx",
                str(higher_version_template),
                "--emit-summary-json",
                str(summary_json),
            ],
            cwd=str(SKILL_ROOT),
            check=True,
        )

        self.assertTrue(output_docx.exists())
        summary = json.loads(summary_json.read_text(encoding="utf-8"))
        self.assertEqual(Path(summary["templateDocx"]), higher_version_template.resolve())
        self.assertNotEqual(Path(summary["templateDocx"]), lower_version_template.resolve())


if __name__ == "__main__":
    unittest.main()
