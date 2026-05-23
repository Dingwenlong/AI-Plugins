from __future__ import annotations

import subprocess
import sys
import unittest
from pathlib import Path

from docx import Document

from test_support import (
    ASSETS_DIR,
    SKILL_ROOT,
    build_temp_manifest,
    create_sample_docx,
    read_json,
    temp_dir,
    write_json,
)
from collect_unittest_results import prepare_clean_workspace
from postman_mcp_evidence import render_status_screenshot


def create_api_runtime_artifacts(
    workspace: Path,
    *,
    status_code: int = 200,
    include_request: bool = True,
    include_response: bool = True,
    include_screenshot: bool = True,
    invalid_request_json: bool = False,
    unmasked_authorization: bool = False,
) -> tuple[Path, Path, Path]:
    artifact_root = workspace / ".agent" / "context" / "N.006" / "ut-report" / "postman-mcp" / "Setting.Query" / "success"
    artifact_root.mkdir(parents=True, exist_ok=True)
    request_path = artifact_root / "request.json"
    response_path = artifact_root / "response.json"
    screenshot_path = artifact_root / "status.png"
    request_payload = {
        "method": "GET",
        "url": "https://example.invalid/api/setting/query",
        "headers": {
            "Authorization": "Bearer secret-token" if unmasked_authorization else "***",
            "Cookie": "***",
        },
    }
    response_payload = {
        "statusCode": status_code,
        "headers": {},
        "body": {
            "isSuccess": status_code == 200,
            "responseCode": "0000" if status_code == 200 else "9999",
        },
    }
    if include_request:
        if invalid_request_json:
            request_path.write_text("{ invalid json", encoding="utf-8")
        else:
            write_json(request_path, request_payload)
    if include_response:
        write_json(response_path, response_payload)
    if include_screenshot:
        render_status_screenshot(request_payload, response_payload, screenshot_path, [200])
    return request_path, response_path, screenshot_path


class RunReportJobTests(unittest.TestCase):
    def test_prepare_clean_workspace_rejects_nested_target_root(self) -> None:
        workspace = temp_dir()
        source_root = workspace / "repo-source"
        source_root.mkdir(parents=True, exist_ok=True)

        with self.assertRaises(SystemExit) as raised:
            prepare_clean_workspace(
                workspace,
                {
                    "sourceRoot": str(source_root),
                    "targetRoot": str(source_root / ".agent" / "report-workspace"),
                    "excludeDirNames": [],
                },
            )

        self.assertIn("targetRoot must be outside sourceRoot", str(raised.exception))

    def test_run_report_job_uses_existing_trx(self) -> None:
        workspace = temp_dir()
        docx_path = create_sample_docx(workspace / "sample.docx")
        output_docx = workspace / "sample.report.docx"
        manifest_path = build_temp_manifest(docx_path, output_docx, ASSETS_DIR / "sample-results.trx")
        manifest = read_json(manifest_path)
        manifest["sections"][0]["items"][0]["mode"] = "unit_test"
        manifest["sections"][0]["items"][0]["enabled"] = True
        manifest["sections"][0]["items"][0]["testBindings"]["testNames"] = ["Tests.Sample.Passes"]
        write_json(manifest_path, manifest)

        results_path = workspace / "run.results.json"
        subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "run_report_job.py"),
                str(manifest_path),
                "--results-json",
                str(results_path),
            ],
            cwd=str(SKILL_ROOT),
            check=True,
        )

        results = read_json(results_path)
        self.assertEqual(results["summary"]["passed"], 1)
        passed_case = next(case for case in results["cases"] if case["caseId"] == "ut-01-001")
        self.assertEqual(passed_case["actualResult"], "通過，詳見 UnitTest 結果")
        table_text = "\n".join(
            cell.text for table in Document(str(output_docx)).tables for row in table.rows for cell in row.cells
        )
        self.assertIn("通過", table_text)

    def test_run_report_job_can_execute_command_before_parsing(self) -> None:
        workspace = temp_dir()
        docx_path = create_sample_docx(workspace / "sample.docx")
        output_docx = workspace / "sample.report.docx"
        manifest_path = build_temp_manifest(docx_path, output_docx)
        trx_target = workspace / "TestResults" / "generated.trx"
        trx_target.parent.mkdir(parents=True, exist_ok=True)
        sample_trx = ASSETS_DIR / "sample-results.trx"
        command = (
            f"\"{sys.executable}\" -c \"from pathlib import Path; import shutil; "
            f"target=Path(r'{trx_target}'); target.parent.mkdir(parents=True, exist_ok=True); "
            f"shutil.copyfile(r'{sample_trx}', target)\""
        )

        manifest = read_json(manifest_path)
        manifest["sections"][0]["items"][0]["mode"] = "unit_test"
        manifest["unitTest"]["resultsDir"] = str(trx_target.parent)
        manifest["unitTest"]["command"] = command
        manifest["sections"][0]["items"][0]["enabled"] = True
        manifest["sections"][0]["items"][0]["testBindings"]["testNames"] = ["Tests.Sample.Passes"]
        write_json(manifest_path, manifest)

        results_path = workspace / "run.results.json"
        subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "run_report_job.py"),
                str(manifest_path),
                "--results-json",
                str(results_path),
            ],
            cwd=str(SKILL_ROOT),
            check=True,
        )

        self.assertTrue(trx_target.exists())
        results = read_json(results_path)
        self.assertTrue(results["trxPath"].endswith("generated.trx"))

    def test_run_report_job_blocks_when_binding_is_missing(self) -> None:
        workspace = temp_dir()
        docx_path = create_sample_docx(workspace / "sample.docx")
        output_docx = workspace / "sample.report.docx"
        manifest_path = build_temp_manifest(docx_path, output_docx, ASSETS_DIR / "sample-results.trx")
        manifest = read_json(manifest_path)
        manifest["sections"][0]["items"][0]["mode"] = "unit_test"
        manifest["sections"][0]["items"][0]["enabled"] = True
        write_json(manifest_path, manifest)

        proc = subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "run_report_job.py"),
                str(manifest_path),
            ],
            cwd=str(SKILL_ROOT),
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(proc.returncode, 0)
        self.assertIn("missing explicit testBindings.testNames", proc.stderr + proc.stdout)

    def test_end_to_end_fixture_uses_generated_docx_and_sample_trx(self) -> None:
        workspace = temp_dir()
        demo_docx = create_sample_docx(workspace / "sample-report.docx")

        bootstrap_proc = subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "bootstrap_manifest.py"),
                str(demo_docx),
            ],
            cwd=str(SKILL_ROOT),
            check=True,
        )
        self.assertEqual(bootstrap_proc.returncode, 0)

        manifest_path = demo_docx.with_name(f"{demo_docx.stem}.job.json")
        manifest = read_json(manifest_path)
        manifest["document"]["outputPath"] = str(workspace / "sample-report.docx")
        manifest["sections"][0]["items"][0]["mode"] = "unit_test"
        manifest["unitTest"]["trxPath"] = str(ASSETS_DIR / "sample-results.trx")
        manifest["sections"] = [
            {
                **manifest["sections"][0],
                "items": [manifest["sections"][0]["items"][0]],
            }
        ]
        manifest["sections"][0]["items"][0]["enabled"] = True
        manifest["sections"][0]["items"][0]["testBindings"]["testNames"] = ["Tests.Sample.Passes"]
        write_json(manifest_path, manifest)

        subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "run_report_job.py"),
                str(manifest_path),
            ],
            cwd=str(SKILL_ROOT),
            check=True,
        )

        output_docx = Path(manifest["document"]["outputPath"])
        self.assertTrue(output_docx.exists())
        table_text = "\n".join(
            cell.text for table in Document(str(output_docx)).tables for row in table.rows for cell in row.cells
        )
        self.assertIn("已執行 1 項，其中通過 1 項、失敗 0 項。", table_text)
        self.assertNotIn("自動化證據摘要", table_text)
        self.assertNotIn("結論：", table_text)
        self.assertNotIn("待補項目：", table_text)
        self.assertIn("通過", table_text)

    def test_run_report_job_supports_integration_test_trx(self) -> None:
        workspace = temp_dir()
        docx_path = create_sample_docx(workspace / "sample.docx")
        output_docx = workspace / "sample.report.docx"
        manifest_path = build_temp_manifest(docx_path, output_docx)
        manifest = read_json(manifest_path)
        manifest["integrationTest"]["trxPath"] = str(ASSETS_DIR / "sample-results.trx")
        manifest["sections"][0]["items"][0]["mode"] = "integration_test"
        manifest["sections"][0]["items"][0]["enabled"] = True
        manifest["sections"][0]["items"][0]["testBindings"]["testNames"] = ["Tests.Sample.Passes"]
        write_json(manifest_path, manifest)

        results_path = workspace / "run.results.json"
        subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "run_report_job.py"),
                str(manifest_path),
                "--results-json",
                str(results_path),
            ],
            cwd=str(SKILL_ROOT),
            check=True,
        )

        results = read_json(results_path)
        passed_case = next(case for case in results["cases"] if case["caseId"] == "ut-01-001")
        self.assertEqual(passed_case["sourceKind"], "integrationTest")
        self.assertEqual(passed_case["actualResult"], "通過，詳見 IntegrationTest 結果")

    def test_run_report_job_supports_code_inspection(self) -> None:
        workspace = temp_dir()
        source_root = workspace / "repo"
        controller_path = source_root / "Controllers" / "SettingController.cs"
        controller_path.parent.mkdir(parents=True, exist_ok=True)
        controller_path.write_text(
            "[HttpPost]\n[Route(\"api/setting/update\")]\npublic IActionResult Update() => Ok();\n",
            encoding="utf-8",
        )

        docx_path = create_sample_docx(workspace / "sample.docx")
        output_docx = workspace / "sample.report.docx"
        manifest_path = build_temp_manifest(docx_path, output_docx)
        manifest = read_json(manifest_path)
        manifest["analysisContext"]["repoRoot"] = str(source_root)
        manifest["sections"][0]["items"][0]["mode"] = "code_inspection"
        manifest["sections"][0]["items"][0]["enabled"] = True
        manifest["sections"][0]["items"][0]["codeInspection"] = {
            "ruleId": "http-contract",
            "evidencePaths": [str(controller_path)],
            "mustContainAny": ["[Http", "Route("],
            "mustContainAll": [],
            "mustNotContainAny": [],
            "passActualResult": "通過，已由代碼定位檢查確認 HTTP 介面契約。",
            "pendingActualResult": "待補，尚未定位到充分的 HTTP 介面契約證據。",
            "failActualResult": "失敗，代碼定位檢查發現衝突證據。",
        }
        write_json(manifest_path, manifest)

        results_path = workspace / "run.results.json"
        subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "run_report_job.py"),
                str(manifest_path),
                "--results-json",
                str(results_path),
            ],
            cwd=str(SKILL_ROOT),
            check=True,
        )

        results = read_json(results_path)
        passed_case = next(case for case in results["cases"] if case["caseId"] == "ut-01-001")
        self.assertEqual(passed_case["sourceKind"], "codeInspection")
        self.assertEqual(passed_case["status"], "passed")

    def test_run_report_job_supports_postman_mcp_api_runtime_call_without_trx(self) -> None:
        workspace = temp_dir()
        request_path, response_path, screenshot_path = create_api_runtime_artifacts(workspace)
        docx_path = create_sample_docx(workspace / "sample.docx")
        output_docx = workspace / "sample.report.docx"
        manifest_path = build_temp_manifest(docx_path, output_docx)
        manifest = read_json(manifest_path)
        manifest["analysisContext"]["contextRoot"] = (workspace / ".agent" / "context" / "N.006").as_posix()
        manifest["sections"][0]["items"][0]["mode"] = "api_runtime_call"
        manifest["sections"][0]["items"][0]["enabled"] = True
        manifest["sections"][0]["items"][0]["testBindings"]["testNames"] = []
        manifest["sections"][0]["items"][0]["apiRuntimeCall"] = {
            "requestPath": request_path.as_posix(),
            "responsePath": response_path.as_posix(),
            "screenshotPath": screenshot_path.as_posix(),
            "expectedStatusCodes": [200],
        }
        write_json(manifest_path, manifest)

        results_path = workspace / "run.results.json"
        subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "run_report_job.py"),
                str(manifest_path),
                "--results-json",
                str(results_path),
            ],
            cwd=str(SKILL_ROOT),
            check=True,
        )

        results = read_json(results_path)
        self.assertEqual(results["summary"]["passed"], 1)
        self.assertEqual(results["apiRuntimeCallSummary"]["passed"], 1)
        self.assertEqual(results["sourceResults"]["apiRuntimeCall"]["calls"][0]["statusCode"], 200)
        passed_case = next(case for case in results["cases"] if case["caseId"] == "ut-01-001")
        self.assertEqual(passed_case["sourceKind"], "apiRuntimeCall")
        self.assertEqual(passed_case["status"], "passed")
        self.assertEqual(
            passed_case["attachmentPaths"],
            [request_path.as_posix(), response_path.as_posix(), screenshot_path.as_posix()],
        )
        self.assertIn("Postman MCP / 真实接口调用", passed_case["actualResult"])
        self.assertNotIn("UnitTest", passed_case["actualResult"])
        table_text = "\n".join(
            cell.text for table in Document(str(output_docx)).tables for row in table.rows for cell in row.cells
        )
        self.assertIn("Postman MCP / 真实接口调用", table_text)
        self.assertNotIn("UnitTest", table_text)

    def test_run_report_job_blocks_when_postman_mcp_screenshot_is_missing(self) -> None:
        workspace = temp_dir()
        request_path, response_path, screenshot_path = create_api_runtime_artifacts(
            workspace,
            include_screenshot=False,
        )
        docx_path = create_sample_docx(workspace / "sample.docx")
        output_docx = workspace / "sample.report.docx"
        manifest_path = build_temp_manifest(docx_path, output_docx)
        manifest = read_json(manifest_path)
        manifest["sections"][0]["items"][0]["mode"] = "api_runtime_call"
        manifest["sections"][0]["items"][0]["enabled"] = True
        manifest["sections"][0]["items"][0]["apiRuntimeCall"] = {
            "requestPath": request_path.as_posix(),
            "responsePath": response_path.as_posix(),
            "screenshotPath": screenshot_path.as_posix(),
            "expectedStatusCodes": [200],
        }
        write_json(manifest_path, manifest)

        proc = subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "run_report_job.py"),
                str(manifest_path),
            ],
            cwd=str(SKILL_ROOT),
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(proc.returncode, 0)
        self.assertIn("Postman MCP status screenshot not found", proc.stderr + proc.stdout)

    def test_run_report_job_blocks_when_postman_mcp_request_json_is_invalid(self) -> None:
        workspace = temp_dir()
        request_path, response_path, screenshot_path = create_api_runtime_artifacts(
            workspace,
            invalid_request_json=True,
        )
        docx_path = create_sample_docx(workspace / "sample.docx")
        output_docx = workspace / "sample.report.docx"
        manifest_path = build_temp_manifest(docx_path, output_docx)
        manifest = read_json(manifest_path)
        manifest["sections"][0]["items"][0]["mode"] = "api_runtime_call"
        manifest["sections"][0]["items"][0]["enabled"] = True
        manifest["sections"][0]["items"][0]["apiRuntimeCall"] = {
            "requestPath": request_path.as_posix(),
            "responsePath": response_path.as_posix(),
            "screenshotPath": screenshot_path.as_posix(),
            "expectedStatusCodes": [200],
        }
        write_json(manifest_path, manifest)

        proc = subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "run_report_job.py"),
                str(manifest_path),
            ],
            cwd=str(SKILL_ROOT),
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(proc.returncode, 0)
        self.assertIn("Postman MCP request JSON is invalid", proc.stderr + proc.stdout)

    def test_run_report_job_blocks_when_postman_mcp_response_json_is_missing(self) -> None:
        workspace = temp_dir()
        request_path, response_path, screenshot_path = create_api_runtime_artifacts(
            workspace,
            include_response=False,
        )
        docx_path = create_sample_docx(workspace / "sample.docx")
        output_docx = workspace / "sample.report.docx"
        manifest_path = build_temp_manifest(docx_path, output_docx)
        manifest = read_json(manifest_path)
        manifest["sections"][0]["items"][0]["mode"] = "api_runtime_call"
        manifest["sections"][0]["items"][0]["enabled"] = True
        manifest["sections"][0]["items"][0]["apiRuntimeCall"] = {
            "requestPath": request_path.as_posix(),
            "responsePath": response_path.as_posix(),
            "screenshotPath": screenshot_path.as_posix(),
            "expectedStatusCodes": [200],
        }
        write_json(manifest_path, manifest)

        proc = subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "run_report_job.py"),
                str(manifest_path),
            ],
            cwd=str(SKILL_ROOT),
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(proc.returncode, 0)
        self.assertIn("Postman MCP response JSON not found", proc.stderr + proc.stdout)

    def test_run_report_job_blocks_when_postman_mcp_artifacts_contain_unmasked_secrets(self) -> None:
        workspace = temp_dir()
        request_path, response_path, screenshot_path = create_api_runtime_artifacts(
            workspace,
            unmasked_authorization=True,
        )
        docx_path = create_sample_docx(workspace / "sample.docx")
        output_docx = workspace / "sample.report.docx"
        manifest_path = build_temp_manifest(docx_path, output_docx)
        manifest = read_json(manifest_path)
        manifest["sections"][0]["items"][0]["mode"] = "api_runtime_call"
        manifest["sections"][0]["items"][0]["enabled"] = True
        manifest["sections"][0]["items"][0]["apiRuntimeCall"] = {
            "requestPath": request_path.as_posix(),
            "responsePath": response_path.as_posix(),
            "screenshotPath": screenshot_path.as_posix(),
            "expectedStatusCodes": [200],
        }
        write_json(manifest_path, manifest)

        proc = subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "run_report_job.py"),
                str(manifest_path),
            ],
            cwd=str(SKILL_ROOT),
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(proc.returncode, 0)
        self.assertIn("unmasked sensitive header", proc.stderr + proc.stdout)

    def test_run_report_job_marks_postman_mcp_status_mismatch_failed(self) -> None:
        workspace = temp_dir()
        request_path, response_path, screenshot_path = create_api_runtime_artifacts(workspace, status_code=500)
        docx_path = create_sample_docx(workspace / "sample.docx")
        output_docx = workspace / "sample.report.docx"
        manifest_path = build_temp_manifest(docx_path, output_docx)
        manifest = read_json(manifest_path)
        manifest["sections"][0]["items"][0]["mode"] = "api_runtime_call"
        manifest["sections"][0]["items"][0]["enabled"] = True
        manifest["sections"][0]["items"][0]["apiRuntimeCall"] = {
            "requestPath": request_path.as_posix(),
            "responsePath": response_path.as_posix(),
            "screenshotPath": screenshot_path.as_posix(),
            "expectedStatusCodes": [200],
        }
        write_json(manifest_path, manifest)

        results_path = workspace / "run.results.json"
        subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "run_report_job.py"),
                str(manifest_path),
                "--results-json",
                str(results_path),
            ],
            cwd=str(SKILL_ROOT),
            check=True,
        )

        results = read_json(results_path)
        self.assertEqual(results["summary"]["failed"], 1)
        self.assertEqual(results["apiRuntimeCallSummary"]["failed"], 1)
        failed_case = next(case for case in results["cases"] if case["caseId"] == "ut-01-001")
        self.assertEqual(failed_case["status"], "failed")
        self.assertIn("HTTP 500", failed_case["actualResult"])
        self.assertIn("Postman MCP / 真实接口调用", failed_case["actualResult"])

    def test_run_report_job_can_execute_integration_test_in_clean_workspace(self) -> None:
        workspace = temp_dir()
        repo_source = workspace / "repo-source"
        repo_target = workspace / "repo-target"
        repo_source.mkdir(parents=True, exist_ok=True)
        (repo_source / "bin").mkdir(parents=True, exist_ok=True)
        (repo_source / "bin" / "ignored.txt").write_text("ignore", encoding="utf-8")
        sample_trx = ASSETS_DIR / "sample-results.trx"
        command = (
            f"\"{sys.executable}\" -c \"from pathlib import Path; import shutil; "
            f"target=Path(r'{{workspaceRoot}}') / '.agent' / 'it-results' / 'integration.trx'; "
            f"target.parent.mkdir(parents=True, exist_ok=True); "
            f"shutil.copyfile(r'{sample_trx}', target)\""
        )

        docx_path = create_sample_docx(workspace / "sample.docx")
        output_docx = workspace / "sample.report.docx"
        manifest_path = build_temp_manifest(docx_path, output_docx)
        manifest = read_json(manifest_path)
        manifest["sections"][0]["items"][0]["mode"] = "integration_test"
        manifest["sections"][0]["items"][0]["enabled"] = True
        manifest["sections"][0]["items"][0]["testBindings"]["testNames"] = ["Tests.Sample.Passes"]
        manifest["integrationTest"]["command"] = command
        manifest["integrationTest"]["workingDirectory"] = "{workspaceRoot}"
        manifest["integrationTest"]["resultsDir"] = "{workspaceRoot}/.agent/it-results"
        manifest["integrationTest"]["failIfTrxMissing"] = True
        manifest["integrationTest"]["cleanWorkspace"] = {
            "enabled": True,
            "sourceRoot": str(repo_source),
            "targetRoot": str(repo_target),
            "excludeDirNames": ["bin"],
        }
        write_json(manifest_path, manifest)

        results_path = workspace / "run.results.json"
        subprocess.run(
            [
                sys.executable,
                str(SKILL_ROOT / "scripts" / "run_report_job.py"),
                str(manifest_path),
                "--results-json",
                str(results_path),
            ],
            cwd=str(SKILL_ROOT),
            check=True,
        )

        self.assertTrue((repo_target / ".agent" / "it-results" / "integration.trx").exists())
        self.assertFalse((repo_target / "bin").exists())
        results = read_json(results_path)
        self.assertEqual(results["integrationTestSummary"]["passed"], 1)


if __name__ == "__main__":
    unittest.main()
